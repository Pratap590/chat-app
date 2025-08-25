#!/usr/bin/env python3
"""
FastAPI GUI Application for Multi-Agent Chatbot System
Provides web interface for all agents with file upload, download, and interactive features
"""

import os
import json
import tempfile
import shutil
import sqlite3
import secrets
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

# Set up logger
logger = logging.getLogger(__name__)

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
import uvicorn
import pandas as pd

# Import our main chatbot system
from main import (
    chat_with_agent, create_tenant, get_tenant_config,
    ingest_documents_from_dir, get_system_stats, get_tool_stats,
    FORM_GENERATOR, _json_to_professional_form,
    node_form_gen, node_doc_qa, node_api_exec, node_analytics, node_escalate,
    CURRENT_TENANT_ID, set_current_tenant, create_session, MessagesState
)

# Initialize FastAPI app
app = FastAPI(
    title="Multi-Agent Chatbot System",
    description="Web interface for document RAG, form generation, API execution, and analytics",
    version="1.0.0"
)

# Create directories for static files and uploads
static_dir = Path("static")
templates_dir = Path("templates")
uploads_dir = Path("uploads")
downloads_dir = Path("downloads")

for directory in [static_dir, templates_dir, uploads_dir, downloads_dir]:
    directory.mkdir(exist_ok=True)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/downloads", StaticFiles(directory="downloads"), name="downloads")

# Templates
templates = Jinja2Templates(directory="templates")

# Pydantic models for API requests
class ChatRequest(BaseModel):
    message: str
    agent_type: str
    tenant_id: str = "default"

class FormGenerationRequest(BaseModel):
    description: str
    format: str = "pdf"  # pdf or docx
    tenant_id: str = "default"

class TenantRequest(BaseModel):
    tenant_id: str
    name: str
    permissions: List[str] = ["read_documents", "use_tools", "generate_forms"]

# Global variables for session management
current_sessions = {}

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Main dashboard page"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/admin", response_class=HTMLResponse)
async def admin_dashboard(request: Request):
    """Admin dashboard page"""
    return templates.TemplateResponse("admin_dashboard.html", {"request": request})

@app.get("/console", response_class=HTMLResponse)
async def agent_console(request: Request):
    """Agent console page for real-time chat monitoring"""
    return templates.TemplateResponse("agent_console.html", {"request": request})

@app.get("/api/agents")
async def get_agents():
    """Get available agents and their descriptions"""
    agents = {
        "doc_qa": {
            "name": "Document Q&A",
            "description": "Upload documents and ask questions based on their content",
            "icon": "📄",
            "features": ["Document Upload", "RAG Search", "Contextual Answers"]
        },
        "form_gen": {
            "name": "Form Generator", 
            "description": "Generate professional forms with PDF/DOC export",
            "icon": "📝",
            "features": ["PDF Export", "DOCX Export", "Professional Templates"]
        },
        "api_exec": {
            "name": "API Executor",
            "description": "Execute API calls and external tool operations",
            "icon": "🔧",
            "features": ["Weather API", "Web Search", "Custom Tools"]
        },
        "analytics": {
            "name": "Analytics",
            "description": "System analytics and data insights",
            "icon": "📊", 
            "features": ["Usage Statistics", "Performance Metrics", "Reports"]
        },
        "escalate": {
            "name": "Escalation",
            "description": "Human support and ticket management",
            "icon": "🆘",
            "features": ["Ticket Creation", "Human Handoff", "Support Queue"]
        }
    }
    return agents

@app.post("/api/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    tenant_id: str = Form("default"),
    user_id: str = Form(None)
):
    """Enhanced upload and process documents for RAG with multiple document support"""
    try:
        # Validate file type
        allowed_extensions = {'.pdf', '.docx', '.txt', '.md', '.csv', '.json'}
        file_extension = Path(file.filename).suffix.lower()

        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )

        # Set current tenant context
        set_current_tenant(tenant_id)

        # Create tenant-specific upload directory
        tenant_upload_dir = uploads_dir / tenant_id
        tenant_upload_dir.mkdir(exist_ok=True)

        # Save uploaded file with unique name to prevent conflicts
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = tenant_upload_dir / safe_filename

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Ensure tenant exists in registry
        from main import _tenant_registry, create_tenant
        if tenant_id not in _tenant_registry:
            create_tenant(tenant_id, f"Tenant {tenant_id}", ["read_documents", "use_tools", "generate_forms"])

        # Process single document with enhanced metadata
        from main import ingest_single_document
        result = ingest_single_document(tenant_id, str(file_path), user_id)

        if not result["success"]:
            # Clean up file if processing failed
            if file_path.exists():
                file_path.unlink()
            raise Exception(result["message"])

        return {
            "success": True,
            "message": result["message"],
            "document_id": result["document_id"],
            "filename": file.filename,
            "file_path": str(file_path),
            "tenant_id": tenant_id,
            "indexed": True,
            "chunks": result.get("chunks", 0),
            "duplicate": result.get("duplicate", False)
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/api/upload-multiple-documents")
async def upload_multiple_documents(
    files: List[UploadFile] = File(...),
    tenant_id: str = Form("default"),
    user_id: str = Form(None)
):
    """Upload and process multiple documents simultaneously"""
    try:
        # Validate all files first
        allowed_extensions = {'.pdf', '.docx', '.txt', '.md', '.csv', '.json'}
        for file in files:
            file_extension = Path(file.filename).suffix.lower()
            if file_extension not in allowed_extensions:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type '{file.filename}'. Allowed: {', '.join(allowed_extensions)}"
                )

        # Set current tenant context
        set_current_tenant(tenant_id)

        # Create tenant-specific upload directory
        tenant_upload_dir = uploads_dir / tenant_id
        tenant_upload_dir.mkdir(exist_ok=True)

        # Save all files
        file_paths = []
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        for i, file in enumerate(files):
            safe_filename = f"{timestamp}_{i:03d}_{file.filename}"
            file_path = tenant_upload_dir / safe_filename

            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            file_paths.append(str(file_path))

        # Ensure tenant exists in registry
        from main import _tenant_registry, create_tenant
        if tenant_id not in _tenant_registry:
            create_tenant(tenant_id, f"Tenant {tenant_id}", ["read_documents", "use_tools", "generate_forms"])

        # Process multiple documents
        from main import ingest_multiple_documents
        result = ingest_multiple_documents(tenant_id, file_paths, user_id)

        return {
            "success": True,
            "message": f"Processed {result['total_files']} files",
            "total_files": result["total_files"],
            "successful": result["successful"],
            "failed": result["failed"],
            "duplicates": result["duplicates"],
            "results": result["results"],
            "tenant_id": tenant_id
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Multiple upload failed: {str(e)}")

@app.get("/api/documents/{tenant_id}")
async def get_documents(tenant_id: str):
    """Get all documents for a tenant"""
    try:
        from main import document_storage
        documents = document_storage.get_documents_by_tenant(tenant_id)

        return {
            "success": True,
            "documents": [
                {
                    "document_id": doc.document_id,
                    "filename": doc.filename,
                    "original_name": doc.original_name,
                    "file_size": doc.file_size,
                    "file_type": doc.file_type,
                    "upload_timestamp": doc.upload_timestamp,
                    "chunk_count": doc.chunk_count,
                    "indexed": doc.indexed,
                    "tags": doc.tags
                }
                for doc in documents
            ],
            "total": len(documents)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get documents: {str(e)}")

@app.delete("/api/documents/{tenant_id}/{document_id}")
async def delete_document(tenant_id: str, document_id: str):
    """Delete a specific document"""
    try:
        from main import document_storage
        documents = document_storage.get_documents_by_tenant(tenant_id)

        # Find the document
        doc_to_delete = None
        for doc in documents:
            if doc.document_id == document_id:
                doc_to_delete = doc
                break

        if not doc_to_delete:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete file from filesystem
        try:
            if os.path.exists(doc_to_delete.file_path):
                os.remove(doc_to_delete.file_path)
                logger.info(f"Deleted file: {doc_to_delete.file_path}")
        except Exception as e:
            logger.warning(f"Failed to delete file {doc_to_delete.file_path}: {e}")

        # Remove from database
        try:
            conn = sqlite3.connect(document_storage.db_path)
            cursor = conn.cursor()
            cursor.execute("DELETE FROM documents WHERE document_id = ?", (document_id,))
            conn.commit()
            conn.close()
            logger.info(f"Removed document {document_id} from database")
        except Exception as e:
            logger.error(f"Failed to remove document from database: {e}")
            raise Exception(f"Database deletion failed: {e}")

        # Remove from vector store by rebuilding index
        vector_rebuild_success = False
        try:
            from main import ingest_documents_from_dir
            # Get tenant's upload directory
            tenant_upload_dir = Path("uploads") / tenant_id
            if tenant_upload_dir.exists():
                # Rebuild the vector store index for this tenant
                ingest_documents_from_dir(str(tenant_upload_dir), tenant_id)
                vector_rebuild_success = True
                logger.info(f"Rebuilt vector store index for tenant {tenant_id} after document deletion")
        except Exception as e:
            logger.warning(f"Failed to rebuild vector store after deletion: {e}")
            # Don't fail the entire operation if vector store rebuild fails

        success_message = f"Document '{doc_to_delete.filename}' deleted successfully"
        if vector_rebuild_success:
            success_message += " and index updated"
        else:
            success_message += " (index rebuild failed but document removed)"

        return {
            "success": True,
            "message": success_message
        }
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Document deletion failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@app.get("/api/chat-history/{tenant_id}/{session_id}")
async def get_chat_history(tenant_id: str, session_id: str, limit: int = 50):
    """Get chat history for a session"""
    try:
        from main import document_storage
        messages = document_storage.get_chat_history(session_id, limit)

        return {
            "success": True,
            "messages": [
                {
                    "message_id": msg.message_id,
                    "role": msg.role,
                    "content": msg.content,
                    "timestamp": msg.timestamp,
                    "agent_type": msg.agent_type,
                    "document_references": msg.document_references
                }
                for msg in messages
            ],
            "total": len(messages)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get chat history: {str(e)}")

@app.post("/api/chat")
async def chat_endpoint(request: ChatRequest):
    """Main chat endpoint for all agents"""
    try:
        # Set current tenant context
        set_current_tenant(request.tenant_id)

        # Ensure tenant exists in registry
        from main import _tenant_registry, create_tenant
        if request.tenant_id not in _tenant_registry:
            create_tenant(request.tenant_id, f"Tenant {request.tenant_id}", ["read_documents", "use_tools", "generate_forms"])

        # For doc_qa agent, check if documents are indexed
        if request.agent_type == "doc_qa":
            from main import get_retriever_for_tenant
            retriever = get_retriever_for_tenant(request.tenant_id)
            if retriever is None:
                return {
                    "success": True,
                    "response": f"No documents indexed for tenant '{request.tenant_id}'. Please upload documents first using the upload area above.",
                    "agent": request.agent_type,
                    "timestamp": datetime.now().isoformat()
                }

        # Route to appropriate agent
        if request.agent_type == "doc_qa":
            # For document Q&A, call the function directly instead of using chat_with_agent
            
            # Set up proper session context
            set_current_tenant(request.tenant_id)
            session = create_session(request.tenant_id)
            
            # Create proper message state with session context
            state = MessagesState(messages=[("user", request.message)])
            
            # Call the document Q&A function
            result = node_doc_qa(state)
            
            # Extract response from the result
            if isinstance(result, dict) and "messages" in result:
                messages = result["messages"]
                if messages:
                    last_message = messages[-1]
                    if hasattr(last_message, 'content'):
                        response = last_message.content
                    elif isinstance(last_message, tuple) and len(last_message) > 1:
                        response = last_message[1]
                    else:
                        response = str(last_message)
                else:
                    response = "No response generated"
            else:
                response = str(result)
        elif request.agent_type == "form_gen":
            response = await handle_form_generation(request.message, request.tenant_id)
        elif request.agent_type == "api_exec":
            response = chat_with_agent(request.message, request.tenant_id)
        elif request.agent_type == "analytics":
            response = chat_with_agent(request.message, request.tenant_id)
        elif request.agent_type == "escalate":
            response = chat_with_agent(request.message, request.tenant_id)
        else:
            response = chat_with_agent(request.message, request.tenant_id)

        # Handle new form generation response with preview
        if isinstance(response, dict) and response.get("form_generated"):
            # Extract response text from messages
            response_text = ""
            messages = response.get("messages", [])
            if messages and isinstance(messages[0], tuple) and len(messages[0]) > 1:
                response_text = messages[0][1]
            elif messages:
                response_text = str(messages[0])
            
            result = {
                "success": True,
                "response": response_text,
                "form_generated": True,
                "preview": response.get("preview"),
                "content_type": response.get("content_type"),
                "file_format": response.get("file_format"),
                "file_size": response.get("file_size", 0),
                "filename": response.get("filename"),
                "agent": request.agent_type,
                "timestamp": datetime.now().isoformat()
            }
            
            # Handle HTML content vs binary file content
            if response.get("html_content"):
                result["html_content"] = response["html_content"]
                result["interactive"] = response.get("interactive", False)
            elif response.get("file_content"):
                result["file_content"] = response["file_content"]
                
            return result

        # Handle legacy auto-download response for form generation
        elif isinstance(response, dict) and response.get("auto_download"):
            # Store file content temporarily for download
            download_id = secrets.token_hex(16)
            current_sessions[download_id] = {
                "file_content": response["file_content"],
                "filename": response["filename"],
                "content_type": response["content_type"],
                "timestamp": datetime.now()
            }

            return {
                "success": True,
                "response": response["response"],
                "auto_download": True,
                "download_id": download_id,
                "filename": response["filename"],
                "agent": request.agent_type,
                "timestamp": datetime.now().isoformat()
            }

        return {
            "success": True,
            "response": response,
            "agent": request.agent_type,
            "timestamp": datetime.now().isoformat()
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

async def handle_form_generation(message: str, tenant_id: str) -> Dict[str, Any]:
    """Handle form generation with automatic download"""
    try:
        # Set tenant context
        from main import set_current_tenant, node_form_gen, MessagesState
        set_current_tenant(tenant_id)

        # Detect file format preference from the message
        file_format = "pdf"  # default
        if any(word in message.lower() for word in ["html", "web", "format: html"]):
            file_format = "html"
        elif any(word in message.lower() for word in ["doc", "docx", "word"]):
            file_format = "docx"
        
        # Don't override the format if it's already specified in the message
        if "(format:" not in message.lower():
            message = f"{message} (format: {file_format})"

        # Generate form directly using the form generation node
        state = MessagesState(messages=[("user", message)])
        form_result = node_form_gen(state)

        # Check if form was generated successfully with file content
        if isinstance(form_result, dict) and form_result.get("form_generated"):
            # Extract the response message
            messages = form_result.get("messages", [])
            response_text = ""
            if messages:
                last_message = messages[-1]
                if hasattr(last_message, 'content'):
                    response_text = last_message.content
                elif isinstance(last_message, tuple) and len(last_message) > 1:
                    response_text = last_message[1]
                else:
                    response_text = str(last_message)

            # Create a clean preview message with the actual form content
            preview_message = f"""
            <div class="form-preview-container">
                <div class="form-preview-header">
                    <h3>📝 Form Generated Successfully!</h3>
                    <div class="form-stats">
                        <span class="stat"><i class="fas fa-file-{form_result.get('file_format', 'pdf')}"></i> {form_result.get('file_format', 'PDF').upper()}</span>
                        <span class="stat"><i class="fas fa-download"></i> {form_result.get('file_size', 0)} bytes</span>
                    </div>
                </div>
                <div class="form-preview-content">
                    {response_text}
                </div>
            </div>
            """

            # Return the appropriate response based on format
            if form_result.get("html_content"):
                # For HTML format, return the HTML content directly
                return {
                    "response": response_text,
                    "form_generated": True,
                    "html_content": form_result.get("html_content"),
                    "filename": form_result.get("filename"),
                    "content_type": form_result.get("content_type"),
                    "file_format": form_result.get("file_format"),
                    "file_size": form_result.get("file_size", 0),
                    "preview": form_result.get("preview"),
                    "interactive": form_result.get("interactive", False),
                    "success": True
                }
            else:
                # For PDF/DOCX format, return with auto-download
                return {
                    "response": preview_message,
                    "form_generated": True,
                    "auto_download": True,
                    "file_content": form_result.get("file_content"),
                    "file_format": form_result.get("file_format", file_format),
                    "filename": form_result.get("filename"),
                    "content_type": form_result.get("content_type"),
                    "file_size": form_result.get("file_size", 0),
                    "preview": form_result.get("preview"),
                    "success": True
                }

        # If no file was generated, extract the text response
        messages = form_result.get("messages", [])
        response_text = ""
        if messages:
            last_message = messages[-1]
            if hasattr(last_message, 'content'):
                response_text = last_message.content
            elif isinstance(last_message, tuple) and len(last_message) > 1:
                response_text = last_message[1]
            else:
                response_text = str(last_message)

        return {
            "response": response_text,
            "success": True
        }

    except Exception as e:
        return {
            "response": f"❌ Form generation error: {str(e)}",
            "success": False
        }

@app.post("/api/generate-form")
async def generate_form_endpoint(request: FormGenerationRequest):
    """Dedicated form generation endpoint"""
    try:
        result = await handle_form_generation(request.description, request.tenant_id)
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Form generation failed: {str(e)}")

@app.get("/api/download/{download_id}")
async def download_generated_file(download_id: str):
    """Download auto-generated files using download ID"""
    if download_id not in current_sessions:
        raise HTTPException(status_code=404, detail="Download not found or expired")

    session_data = current_sessions[download_id]

    # Clean up the session after download
    del current_sessions[download_id]

    # Return the file content as a response
    from fastapi.responses import Response
    return Response(
        content=session_data["file_content"],
        media_type=session_data["content_type"],
        headers={
            "Content-Disposition": f"attachment; filename={session_data['filename']}"
        }
    )

@app.get("/downloads/{filename}")
async def download_file(filename: str):
    """Download generated files from downloads folder (legacy)"""
    file_path = downloads_dir / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type='application/octet-stream'
    )

@app.post("/api/convert-form")
async def convert_form(request: Request):
    """Convert HTML form to PDF or DOCX format with filled data"""
    try:
        data = await request.json()
        html_content = data.get('htmlContent', '')
        filename = data.get('filename', 'form')
        format_type = data.get('format', 'pdf').lower()
        
        logger.info(f"Convert form request: format={format_type}, filename={filename}, content_length={len(html_content)}")
        
        if not html_content:
            logger.error("No HTML content provided for conversion")
            raise HTTPException(status_code=400, detail="HTML content is required")
        
        if format_type not in ['pdf', 'docx']:
            logger.error(f"Invalid format requested: {format_type}")
            raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")
        
        logger.info(f"Converting HTML form to {format_type.upper()} format")
        
        # Import required modules for conversion
        from main import FORM_GENERATOR
        import tempfile
        import os
        
        try:
            # Parse HTML to extract form structure and data
            form_structure = parse_html_form(html_content)
            logger.info(f"Parsed form structure: title='{form_structure['title']}', company='{form_structure['company_name']}', fields={len(form_structure['fields'])}")
            
            # Create a temporary HTML file for conversion
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_file:
                temp_file.write(html_content)
                temp_html_path = temp_file.name
            
            logger.info(f"Created temporary HTML file: {temp_html_path}")
            
            try:
                # Generate the requested format using HTML to PDF/DOCX conversion
                if format_type == 'pdf':
                    logger.info("Starting PDF conversion...")
                    output_path = convert_html_to_pdf(temp_html_path, filename)
                    content_type = "application/pdf"
                else:  # docx
                    logger.info("Starting DOCX conversion...")
                    output_path = convert_html_to_docx(temp_html_path, filename, form_structure)
                    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                
                logger.info(f"Conversion completed, output file: {output_path}")
                
                # Check if output file exists and has content
                if not os.path.exists(output_path):
                    raise Exception(f"Output file not created: {output_path}")
                
                file_size = os.path.getsize(output_path)
                if file_size == 0:
                    raise Exception(f"Output file is empty: {output_path}")
                
                logger.info(f"Output file size: {file_size} bytes")
                
                # Read the generated file
                with open(output_path, 'rb') as f:
                    file_content = f.read()
                
                logger.info(f"Read {len(file_content)} bytes from output file")
                
                # Clean up temporary files
                try:
                    os.unlink(temp_html_path)
                    os.unlink(output_path)
                    logger.info("Cleaned up temporary files")
                except Exception as cleanup_error:
                    logger.warning(f"Error cleaning up temporary files: {cleanup_error}")
                
                # Return the file as a download
                from fastapi.responses import Response
                response = Response(
                    content=file_content,
                    media_type=content_type,
                    headers={"Content-Disposition": f"attachment; filename={filename}_completed.{format_type}"}
                )
                
                logger.info(f"Successfully returning {format_type.upper()} file download")
                return response
                
            except Exception as conversion_error:
                logger.error(f"Conversion error: {conversion_error}")
                # Clean up temp file on error
                if os.path.exists(temp_html_path):
                    try:
                        os.unlink(temp_html_path)
                    except:
                        pass
                raise conversion_error
            
        except Exception as conversion_error:
            logger.error(f"Form conversion error: {conversion_error}")
            # Fallback to simple text-based conversion
            logger.info("Attempting fallback conversion...")
            return await fallback_form_conversion(html_content, filename, format_type)
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Convert form API error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

# Helper function to parse HTML form structure
def parse_html_form(html_content):
    """Parse HTML to extract form structure and filled data using regex"""
    try:
        import re
        
        # Extract title from h1 tags
        title_match = re.search(r'<h1[^>]*>([^<]+)</h1>', html_content, re.IGNORECASE)
        title = title_match.group(1).strip() if title_match else "Form"
        
        # Extract company name - try multiple patterns
        company_patterns = [
            r'<span[^>]*color[^>]*>([^<]+)</span>',  # Blue colored spans
            r'<h2[^>]*>([^<]+)</h2>',  # H2 headings
            r'<div[^>]*class="[^"]*company[^>]*>([^<]+)</div>',  # Company class divs
            r'<p[^>]*>\s*([^<]*(?:Company|Corp|Inc|LLC)[^<]*)\s*</p>'  # Company in paragraphs
        ]
        
        company_name = "Your Company"
        for pattern in company_patterns:
            match = re.search(pattern, html_content, re.IGNORECASE)
            if match:
                potential_company = match.group(1).strip()
                # Filter out common non-company texts
                if (potential_company and 
                    potential_company != title and 
                    len(potential_company) > 2 and
                    not potential_company.lower().startswith('form') and
                    not potential_company.lower().startswith('submit')):
                    company_name = potential_company
                    break
        
        # Extract logo if present
        logo_data = None
        logo_pattern = r'<img[^>]*src=["\']([^"\'>]*)["\'][^>]*alt=["\'][^"\'>]*[Ll]ogo[^"\'>]*["\'][^>]*>'
        logo_match = re.search(logo_pattern, html_content, re.IGNORECASE)
        if logo_match:
            logo_data = logo_match.group(1)
            logger.info(f"Found logo in HTML: {logo_data[:50]}...")
        
        # Extract form fields and their values
        form_fields = []
        
        # Find input fields with values (improved regex)
        input_patterns = [
            r'<input[^>]*name=["\']([^"\'>]+)["\'][^>]*value=["\']([^"\'>]*)["\'][^>]*>',
            r'<input[^>]*value=["\']([^"\'>]*)["\'][^>]*name=["\']([^"\'>]+)["\'][^>]*>'
        ]
        
        for pattern in input_patterns:
            matches = re.findall(pattern, html_content, re.IGNORECASE)
            for match in matches:
                if len(match) == 2:
                    name, value = match if 'name=' in pattern[:30] else (match[1], match[0])
                    if value and value.strip() and name and name.strip():
                        # Clean up the field name
                        clean_name = name.replace('_', ' ').replace('-', ' ').title()
                        form_fields.append({
                            'label': clean_name,
                            'value': value.strip()
                        })
        
        # Find textarea fields with content
        textarea_pattern = r'<textarea[^>]*name=["\']([^"\'>]+)["\'][^>]*>([^<]*)</textarea>'
        textarea_matches = re.findall(textarea_pattern, html_content, re.IGNORECASE)
        
        for name, value in textarea_matches:
            if value and value.strip():
                clean_name = name.replace('_', ' ').replace('-', ' ').title()
                form_fields.append({
                    'label': clean_name,
                    'value': value.strip()
                })
        
        # Find selected options in select elements (Enhanced)
        # Pattern 1: Selected options
        select_pattern = r'<select[^>]*name=["\']([^"\'>]+)["\'][^>]*>.*?<option[^>]*selected[^>]*(?:value=["\']([^"\'>]*)["\'])?[^>]*>([^<]*)</option>.*?</select>'
        select_matches = re.findall(select_pattern, html_content, re.IGNORECASE | re.DOTALL)
        
        for name, value, text in select_matches:
            if (value and value.strip()) or (text and text.strip()):
                clean_name = name.replace('_', ' ').replace('-', ' ').title()
                display_value = value.strip() if value.strip() else text.strip()
                form_fields.append({
                    'label': clean_name,
                    'value': display_value
                })
        
        # Pattern 2: Find all select elements and their selected values
        select_element_pattern = r'<select[^>]*name=["\']([^"\'>]+)["\'][^>]*>(.*?)</select>'
        select_elements = re.findall(select_element_pattern, html_content, re.IGNORECASE | re.DOTALL)
        
        for name, select_content in select_elements:
            # Look for selected option within this select
            selected_option = re.search(r'<option[^>]*selected[^>]*(?:value=["\']([^"\'>]*)["\'])?[^>]*>([^<]*)</option>', select_content, re.IGNORECASE)
            if selected_option:
                value, text = selected_option.groups()
                if (value and value.strip()) or (text and text.strip()):
                    clean_name = name.replace('_', ' ').replace('-', ' ').title()
                    display_value = value.strip() if value and value.strip() else text.strip()
                    # Avoid duplicates
                    if not any(field['label'] == clean_name for field in form_fields):
                        form_fields.append({
                            'label': clean_name,
                            'value': display_value
                        })
        
        # Find checked checkboxes and radio buttons (Enhanced)
        checked_patterns = [
            r'<input[^>]*type=["\'](?:checkbox|radio)["\'][^>]*name=["\']([^"\'>]+)["\'][^>]*value=["\']([^"\'>]*)["\'][^>]*checked[^>]*>',
            r'<input[^>]*name=["\']([^"\'>]+)["\'][^>]*type=["\'](?:checkbox|radio)["\'][^>]*value=["\']([^"\'>]*)["\'][^>]*checked[^>]*>',
            r'<input[^>]*checked[^>]*name=["\']([^"\'>]+)["\'][^>]*type=["\'](?:checkbox|radio)["\'][^>]*value=["\']([^"\'>]*)["\'][^>]*>'
        ]
        
        for pattern in checked_patterns:
            checked_matches = re.findall(pattern, html_content, re.IGNORECASE)
            for name, value in checked_matches:
                if value and value.strip():
                    clean_name = name.replace('_', ' ').replace('-', ' ').title()
                    # Avoid duplicates
                    if not any(field['label'] == clean_name and field['value'] == value.strip() for field in form_fields):
                        form_fields.append({
                            'label': clean_name,
                            'value': value.strip()
                        })
        
        # Alternative approach: Look for any input with checked attribute
        general_checked_pattern = r'<input[^>]*checked[^>]*>'
        checked_inputs = re.findall(general_checked_pattern, html_content, re.IGNORECASE)
        
        for input_tag in checked_inputs:
            # Extract name and value from this specific input
            name_match = re.search(r'name=["\']([^"\'>]+)["\']', input_tag, re.IGNORECASE)
            value_match = re.search(r'value=["\']([^"\'>]*)["\']', input_tag, re.IGNORECASE)
            
            if name_match:
                name = name_match.group(1)
                value = value_match.group(1) if value_match else 'Yes'
                clean_name = name.replace('_', ' ').replace('-', ' ').title()
                # Avoid duplicates
                if not any(field['label'] == clean_name for field in form_fields):
                    form_fields.append({
                        'label': clean_name,
                        'value': value.strip() if value else 'Yes'
                    })
        
        # Remove duplicates while preserving order
        seen = set()
        unique_fields = []
        for field in form_fields:
            field_key = f"{field['label']}:{field['value']}"
            if field_key not in seen:
                seen.add(field_key)
                unique_fields.append(field)
        
        logger.info(f"Parsed form: title='{title}', company='{company_name}', fields={len(unique_fields)}")
        
        # Debug: Log all extracted fields
        for i, field in enumerate(unique_fields):
            logger.info(f"Field {i+1}: {field['label']} = {field['value']}")
        
        return {
            'title': title,
            'company_name': company_name,
            'logo_data': logo_data,
            'fields': unique_fields
        }
        
    except Exception as e:
        logger.error(f"Error parsing HTML form: {e}")
        return {
            'title': 'Form',
            'company_name': 'Your Company',
            'logo_data': None,
            'fields': []
        }

# Helper function to convert HTML to PDF using reportlab
def convert_html_to_pdf(html_path, filename):
    """Convert HTML content to PDF using reportlab"""
    import tempfile
    import os
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    import re
    
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}_completed.pdf")
    
    try:
        # Read HTML content
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse the HTML to extract content
        form_structure = parse_html_form(html_content)
        
        logger.info(f"Creating PDF with: {len(form_structure['fields'])} fields")
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            textColor=colors.HexColor('#2c5aa0'),
            alignment=1  # Center alignment
        )
        
        company_style = ParagraphStyle(
            'CompanyName',
            parent=styles['Normal'],
            fontSize=16,
            textColor=colors.HexColor('#2c5aa0'),
            spaceAfter=20,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold'
        )
        
        field_style = ParagraphStyle(
            'FieldStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=8,
            leftIndent=20
        )
        
        # Build PDF content
        story = []
        
        # Add logo if present
        if form_structure.get('logo_data') and form_structure['logo_data'].startswith('data:image'):
            try:
                import base64
                from io import BytesIO
                from reportlab.platypus import Image
                
                # Extract image data from base64
                header, encoded = form_structure['logo_data'].split(',', 1)
                image_data = base64.b64decode(encoded)
                
                # Create image from data
                image_stream = BytesIO(image_data)
                logo_image = Image(image_stream, width=2*inch, height=1*inch, hAlign='CENTER')
                story.append(logo_image)
                story.append(Spacer(1, 12))
                
                logger.info("Added logo to PDF")
            except Exception as logo_error:
                logger.warning(f"Could not add logo to PDF: {logo_error}")
        
        # Add title
        story.append(Paragraph(form_structure['title'], title_style))
        story.append(Spacer(1, 12))
        
        # Add company name
        story.append(Paragraph(form_structure['company_name'], company_style))
        story.append(Spacer(1, 20))
        
        # Add form fields
        if form_structure['fields']:
            story.append(Paragraph("Form Information:", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Group fields by categories for better organization
            personal_fields = []
            preference_fields = []
            other_fields = []
            
            for field in form_structure['fields']:
                field_label_lower = field['label'].lower()
                if any(word in field_label_lower for word in ['name', 'email', 'phone', 'contact']):
                    personal_fields.append(field)
                elif any(word in field_label_lower for word in ['preference', 'session', 'meal', 'newsletter']):
                    preference_fields.append(field)
                else:
                    other_fields.append(field)
            
            # Create sections for better organization
            sections = []
            if personal_fields:
                sections.append(("Personal Information", personal_fields))
            if preference_fields:
                sections.append(("Preferences & Selections", preference_fields))
            if other_fields:
                sections.append(("Additional Information", other_fields))
            
            if not sections:  # Fallback if categorization fails
                sections = [("Form Data", form_structure['fields'])]
            
            for section_title, section_fields in sections:
                # Add section heading
                section_style = ParagraphStyle(
                    'SectionHeading',
                    parent=styles['Heading3'],
                    fontSize=14,
                    spaceAfter=12,
                    textColor=colors.HexColor('#2c5aa0'),
                    fontName='Helvetica-Bold'
                )
                story.append(Paragraph(section_title, section_style))
                story.append(Spacer(1, 8))
                
                # Create table data for this section
                table_data = []
                for field in section_fields:
                    # Clean up field values for PDF display
                    field_label = field['label'][:50] + '...' if len(field['label']) > 50 else field['label']
                    field_value = field['value'][:100] + '...' if len(field['value']) > 100 else field['value']
                    
                    table_data.append([field_label + ':', field_value])
                
                if table_data:
                    # Create table with proper styling
                    table = Table(table_data, colWidths=[2.5*inch, 3.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f7fafc')),
                        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#2d3748')),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 11),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 20))
        else:
            story.append(Paragraph("No form data was captured. This is a blank form template.", styles['Normal']))
        
        # Add footer information
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            alignment=1  # Center alignment
        )
        
        from datetime import datetime
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
        story.append(Paragraph("This document was automatically generated from an HTML form.", footer_style))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF successfully created at: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        # Try to clean up partial file
        if os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except:
                pass
        raise Exception(f"PDF conversion failed: {str(e)}")
        
        # Build PDF
        doc.build(story)
        return output_path
        
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        raise Exception(f"PDF conversion failed: {str(e)}")

# Helper function to convert HTML to DOCX
def convert_html_to_docx(html_path, filename, form_structure):
    """Convert HTML to DOCX format using python-docx"""
    import tempfile
    import os
    
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}_completed.docx")
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Create a new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add logo if present
        if form_structure.get('logo_data') and form_structure['logo_data'].startswith('data:image'):
            try:
                import base64
                from io import BytesIO
                
                # Extract image data from base64
                header, encoded = form_structure['logo_data'].split(',', 1)
                image_data = base64.b64decode(encoded)
                
                # Create image from data and add to document
                image_stream = BytesIO(image_data)
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
                logo_run.add_picture(image_stream, width=Inches(2.0))
                
                doc.add_paragraph()  # Add space after logo
                logger.info("Added logo to DOCX")
            except Exception as logo_error:
                logger.warning(f"Could not add logo to DOCX: {logo_error}")
        
        # Add title with styling
        title_para = doc.add_heading(form_structure['title'], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        
        # Add company name with styling
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(form_structure['company_name'])
        company_run.bold = True
        company_run.font.size = Pt(16)
        company_run.font.color.rgb = RGBColor(44, 90, 160)  # Blue color
        
        # Add space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add form fields
        if form_structure['fields']:
            # Add section heading
            doc.add_heading('Form Information:', level=1)
            
            # Group fields by categories for better organization
            personal_fields = []
            preference_fields = []
            other_fields = []
            
            for field in form_structure['fields']:
                field_label_lower = field['label'].lower()
                if any(word in field_label_lower for word in ['name', 'email', 'phone', 'contact']):
                    personal_fields.append(field)
                elif any(word in field_label_lower for word in ['preference', 'session', 'meal', 'newsletter']):
                    preference_fields.append(field)
                else:
                    other_fields.append(field)
            
            # Create sections for better organization
            sections = []
            if personal_fields:
                sections.append(("Personal Information", personal_fields))
            if preference_fields:
                sections.append(("Preferences & Selections", preference_fields))
            if other_fields:
                sections.append(("Additional Information", other_fields))
            
            if not sections:  # Fallback if categorization fails
                sections = [("Form Data", form_structure['fields'])]
            
            for section_title, section_fields in sections:
                # Add section subheading
                section_heading = doc.add_heading(section_title, level=2)
                
                # Create a table for this section
                table = doc.add_table(rows=1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Light Grid Accent 1'
                
                # Set column widths
                table.columns[0].width = Inches(2.5)
                table.columns[1].width = Inches(4.0)
                
                # Add header row for this section
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Field'
                header_cells[1].text = 'Value'
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows for this section
                for field in section_fields:
                    row_cells = table.add_row().cells
                    row_cells[0].text = field['label']
                    row_cells[1].text = field['value']
                
                # Add spacing after each section
                doc.add_paragraph()
        else:
            # Add message for empty form
            no_data_para = doc.add_paragraph('No form data was captured. This is a blank form template.')
            no_data_para.italic = True
        
        # Add footer section
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add generation info
        from datetime import datetime
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        footer_run.font.size = Pt(10)
        footer_run.italic = True
        
        footer_para2 = doc.add_paragraph()
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run2 = footer_para2.add_run("This document was automatically generated from an HTML form.")
        footer_run2.font.size = Pt(10)
        footer_run2.italic = True
        
        # Save the document
        doc.save(output_path)
        
        logger.info(f"DOCX successfully created at: {output_path}")
        return output_path
        
    except ImportError:
        logger.error("python-docx not available")
        raise Exception("DOCX conversion not available. Please install python-docx.")
    except Exception as e:
        logger.error(f"DOCX conversion error: {e}")
        # Try to clean up partial file
        if os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except:
                pass
        raise Exception(f"DOCX conversion failed: {str(e)}")
# Fallback conversion function
async def fallback_form_conversion(html_content, filename, format_type):
    """Fallback conversion using the existing form generator"""
    try:
        from main import FORM_GENERATOR, _json_to_professional_form
        
        # Create a basic form structure from HTML
        form_data = {
            "title": filename.replace('_', ' ').title(),
            "description": "Form converted from HTML",
            "company_name": "Your Company",
            "form_type": "other",
            "sections": [
                {
                    "title": "Form Content", 
                    "description": "Content from HTML form",
                    "fields": [
                        {
                            "name": "content",
                            "label": "Form Content",
                            "field_type": "textarea",
                            "required": False,
                            "description": "Original form content"
                        }
                    ]
                }
            ],
            "footer_text": "Converted from HTML format"
        }
        
        # Convert to ProfessionalForm object
        professional_form = _json_to_professional_form(form_data)
        
        # Generate the requested format
        if format_type == 'pdf':
            file_path = FORM_GENERATOR.create_pdf_form(professional_form)
            content_type = "application/pdf"
        else:  # docx
            file_path = FORM_GENERATOR.create_docx_form(professional_form)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Read the generated file
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        # Clean up the temporary file
        import os
        os.remove(file_path)
        
        # Return the file as a download
        from fastapi.responses import Response
        return Response(
            content=file_content,
            media_type=content_type,
            headers={"Content-Disposition": f"attachment; filename={filename}_fallback.{format_type}"}
        )
        
    except Exception as e:
        logger.error(f"Fallback conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
@app.get("/api/system-stats")
async def get_system_statistics():
    """Get system statistics for analytics"""
    try:
        system_stats = get_system_stats()
        tool_stats = get_tool_stats()
        
        return {
            "system": system_stats,
            "tools": tool_stats,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Stats retrieval failed: {str(e)}")

@app.post("/api/create-tenant")
async def create_tenant_endpoint(request: TenantRequest):
    """Create a new tenant"""
    try:
        tenant_config = create_tenant(
            request.tenant_id,
            request.name,
            request.permissions
        )
        
        return {
            "success": True,
            "tenant_id": tenant_config.tenant_id,
            "name": tenant_config.name,
            "permissions": tenant_config.permissions
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Tenant creation failed: {str(e)}")

@app.get("/api/tenants")
async def list_tenants():
    """List all available tenants"""
    from main import _tenant_registry
    tenants = []

    for tenant_id, config in _tenant_registry.items():
        tenants.append({
            "tenant_id": tenant_id,
            "name": config.name,
            "permissions": config.permissions
        })

    # Ensure default tenant exists
    if not tenants:
        tenants.append({
            "tenant_id": "default",
            "name": "Default Tenant",
            "permissions": ["read_documents", "use_tools", "generate_forms"]
        })

    return {"tenants": tenants}

@app.get("/api/debug/tenant/{tenant_id}")
async def debug_tenant_status(tenant_id: str):
    """Debug endpoint to check tenant document indexing status"""
    try:
        from main import get_retriever_for_tenant, _tenant_index_path
        import os

        # Check if tenant exists
        from main import _tenant_registry
        tenant_exists = tenant_id in _tenant_registry

        # Check if index directory exists
        index_path = _tenant_index_path(tenant_id)
        index_exists = os.path.exists(index_path)

        # Check if retriever can be created
        retriever = get_retriever_for_tenant(tenant_id)
        retriever_available = retriever is not None

        # List files in upload directory
        upload_path = uploads_dir / tenant_id
        uploaded_files = []
        if upload_path.exists():
            uploaded_files = [f.name for f in upload_path.iterdir() if f.is_file()]

        return {
            "tenant_id": tenant_id,
            "tenant_exists": tenant_exists,
            "index_path": index_path,
            "index_exists": index_exists,
            "retriever_available": retriever_available,
            "uploaded_files": uploaded_files,
            "upload_path": str(upload_path)
        }

    except Exception as e:
        return {"error": str(e)}

@app.get("/api/tools/{tenant_id}")
async def get_tenant_tools(tenant_id: str):
    """Get available tools for a tenant"""
    try:
        from main import get_tenant_tools
        tools = get_tenant_tools(tenant_id)

        return {
            "success": True,
            "tools": [
                {
                    "name": tool.name,
                    "description": tool.description,
                    "args": tool.args if hasattr(tool, 'args') else {}
                }
                for tool in tools
            ],
            "total": len(tools)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get tools: {str(e)}")

@app.post("/api/dynamic-apis/{tenant_id}")
async def register_dynamic_api(tenant_id: str, api_config: dict):
    """Register a new dynamic API for a tenant"""
    try:
        from main import DYNAMIC_API_MANAGER, DynamicAPI
        
        # Create DynamicAPI object
        api = DynamicAPI(
            name=api_config.get("name", ""),
            base_url=api_config.get("base_url", ""),
            method=api_config.get("method", "GET"),
            headers=api_config.get("headers", {}),
            auth_type=api_config.get("auth_type", "none"),
            auth_value=api_config.get("auth_value", ""),
            description=api_config.get("description", ""),
            parameters=api_config.get("parameters", {})
        )
        
        # Register the API
        success = DYNAMIC_API_MANAGER.register_api(api)
        
        if success:
            return {
                "success": True,
                "message": f"Dynamic API '{api.name}' registered successfully",
                "api_name": api.name
            }
        else:
            raise HTTPException(status_code=400, detail="Failed to register API")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to register dynamic API: {str(e)}")

@app.get("/api/dynamic-apis/{tenant_id}")
async def get_dynamic_apis(tenant_id: str):
    """Get all registered dynamic APIs for a tenant"""
    try:
        from main import DYNAMIC_API_MANAGER
        
        apis = []
        for name, api in DYNAMIC_API_MANAGER.apis.items():
            apis.append({
                "name": api.name,
                "base_url": api.base_url,
                "method": api.method,
                "description": api.description,
                "auth_type": api.auth_type,
                "parameters": api.parameters
            })
        
        return {"apis": apis}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get dynamic APIs: {str(e)}")

@app.delete("/api/dynamic-apis/{tenant_id}/{api_name}")
async def remove_dynamic_api(tenant_id: str, api_name: str):
    """Remove a dynamic API"""
    try:
        from main import DYNAMIC_API_MANAGER
        
        success = DYNAMIC_API_MANAGER.remove_api(api_name)
        
        if success:
            return {
                "success": True,
                "message": f"Dynamic API '{api_name}' removed successfully"
            }
        else:
            raise HTTPException(status_code=404, detail="API not found")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to remove dynamic API: {str(e)}")

@app.post("/api/escalate")
async def create_escalation(request: ChatRequest):
    """Create an escalation ticket"""
    try:
        # Set current tenant context
        set_current_tenant(request.tenant_id)

        # Use the escalation agent
        response = chat_with_agent(request.message, request.tenant_id)

        return {
            "success": True,
            "response": response,
            "agent": "escalate",
            "timestamp": datetime.now().isoformat()
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Escalation failed: {str(e)}")

# -----------------------------
# Admin API Endpoints
# -----------------------------

@app.get("/api/admin/stats")
async def get_admin_stats():
    """Get admin dashboard statistics"""
    try:
        from main import _tenant_registry, _active_sessions, document_storage

        # Get tenant count
        total_tenants = len(_tenant_registry)
        active_tenants = sum(1 for config in _tenant_registry.values() if config.is_active)

        # Get active sessions count
        active_sessions = len(_active_sessions)

        # Get document count
        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM documents")
        total_documents = cursor.fetchone()[0]

        # Get open tickets count
        cursor.execute("SELECT COUNT(*) FROM escalation_tickets WHERE status = 'open'")
        open_tickets = cursor.fetchone()[0]

        conn.close()

        return {
            "total_tenants": total_tenants,
            "active_tenants": active_tenants,
            "active_sessions": active_sessions,
            "total_documents": total_documents,
            "open_tickets": open_tickets
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get stats: {str(e)}")

@app.get("/api/admin/tenants")
async def get_admin_tenants():
    """Get all tenants for admin dashboard"""
    try:
        from main import _tenant_registry, document_storage

        tenants = []
        for tenant_id, config in _tenant_registry.items():
            # Get document count for this tenant
            conn = sqlite3.connect(document_storage.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM documents WHERE tenant_id = ?", (tenant_id,))
            document_count = cursor.fetchone()[0]
            conn.close()

            tenants.append({
                "tenant_id": tenant_id,
                "name": config.name,
                "is_active": config.is_active,
                "permissions": config.permissions,
                "document_count": document_count,
                "created_at": config.created_at
            })

        return {"tenants": tenants}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get tenants: {str(e)}")

@app.post("/api/admin/tenants")
async def create_admin_tenant(request: TenantRequest):
    """Create a new tenant via admin dashboard"""
    try:
        from main import create_tenant, document_storage

        tenant_config = create_tenant(
            request.tenant_id,
            request.name,
            request.permissions
        )

        # Initialize tenant customization
        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO tenant_customization
            (tenant_id, created_at, updated_at)
            VALUES (?, ?, ?)
        ''', (request.tenant_id, datetime.now().isoformat(), datetime.now().isoformat()))
        conn.commit()
        conn.close()

        return {
            "success": True,
            "tenant_id": tenant_config.tenant_id,
            "name": tenant_config.name,
            "permissions": tenant_config.permissions
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Tenant creation failed: {str(e)}")

@app.get("/api/admin/sessions")
async def get_admin_sessions():
    """Get all active sessions for admin dashboard"""
    try:
        from main import _active_sessions, document_storage

        sessions = []
        for session_id, session in _active_sessions.items():
            sessions.append({
                "session_id": session_id,
                "tenant_id": session.tenant_id,
                "user_id": session.user_id,
                "created_at": session.created_at,
                "last_activity": session.last_activity,
                "status": "active",
                "permissions": session.permissions
            })

        return {"sessions": sessions}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get sessions: {str(e)}")

@app.get("/api/admin/tickets")
async def get_admin_tickets():
    """Get all escalation tickets for admin dashboard"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT ticket_id, session_id, tenant_id, user_id, title, description,
                   status, priority, assigned_to, created_at, updated_at
            FROM escalation_tickets
            ORDER BY created_at DESC
        ''')

        tickets = []
        for row in cursor.fetchall():
            tickets.append({
                "ticket_id": row[0],
                "session_id": row[1],
                "tenant_id": row[2],
                "user_id": row[3],
                "title": row[4],
                "description": row[5],
                "status": row[6],
                "priority": row[7],
                "assigned_to": row[8],
                "created_at": row[9],
                "updated_at": row[10]
            })

        conn.close()
        return {"tickets": tickets}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get tickets: {str(e)}")

@app.get("/api/admin/langsmith-status")
async def get_langsmith_status():
    """Get LangSmith integration status"""
    try:
        from main import LANGSMITH_CLIENT

        if LANGSMITH_CLIENT:
            return {
                "enabled": True,
                "project": os.getenv("LANGSMITH_PROJECT", "default"),
                "endpoint": os.getenv("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com")
            }
        else:
            return {"enabled": False}

    except Exception as e:
        return {"enabled": False, "error": str(e)}

class TenantCustomizationRequest(BaseModel):
    theme_color: str = "#667eea"
    chat_background_color: str = "#ffffff"
    welcome_message: str = "Hello! How can I help you today?"
    logo_url: Optional[str] = None
    widget_position: str = "bottom-right"
    custom_css: Optional[str] = None

class MeetingScheduleRequest(BaseModel):
    title: str
    description: Optional[str] = None
    scheduled_time: str  # ISO format datetime
    duration_minutes: int = 30
    meeting_type: str = "general"
    calendar_provider: Optional[str] = None
    session_id: Optional[str] = None
    tenant_id: str
    user_id: Optional[str] = None

class EscalationTicketRequest(BaseModel):
    title: str
    description: str
    priority: str = "medium"
    session_id: Optional[str] = None
    tenant_id: str
    user_id: Optional[str] = None
    chat_context: Optional[str] = None

@app.get("/api/admin/tenants/{tenant_id}/customization")
async def get_tenant_customization(tenant_id: str):
    """Get tenant customization settings"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT theme_color, logo_url, chat_background_color, widget_position,
                   welcome_message, custom_css
            FROM tenant_customization WHERE tenant_id = ?
        ''', (tenant_id,))

        row = cursor.fetchone()
        conn.close()

        if row:
            return {
                "theme_color": row[0],
                "logo_url": row[1],
                "chat_background_color": row[2],
                "widget_position": row[3],
                "welcome_message": row[4],
                "custom_css": row[5]
            }
        else:
            # Return defaults
            return {
                "theme_color": "#667eea",
                "logo_url": None,
                "chat_background_color": "#ffffff",
                "widget_position": "bottom-right",
                "welcome_message": "Hello! How can I help you today?",
                "custom_css": None
            }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get customization: {str(e)}")

@app.post("/api/admin/tenants/{tenant_id}/customization")
async def save_tenant_customization(tenant_id: str, request: TenantCustomizationRequest):
    """Save tenant customization settings"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO tenant_customization
            (tenant_id, theme_color, logo_url, chat_background_color, widget_position,
             welcome_message, custom_css, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?,
                    COALESCE((SELECT created_at FROM tenant_customization WHERE tenant_id = ?), ?),
                    ?)
        ''', (
            tenant_id, request.theme_color, request.logo_url, request.chat_background_color,
            request.widget_position, request.welcome_message, request.custom_css,
            tenant_id, datetime.now().isoformat(), datetime.now().isoformat()
        ))
        conn.commit()
        conn.close()

        return {"success": True, "message": "Customization saved successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save customization: {str(e)}")

@app.get("/api/tenants/{tenant_id}/widget")
async def get_tenant_widget_config(tenant_id: str, request: Request):
    """Get tenant widget configuration for embedding"""
    try:
        from main import document_storage, get_tenant_config

        # Check if tenant exists
        tenant_config = get_tenant_config(tenant_id)
        if not tenant_config:
            raise HTTPException(status_code=404, detail="Tenant not found")

        # Get customization
        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT theme_color, logo_url, chat_background_color, widget_position,
                   welcome_message, custom_css
            FROM tenant_customization WHERE tenant_id = ?
        ''', (tenant_id,))

        row = cursor.fetchone()
        conn.close()

        customization = {
            "theme_color": row[0] if row else "#667eea",
            "logo_url": row[1] if row else None,
            "chat_background_color": row[2] if row else "#ffffff",
            "widget_position": row[3] if row else "bottom-right",
            "welcome_message": row[4] if row else "Hello! How can I help you today?",
            "custom_css": row[5] if row else None
        }

        return {
            "tenant_id": tenant_id,
            "tenant_name": tenant_config.name,
            "customization": customization,
            "api_endpoint": f"{request.url.scheme}://{request.url.netloc}/api/chat"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get widget config: {str(e)}")

@app.get("/widget/{tenant_id}")
async def get_chat_widget(request: Request, tenant_id: str):
    """Get embeddable chat widget for a tenant"""
    try:
        from main import document_storage, get_tenant_config

        # Check if tenant exists
        tenant_config = get_tenant_config(tenant_id)
        if not tenant_config:
            raise HTTPException(status_code=404, detail="Tenant not found")

        # Get customization
        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT theme_color, logo_url, chat_background_color, widget_position,
                   welcome_message, custom_css
            FROM tenant_customization WHERE tenant_id = ?
        ''', (tenant_id,))

        row = cursor.fetchone()
        conn.close()

        customization = {
            "theme_color": row[0] if row else "#667eea",
            "logo_url": row[1] if row else None,
            "chat_background_color": row[2] if row else "#ffffff",
            "widget_position": row[3] if row else "bottom-right",
            "welcome_message": row[4] if row else "Hello! How can I help you today?",
            "custom_css": row[5] if row else None
        }

        widget_config = {
            "tenant_id": tenant_id,
            "tenant_name": tenant_config.name,
            "customization": customization,
            "api_endpoint": f"{request.url.scheme}://{request.url.netloc}/api/chat"
        }

        return templates.TemplateResponse("chat_widget.html", {
            "request": request,
            "config": widget_config
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get widget: {str(e)}")

@app.get("/widget/{tenant_id}/embed.js")
async def get_widget_embed_script(request: Request, tenant_id: str):
    """Get JavaScript embedding script for the chat widget"""
    try:
        widget_url = f"{request.url.scheme}://{request.url.netloc}/widget/{tenant_id}"

        embed_script = f"""
(function() {{
    // Create iframe for the chat widget
    var iframe = document.createElement('iframe');
    iframe.src = '{widget_url}';
    iframe.style.cssText = `
        position: fixed !important;
        bottom: 0 !important;
        right: 0 !important;
        width: 100% !important;
        height: 100% !important;
        border: none !important;
        z-index: 9999 !important;
        pointer-events: none !important;
        background: transparent !important;
    `;
    iframe.id = 'chatbot-widget-iframe';

    // Allow pointer events only on the widget area
    iframe.onload = function() {{
        iframe.style.pointerEvents = 'auto';
    }};

    // Add to page
    document.body.appendChild(iframe);

    // Handle responsive behavior
    function updateIframeSize() {{
        if (window.innerWidth <= 480) {{
            iframe.style.width = '100%';
            iframe.style.height = '100%';
        }} else {{
            iframe.style.width = '400px';
            iframe.style.height = '600px';
        }}
    }}

    window.addEventListener('resize', updateIframeSize);
    updateIframeSize();
}})();
"""

        return Response(
            content=embed_script,
            media_type="application/javascript",
            headers={"Cache-Control": "public, max-age=3600"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate embed script: {str(e)}")

# -----------------------------
# Meeting Scheduling Endpoints
# -----------------------------

@app.post("/api/schedule-meeting")
async def schedule_meeting(request: MeetingScheduleRequest):
    """Schedule a meeting with human agent"""
    try:
        from main import document_storage

        meeting_id = secrets.token_urlsafe(16)

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO meeting_schedules
            (meeting_id, session_id, tenant_id, user_id, title, description,
             scheduled_time, duration_minutes, meeting_type, calendar_provider,
             status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'scheduled', ?, ?)
        ''', (
            meeting_id, request.session_id, request.tenant_id, request.user_id,
            request.title, request.description, request.scheduled_time,
            request.duration_minutes, request.meeting_type, request.calendar_provider,
            datetime.now().isoformat(), datetime.now().isoformat()
        ))
        conn.commit()
        conn.close()

        # TODO: Integrate with Google Calendar/Outlook API
        # For now, just return success

        return {
            "success": True,
            "meeting_id": meeting_id,
            "message": "Meeting scheduled successfully",
            "scheduled_time": request.scheduled_time,
            "calendar_link": f"/api/meetings/{meeting_id}/calendar"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to schedule meeting: {str(e)}")

@app.get("/api/meetings/{meeting_id}")
async def get_meeting(meeting_id: str):
    """Get meeting details"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT meeting_id, session_id, tenant_id, user_id, title, description,
                   scheduled_time, duration_minutes, meeting_type, calendar_provider,
                   calendar_event_id, status, created_at, updated_at
            FROM meeting_schedules WHERE meeting_id = ?
        ''', (meeting_id,))

        row = cursor.fetchone()
        conn.close()

        if not row:
            raise HTTPException(status_code=404, detail="Meeting not found")

        return {
            "meeting_id": row[0],
            "session_id": row[1],
            "tenant_id": row[2],
            "user_id": row[3],
            "title": row[4],
            "description": row[5],
            "scheduled_time": row[6],
            "duration_minutes": row[7],
            "meeting_type": row[8],
            "calendar_provider": row[9],
            "calendar_event_id": row[10],
            "status": row[11],
            "created_at": row[12],
            "updated_at": row[13]
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get meeting: {str(e)}")

@app.post("/api/create-ticket")
async def create_escalation_ticket(request: EscalationTicketRequest):
    """Create an escalation ticket"""
    try:
        from main import document_storage

        ticket_id = secrets.token_urlsafe(16)

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO escalation_tickets
            (ticket_id, session_id, tenant_id, user_id, title, description,
             status, priority, chat_context, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, 'open', ?, ?, ?, ?)
        ''', (
            ticket_id, request.session_id, request.tenant_id, request.user_id,
            request.title, request.description, request.priority,
            request.chat_context, datetime.now().isoformat(), datetime.now().isoformat()
        ))
        conn.commit()
        conn.close()

        return {
            "success": True,
            "ticket_id": ticket_id,
            "message": "Escalation ticket created successfully",
            "status": "open",
            "priority": request.priority
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create ticket: {str(e)}")

@app.get("/api/tickets/{ticket_id}")
async def get_escalation_ticket(ticket_id: str):
    """Get escalation ticket details"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT ticket_id, session_id, tenant_id, user_id, title, description,
                   status, priority, assigned_to, chat_context, created_at, updated_at, resolved_at
            FROM escalation_tickets WHERE ticket_id = ?
        ''', (ticket_id,))

        row = cursor.fetchone()
        conn.close()

        if not row:
            raise HTTPException(status_code=404, detail="Ticket not found")

        return {
            "ticket_id": row[0],
            "session_id": row[1],
            "tenant_id": row[2],
            "user_id": row[3],
            "title": row[4],
            "description": row[5],
            "status": row[6],
            "priority": row[7],
            "assigned_to": row[8],
            "chat_context": row[9],
            "created_at": row[10],
            "updated_at": row[11],
            "resolved_at": row[12]
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get ticket: {str(e)}")

# -----------------------------
# Agent Console Endpoints
# -----------------------------

@app.get("/api/admin/sessions/{session_id}/messages")
async def get_session_messages(session_id: str):
    """Get chat messages for a specific session"""
    try:
        from main import document_storage

        conn = sqlite3.connect(document_storage.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT message_id, role, content, timestamp, agent_type
            FROM chat_messages
            WHERE session_id = ?
            ORDER BY timestamp ASC
        ''', (session_id,))

        messages = []
        for row in cursor.fetchall():
            messages.append({
                "message_id": row[0],
                "role": row[1],
                "content": row[2],
                "timestamp": row[3],
                "agent_type": row[4]
            })

        conn.close()
        return {"messages": messages}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get messages: {str(e)}")

class AgentMessageRequest(BaseModel):
    session_id: str
    message: str
    agent_id: str

@app.post("/api/agent/send-message")
async def send_agent_message(request: AgentMessageRequest):
    """Send a message as an agent in a chat session"""
    try:
        from main import document_storage, save_chat_message_to_history

        # Save the agent message to chat history
        save_chat_message_to_history(
            session_id=request.session_id,
            tenant_id="default",  # This should be determined from session
            role="agent",
            content=request.message,
            agent_type="human_agent"
        )

        return {
            "success": True,
            "message": "Message sent successfully",
            "timestamp": datetime.now().isoformat()
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send message: {str(e)}")

@app.post("/api/agent/takeover/{session_id}")
async def takeover_session(session_id: str):
    """Take over a chat session"""
    try:
        from main import _active_sessions

        session = _active_sessions.get(session_id)
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")

        # Mark session as taken over by agent
        # This would be enhanced with proper agent tracking

        return {
            "success": True,
            "message": "Session takeover successful",
            "session_id": session_id
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to takeover session: {str(e)}")

@app.post("/api/agent/end-takeover/{session_id}")
async def end_takeover_session(session_id: str):
    """End takeover of a chat session"""
    try:
        from main import _active_sessions

        session = _active_sessions.get(session_id)
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")

        # Remove agent takeover status
        # This would be enhanced with proper agent tracking

        return {
            "success": True,
            "message": "Session takeover ended",
            "session_id": session_id
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to end takeover: {str(e)}")

# -----------------------------
# Dynamic API Tool Registration
# -----------------------------

class DynamicToolRequest(BaseModel):
    tenant_id: str = "default"
    tool_name: str
    tool_type: str = "GET"  # GET or POST
    description: str
    base_url: str
    api_key: Optional[str] = None

@app.post("/api/register-tool")
async def register_api_tool(request: DynamicToolRequest):
    """Register a dynamic API tool for a tenant"""
    try:
        from main import register_dynamic_tool, make_http_get_tool, make_http_post_tool
        
        # Set current tenant context
        set_current_tenant(request.tenant_id)
        
        # Set environment variables for the tool
        base_url_env = f"{request.tool_name.upper()}_BASE_URL"
        api_key_env = f"{request.tool_name.upper()}_API_KEY" if request.api_key else None
        
        os.environ[base_url_env] = request.base_url
        if request.api_key:
            os.environ[api_key_env] = request.api_key
        
        # Create the appropriate tool
        if request.tool_type.upper() == "POST":
            tool = make_http_post_tool(
                name=request.tool_name,
                description=request.description,
                base_url_env=base_url_env,
                api_key_env=api_key_env
            )
        else:
            tool = make_http_get_tool(
                name=request.tool_name,
                description=request.description,
                base_url_env=base_url_env,
                api_key_env=api_key_env
            )
        
        # Register the tool
        register_dynamic_tool(request.tenant_id, tool)
        
        return {
            "success": True,
            "message": f"Successfully registered {request.tool_type} tool '{request.tool_name}' for tenant '{request.tenant_id}'",
            "tool_name": request.tool_name,
            "tool_type": request.tool_type,
            "tenant_id": request.tenant_id
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to register tool: {str(e)}")

@app.get("/api/tools/{tenant_id}")
async def get_tenant_tools_endpoint(tenant_id: str):
    """Get all available tools for a tenant"""
    try:
        from main import get_tenant_tools
        tools = get_tenant_tools(tenant_id)
        
        return {
            "success": True,
            "tools": [
                {
                    "name": tool.name,
                    "description": tool.description,
                    "args": getattr(tool, 'args', {})
                }
                for tool in tools
            ],
            "total": len(tools)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get tools: {str(e)}")

if __name__ == "__main__":
    print("🚀 Starting Multi-Agent Chatbot Web Interface...")
    print("📊 Dashboard: http://localhost:8000")
    print("🔧 Admin Panel: http://localhost:8000/admin")
    print("👥 Agent Console: http://localhost:8000/console")
    print("📚 API Docs: http://localhost:8000/docs")
    print("🎨 Widget Example: http://localhost:8000/widget/default")

    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
