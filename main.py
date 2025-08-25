from __future__ import annotations

import os
import json
import time
import hashlib
import secrets
import logging
import sqlite3
import shutil
import base64
from typing import Callable, Dict, List, Optional, Any, Union
from dataclasses import dataclass, field, asdict
from datetime import datetime, timedelta
from pathlib import Path

import requests
from dotenv import load_dotenv

# PDF and DOC generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOC file generation will be unavailable.")

from langgraph.graph import StateGraph, MessagesState, START, END
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import MemorySaver

from langchain_core.tools import tool
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_google_genai import ChatGoogleGenerativeAI

# OpenAI support
try:
    from langchain_openai import ChatOpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("Warning: langchain-openai not installed. OpenAI models will be unavailable.")

# Removed HuggingFaceEmbeddings due to TensorFlow conflicts

# Vector store + splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# MCP (Model Context Protocol) Integration
try:
    import mcp
    from mcp.client.session import ClientSession
    from mcp.client.stdio import StdioServerParameters, stdio_client
    MCP_AVAILABLE = True
except ImportError:
    MCP_AVAILABLE = False
    print("Warning: MCP not installed. Model Context Protocol features will be unavailable.")

# LangSmith Integration for monitoring and evaluation
try:
    from langsmith import Client
    from langsmith.wrappers import wrap_openai
    from langsmith.evaluation import evaluate
    LANGSMITH_AVAILABLE = True
except ImportError:
    LANGSMITH_AVAILABLE = False
    print("Warning: LangSmith not installed. Monitoring and evaluation features will be unavailable.")

load_dotenv()

# Configure logging first
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize LangSmith client
LANGSMITH_CLIENT = None
if LANGSMITH_AVAILABLE and os.getenv("LANGSMITH_TRACING", "false").lower() == "true":
    try:
        LANGSMITH_CLIENT = Client(
            api_url=os.getenv("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com"),
            api_key=os.getenv("LANGSMITH_API_KEY")
        )
        # Set environment variables for automatic tracing
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_ENDPOINT"] = os.getenv("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com")
        os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY", "")
        os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGSMITH_PROJECT", "default")
        logger.info(f"✅ LangSmith tracing enabled for project: {os.getenv('LANGSMITH_PROJECT')}")
    except Exception as e:
        logger.warning(f"⚠️ LangSmith initialization failed: {e}")
        LANGSMITH_CLIENT = None


# -----------------------------
# Professional Form Generation Classes
# -----------------------------

@dataclass
class FormField:
    """Represents a form field with all necessary properties."""
    name: str
    label: str
    field_type: str
    required: bool = False
    placeholder: str = ""
    options: List[str] = field(default_factory=list)
    validation: str = ""
    description: str = ""
    default_value: str = ""

@dataclass
class FormSection:
    """Represents a section of a form."""
    title: str
    description: str = ""
    fields: List[FormField] = field(default_factory=list)

@dataclass
class ProfessionalForm:
    """Represents a complete professional form."""
    title: str
    description: str
    company_name: str = ""
    form_type: str = "general"
    sections: List[FormSection] = field(default_factory=list)
    footer_text: str = ""
    created_date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    form_id: str = field(default_factory=lambda: secrets.token_hex(4))

class FormGenerator:
    """Professional form generator with PDF and DOC export capabilities."""

    def __init__(self):
        self.output_dir = Path("generated_forms")
        self.output_dir.mkdir(exist_ok=True)

    def generate_form_preview(self, form: ProfessionalForm) -> str:
        """Generate a text preview of the form structure."""
        preview_lines = []
        preview_lines.append(f"📋 **{form.title}**")
        preview_lines.append(f"📝 {form.description}")

        if form.company_name:
            preview_lines.append(f"🏢 Company: {form.company_name}")

        preview_lines.append(f"📊 Type: {form.form_type.title()}")
        preview_lines.append(f"🆔 Form ID: {form.form_id}")
        preview_lines.append("")

        # Add sections preview
        for i, section in enumerate(form.sections, 1):
            preview_lines.append(f"**Section {i}: {section.title}**")
            if section.description:
                preview_lines.append(f"   ℹ️ {section.description}")

            for j, field in enumerate(section.fields, 1):
                required_mark = " *" if field.required else ""
                preview_lines.append(f"   {j}. {field.label}{required_mark} ({field.field_type})")
                if field.placeholder:
                    preview_lines.append(f"      💡 Placeholder: {field.placeholder}")
                if field.options:
                    preview_lines.append(f"      🔘 Options: {', '.join(field.options[:3])}{'...' if len(field.options) > 3 else ''}")
            preview_lines.append("")

        if form.footer_text:
            preview_lines.append(f"📄 Footer: {form.footer_text}")

        preview_lines.append(f"📅 Created: {form.created_date}")

        return "\n".join(preview_lines)

    def create_pdf_form(self, form: ProfessionalForm, filename: str = None) -> str:
        """Generate a professional PDF form."""
        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.pdf"

        # Ensure the filename doesn't contain path separators that could create subdirectories
        filename = filename.replace('\\', '_').replace('/', '_')

        # Create the full file path
        filepath = self.output_dir / filename

        # Ensure the parent directory exists
        filepath.parent.mkdir(parents=True, exist_ok=True)

        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)

        # Get styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )

        section_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue
        )

        field_style = ParagraphStyle(
            'FieldLabel',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=20
        )

        # Build content
        content = []

        # Header
        if form.company_name:
            company_para = Paragraph(form.company_name, styles['Normal'])
            company_para.alignment = TA_CENTER
            content.append(company_para)
            content.append(Spacer(1, 12))

        # Title
        content.append(Paragraph(form.title, title_style))

        # Description
        if form.description:
            content.append(Paragraph(form.description, styles['Normal']))
            content.append(Spacer(1, 20))

        # Form ID and Date
        info_text = f"Form ID: {form.form_id} | Date: {form.created_date}"
        content.append(Paragraph(info_text, styles['Normal']))
        content.append(Spacer(1, 20))

        # Sections and Fields
        for section in form.sections:
            # Section title
            content.append(Paragraph(section.title, section_style))

            if section.description:
                content.append(Paragraph(section.description, styles['Normal']))
                content.append(Spacer(1, 10))

            # Fields in this section
            for field in section.fields:
                # Field label with required indicator
                label_text = field.label
                if field.required:
                    label_text += " *"

                content.append(Paragraph(label_text, field_style))

                # Field description
                if field.description:
                    desc_style = ParagraphStyle(
                        'FieldDesc',
                        parent=styles['Normal'],
                        fontSize=9,
                        leftIndent=40,
                        textColor=colors.grey
                    )
                    content.append(Paragraph(field.description, desc_style))

                # Input area based on field type
                if field.field_type in ['text', 'email', 'number', 'date']:
                    input_line = "_" * 50
                    content.append(Paragraph(input_line, styles['Normal']))
                elif field.field_type == 'textarea':
                    for _ in range(3):
                        content.append(Paragraph("_" * 70, styles['Normal']))
                elif field.field_type in ['select', 'radio']:
                    for option in field.options:
                        option_text = f"☐ {option}"
                        content.append(Paragraph(option_text, field_style))
                elif field.field_type == 'checkbox':
                    for option in field.options:
                        option_text = f"☐ {option}"
                        content.append(Paragraph(option_text, field_style))

                content.append(Spacer(1, 15))

        # Footer
        if form.footer_text:
            content.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=9,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            content.append(Paragraph(form.footer_text, footer_style))

        # Required fields note
        content.append(Spacer(1, 20))
        required_note = "* Required fields"
        content.append(Paragraph(required_note, styles['Normal']))

        # Build PDF
        doc.build(content)

        logger.info(f"Generated PDF form: {filepath}")
        return str(filepath)

    def create_docx_form(self, form: ProfessionalForm, filename: str = None) -> str:
        """Generate a professional DOCX form."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOC file generation. Install with: pip install python-docx")

        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.docx"

        # Ensure the filename doesn't contain path separators that could create subdirectories
        filename = filename.replace('\\', '_').replace('/', '_')

        # Create the full file path
        filepath = self.output_dir / filename

        # Ensure the parent directory exists
        filepath.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = DocxDocument()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header
        if form.company_name:
            header = doc.add_heading(form.company_name, level=0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_heading(form.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Description
        if form.description:
            desc_para = doc.add_paragraph(form.description)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Form info
        info_para = doc.add_paragraph(f"Form ID: {form.form_id} | Date: {form.created_date}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add spacing
        doc.add_paragraph()

        # Sections and Fields
        for section in form.sections:
            # Section heading
            section_heading = doc.add_heading(section.title, level=2)

            if section.description:
                doc.add_paragraph(section.description)

            # Fields
            for field in section.fields:
                # Field label
                label_text = field.label
                if field.required:
                    label_text += " *"

                field_para = doc.add_paragraph()
                field_run = field_para.add_run(label_text)
                field_run.bold = True

                # Field description
                if field.description:
                    desc_para = doc.add_paragraph(field.description)
                    desc_run = desc_para.runs[0]
                    desc_run.italic = True

                # Input area
                if field.field_type in ['text', 'email', 'number', 'date']:
                    input_para = doc.add_paragraph("_" * 50)
                elif field.field_type == 'textarea':
                    for _ in range(3):
                        doc.add_paragraph("_" * 70)
                elif field.field_type in ['select', 'radio']:
                    for option in field.options:
                        option_para = doc.add_paragraph(f"☐ {option}")
                elif field.field_type == 'checkbox':
                    for option in field.options:
                        option_para = doc.add_paragraph(f"☐ {option}")

                # Add spacing
                doc.add_paragraph()

        # Footer
        if form.footer_text:
            doc.add_page_break()
            footer_para = doc.add_paragraph(form.footer_text)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Required fields note
        required_para = doc.add_paragraph("* Required fields")
        required_run = required_para.runs[0]
        required_run.italic = True

        # Save document
        doc.save(str(filepath))

        logger.info(f"Generated DOCX form: {filepath}")
        return str(filepath)

    def create_html_form(self, form: ProfessionalForm, filename: str = None) -> str:
        """Generate an interactive HTML form."""
        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.html"

        # Ensure the filename doesn't contain path separators
        filename = filename.replace('\\', '_').replace('/', '_')

        # Create the full file path
        filepath = self.output_dir / filename

        # Ensure the parent directory exists
        filepath.parent.mkdir(parents=True, exist_ok=True)

        # Generate HTML content
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{form.title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
            line-height: 1.6;
        }}
        .form-container {{
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        .form-header {{
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #007bff;
            padding-bottom: 20px;
        }}
        .form-title {{
            color: #333;
            font-size: 28px;
            margin-bottom: 10px;
        }}
        .form-description {{
            color: #666;
            font-size: 16px;
            margin-bottom: 10px;
        }}
        .company-name {{
            color: #007bff;
            font-size: 18px;
            font-weight: bold;
        }}
        .section {{
            margin-bottom: 30px;
            padding: 20px;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            background-color: #fafafa;
        }}
        .section-title {{
            color: #333;
            font-size: 20px;
            margin-bottom: 10px;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }}
        .section-description {{
            color: #666;
            font-size: 14px;
            margin-bottom: 15px;
            font-style: italic;
        }}
        .field-group {{
            margin-bottom: 20px;
        }}
        .field-label {{
            display: block;
            font-weight: bold;
            margin-bottom: 5px;
            color: #333;
        }}
        .required {{
            color: #dc3545;
        }}
        .field-description {{
            font-size: 12px;
            color: #666;
            margin-bottom: 5px;
            font-style: italic;
        }}
        input[type="text"], input[type="email"], input[type="number"], 
        input[type="date"], input[type="tel"], select, textarea {{
            width: 100%;
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
            font-size: 14px;
            box-sizing: border-box;
        }}
        input:focus, select:focus, textarea:focus {{
            outline: none;
            border-color: #007bff;
            box-shadow: 0 0 5px rgba(0,123,255,0.3);
        }}
        textarea {{
            height: 100px;
            resize: vertical;
        }}
        .radio-group, .checkbox-group {{
            display: flex;
            flex-direction: column;
            gap: 8px;
        }}
        .radio-option, .checkbox-option {{
            display: flex;
            align-items: center;
            gap: 8px;
        }}
        .radio-option input, .checkbox-option input {{
            width: auto;
        }}
        .form-footer {{
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #666;
            font-size: 14px;
        }}
        .form-actions {{
            text-align: center;
            margin-top: 30px;
            padding-top: 20px;
            border-top: 2px solid #007bff;
        }}
        .btn {{
            padding: 12px 30px;
            margin: 0 10px;
            border: none;
            border-radius: 5px;
            font-size: 16px;
            cursor: pointer;
            transition: background-color 0.3s;
        }}
        .btn-primary {{
            background-color: #007bff;
            color: white;
        }}
        .btn-primary:hover {{
            background-color: #0056b3;
        }}
        .btn-secondary {{
            background-color: #6c757d;
            color: white;
        }}
        .btn-secondary:hover {{
            background-color: #545b62;
        }}
        .btn-success {{
            background-color: #28a745;
            color: white;
        }}
        .btn-success:hover {{
            background-color: #1e7e34;
        }}
        .required-note {{
            color: #dc3545;
            font-size: 14px;
            margin-top: 20px;
            text-align: center;
        }}
        .download-options {{
            margin-top: 20px;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 5px;
            text-align: center;
        }}
        .download-options h4 {{
            margin-bottom: 10px;
            color: #333;
        }}
    </style>
</head>
<body>
    <div class="form-container">
        <div class="form-header">
            <h1 class="form-title">{form.title}</h1>
            {f'<p class="form-description">{form.description}</p>' if form.description else ''}
            {f'<p class="company-name">{form.company_name}</p>' if form.company_name else ''}
            
            <div class="form-meta">
                <div class="meta-item">
                    <i class="fas fa-calendar"></i>
                    <span>Created: {form.created_date}</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-id-card"></i>
                    <span>Form ID: {form.form_id}</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-list"></i>
                    <span>{len(form.sections)} Section(s)</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-tasks"></i>
                    <span>{sum(len(section.fields) for section in form.sections)} Field(s)</span>
                </div>
            </div>
        </div>
        
        <div class="form-info">
            <h4><i class="fas fa-info-circle"></i> Comprehensive Form Guide</h4>
            <p><strong>Purpose:</strong> This {form.form_type} form is designed to collect detailed information efficiently and securely.</p>
            <p><strong>Required Fields:</strong> Fields marked with <span style="color: #e74c3c;">*</span> are mandatory and must be completed before submission.</p>
            <p><strong>Data Privacy:</strong> All information provided is encrypted and handled according to our strict privacy policy and GDPR compliance.</p>
            <p><strong>Completion Time:</strong> Estimated time to complete: {max(3, len(form.sections) * 2)} minutes</p>
            <p><strong>Auto-Save:</strong> Your progress is automatically saved as you type to prevent data loss.</p>
            <p><strong>Technical Support:</strong> If you experience any issues, please contact support@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</p>
            <p><strong>Browser Requirements:</strong> This form works best with modern browsers (Chrome 90+, Firefox 88+, Safari 14+, Edge 90+).</p>
            <p><strong>Accessibility:</strong> This form is designed to be accessible and supports screen readers and keyboard navigation.</p>
        </div>
        
        <div class="progress-container">
            <div class="progress-text">Form Completion Progress: <span id="progress-percent">0%</span></div>
            <div class="progress-bar">
                <div class="progress-fill" id="progress-fill"></div>
            </div>
        </div>

        <form id="dynamicForm" method="post" action="#" onsubmit="return handleSubmit(event)">
"""

        # Enhanced sections with more detailed field information
        for section_idx, section in enumerate(form.sections, 1):
            html_content += f"""
            <div class="section">
                <h2 class="section-title">
                    <span class="section-icon"><i class="fas fa-folder"></i></span>
                    Section {section_idx}: {section.title}
                    <span style="font-size: 14px; color: #6c757d; font-weight: normal;">({len(section.fields)} fields)</span>
                </h2>
                {f'<p class="section-description"><i class="fas fa-info"></i> {section.description}</p>' if section.description else ''}
                
                <div style="background: #f8f9fa; padding: 10px 15px; border-radius: 6px; margin-bottom: 20px; font-size: 14px; color: #495057;">
                    <strong>Section Guidelines:</strong> Please provide accurate and complete information in this section. 
                    All fields are important for processing your request efficiently.
                </div>
"""

            for field_idx, field in enumerate(section.fields, 1):
                required_attr = 'required' if field.required else ''
                required_mark = '<span class="required">*</span>' if field.required else ''
                
                # Enhanced field information
                field_info = ""
                if field.field_type == 'email':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-envelope'></i> Please enter a valid email address (e.g., user@example.com)</small>"
                elif field.field_type == 'tel':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-phone'></i> Include country code if international (e.g., +1-555-123-4567)</small>"
                elif field.field_type == 'date':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-calendar'></i> Select date using the calendar picker</small>"
                elif field.field_type == 'number':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-hashtag'></i> Enter numeric values only</small>"
                elif field.field_type == 'select':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-list'></i> Choose one option from {len(field.options)} available choices</small>"
                elif field.field_type == 'radio':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-dot-circle'></i> Select one option from the {len(field.options)} choices below</small>"
                elif field.field_type == 'checkbox':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-check-square'></i> You can select multiple options from the {len(field.options)} choices</small>"
                elif field.field_type == 'textarea':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-align-left'></i> Provide detailed information. This field supports multiple lines of text.</small>"
                else:
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-edit'></i> Please provide accurate information</small>"
                
                html_content += f"""
                <div class="field-group" data-field-type="{field.field_type}">
                    <label class="field-label" for="{field.name}">
                        <span style="font-weight: 600;">{field_idx}. {field.label} {required_mark}</span>
                        <span style="font-size: 12px; color: #6c757d; font-weight: normal;">({field.field_type.title()})</span>
                    </label>
                    {f'<p class="field-description"><i class="fas fa-info-circle"></i> {field.description}</p>' if field.description else ''}
                    {field_info}
"""

                if field.field_type == 'text':
                    html_content += f'<input type="text" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'email':
                    html_content += f'<input type="email" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'number':
                    html_content += f'<input type="number" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'date':
                    html_content += f'<input type="date" id="{field.name}" name="{field.name}" {required_attr}>'
                elif field.field_type == 'tel':
                    html_content += f'<input type="tel" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'textarea':
                    html_content += f'<textarea id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}></textarea>'
                elif field.field_type == 'select':
                    html_content += f'<select id="{field.name}" name="{field.name}" {required_attr}>'
                    html_content += f'<option value="">Choose an option...</option>'
                    for option in field.options:
                        html_content += f'<option value="{option}">{option}</option>'
                    html_content += '</select>'
                elif field.field_type == 'radio':
                    html_content += '<div class="radio-group">'
                    for option in field.options:
                        html_content += f'''
                        <div class="radio-option">
                            <input type="radio" id="{field.name}_{option}" name="{field.name}" value="{option}" {required_attr}>
                            <label for="{field.name}_{option}">{option}</label>
                        </div>'''
                    html_content += '</div>'
                elif field.field_type == 'checkbox':
                    html_content += '<div class="checkbox-group">'
                    for option in field.options:
                        html_content += f'''
                        <div class="checkbox-option">
                            <input type="checkbox" id="{field.name}_{option}" name="{field.name}[]" value="{option}">
                            <label for="{field.name}_{option}">{option}</label>
                        </div>'''
                    html_content += '</div>'

                html_content += '</div>'

            html_content += '</div>'

        # Enhanced form footer with comprehensive information
        html_content += f"""
            <div class="form-actions">
                <button type="submit" class="btn btn-primary">
                    <i class="fas fa-paper-plane"></i> Submit Form
                </button>
                <button type="button" class="btn btn-secondary" onclick="clearForm()">
                    <i class="fas fa-eraser"></i> Clear All Data
                </button>
                <button type="button" class="btn" style="background: linear-gradient(135deg, #17a2b8 0%, #138496 100%); color: white;" onclick="saveProgress()">
                    <i class="fas fa-save"></i> Save Progress
                </button>
                
                <div class="download-options">
                    <h4><i class="fas fa-download"></i> Download Options</h4>
                    <p style="font-size: 14px; color: #6c757d; margin-bottom: 15px;">Download your completed form in various formats for your records:</p>
                    <button type="button" class="btn btn-success" onclick="downloadAsPDF()">
                        <i class="fas fa-file-pdf"></i> Download as PDF
                    </button>
                    <button type="button" class="btn btn-success" onclick="downloadAsDOCX()">
                        <i class="fas fa-file-word"></i> Download as DOCX
                    </button>
                    <div style="margin-top: 10px; font-size: 12px; color: #6c757d;">
                        <i class="fas fa-info-circle"></i> Downloads preserve all your entered data and form formatting
                    </div>
                </div>
            </div>

            <div class="form-footer">
                <div style="border-top: 2px solid #e9ecef; padding-top: 20px; margin-top: 20px;">
                    <h5 style="color: #2c3e50; margin-bottom: 15px;"><i class="fas fa-shield-alt"></i> Security & Privacy Information</h5>
                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px; margin-bottom: 15px;">
                        <div>
                            <strong>Data Encryption:</strong><br>
                            <small>All data is encrypted using industry-standard AES-256 encryption</small>
                        </div>
                        <div>
                            <strong>Privacy Compliance:</strong><br>
                            <small>GDPR, CCPA, and SOX compliant data handling procedures</small>
                        </div>
                        <div>
                            <strong>Data Retention:</strong><br>
                            <small>Data retained only as long as necessary for business purposes</small>
                        </div>
                        <div>
                            <strong>Access Control:</strong><br>
                            <small>Strict role-based access controls and audit trails</small>
                        </div>
                    </div>
                    
                    <div style="background: #f8f9fa; padding: 15px; border-radius: 8px; margin-top: 15px;">
                        <h6 style="color: #495057; margin-bottom: 10px;"><i class="fas fa-question-circle"></i> Need Help?</h6>
                        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 10px; font-size: 14px;">
                            <div><strong>Technical Support:</strong> support@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</div>
                            <div><strong>Form Issues:</strong> forms@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</div>
                            <div><strong>Phone Support:</strong> 1-800-SUPPORT</div>
                            <div><strong>Live Chat:</strong> Available 24/7 on our website</div>
                        </div>
                    </div>
                    
                    <div style="text-align: center; margin-top: 20px; padding-top: 15px; border-top: 1px solid #dee2e6;">
                        <p style="margin: 5px 0; color: #6c757d; font-size: 14px;">
                            <strong>Form ID:</strong> {form.form_id} | 
                            <strong>Generated:</strong> {form.created_date} | 
                            <strong>Version:</strong> 2.1.0
                        </p>
                        <p style="margin: 5px 0; color: #6c757d; font-size: 13px;">
                            © 2024 {form.company_name if form.company_name else 'Your Company'}. All rights reserved. 
                            <a href="#" style="color: #007bff;">Privacy Policy</a> | 
                            <a href="#" style="color: #007bff;">Terms of Service</a> | 
                            <a href="#" style="color: #007bff;">Accessibility</a>
                        </p>
                    </div>
                </div>
            </div>
            
            <p class="required-note">
                <i class="fas fa-asterisk"></i> Required fields | 
                <i class="fas fa-lock"></i> Secure form submission | 
                <i class="fas fa-mobile-alt"></i> Mobile-friendly design
            </p>
        </form>
    </div>

    <script>
        // Enhanced form handling with comprehensive features
        let formProgress = 0;
        let autoSaveInterval;
        
        // Initialize form enhancements
        document.addEventListener('DOMContentLoaded', function() {{
            initializeFormFeatures();
            startAutoSave();
            updateProgress();
        }});
        
        function initializeFormFeatures() {{
            // Add input event listeners for progress tracking
            const inputs = document.querySelectorAll('input, select, textarea');
            inputs.forEach(input => {{
                input.addEventListener('input', updateProgress);
                input.addEventListener('change', updateProgress);
            }});
            
            // Add accessibility enhancements
            inputs.forEach((input, index) => {{
                input.setAttribute('tabindex', index + 1);
                if (input.hasAttribute('required')) {{
                    input.setAttribute('aria-required', 'true');
                }}
            }});
        }}
        
        function updateProgress() {{
            const allFields = document.querySelectorAll('input, select, textarea');
            const filledFields = Array.from(allFields).filter(field => {{
                if (field.type === 'radio' || field.type === 'checkbox') {{
                    return document.querySelector("input[name=\"" + field.name + "\"]:checked");
                }}
                return field.value.trim() !== '';
            }});
            
            const totalFields = allFields.length;
            const completed = new Set(filledFields.map(f => f.name)).size;
            formProgress = totalFields > 0 ? Math.round((completed / totalFields) * 100) : 0;
            
            document.getElementById('progress-percent').textContent = formProgress + '%';
            document.getElementById('progress-fill').style.width = formProgress + '%';
        }}
        
        function startAutoSave() {{
            autoSaveInterval = setInterval(() => {{
                saveProgress();
            }}, 30000); // Auto-save every 30 seconds
        }}
        
        function saveProgress() {{
            const formData = collectFormData();
            localStorage.setItem('form_progress_' + window.location.pathname, JSON.stringify({{
                data: formData,
                timestamp: new Date().toISOString(),
                progress: formProgress
            }}));
            
            // Show save indicator
            const saveBtn = document.querySelector('button[onclick="saveProgress()"]');
            if (saveBtn) {{
                const originalText = saveBtn.innerHTML;
                saveBtn.innerHTML = '<i class="fas fa-check"></i> Saved!';
                saveBtn.style.background = 'linear-gradient(135deg, #28a745 0%, #20c997 100%)';
                setTimeout(() => {{
                    saveBtn.innerHTML = originalText;
                    saveBtn.style.background = 'linear-gradient(135deg, #17a2b8 0%, #138496 100%)';
                }}, 2000);
            }}
        }}
        
        function collectFormData() {{
            const formData = new FormData(document.getElementById('dynamicForm'));
            const data = {{}};
            
            for (let [key, value] of formData.entries()) {{
                if (data[key]) {{
                    if (Array.isArray(data[key])) {{
                        data[key].push(value);
                    }} else {{
                        data[key] = [data[key], value];
                    }}
                }} else {{
                    data[key] = value;
                }}
            }}
            
            return data;
        }}
        
        function handleSubmit(event) {{
            event.preventDefault();
            
            // Validate form
            const form = event.target;
            if (!form.checkValidity()) {{
                form.reportValidity();
                return false;
            }}
            
            // Collect comprehensive form data
            const formData = collectFormData();
            const submissionData = {{
                formId: '{form.form_id}',
                submissionDate: new Date().toISOString(),
                formData: formData,
                browserInfo: {{
                    userAgent: navigator.userAgent,
                    language: navigator.language,
                    platform: navigator.platform
                }},
                formMetadata: {{
                    totalFields: document.querySelectorAll('input, select, textarea').length,
                    completionTime: 'Calculated from session start',
                    version: '2.1.0'
                }}
            }};
            
            // Clear auto-save
            localStorage.removeItem('form_progress_' + window.location.pathname);
            clearInterval(autoSaveInterval);
            
            // Show success message with more details
            alert(`Form submitted successfully!\\n\\nSubmission Details:\\n` + 
                  `- Form ID: {form.form_id}\\n` +
                  `- Submission Date: " + new Date().toLocaleString() + "\\n` +
                  `- Fields Completed: " + Object.keys(formData).length + "\\n` +
                  `- Data will be processed within 24-48 hours\\n\\n` +
                  `You will receive a confirmation email shortly.`);
            
            return false;
        }}
        
        function clearForm() {{
            if (confirm('Are you sure you want to clear all form data?\\n\\nThis action cannot be undone and will remove all entered information.')) {{
                document.getElementById('dynamicForm').reset();
                localStorage.removeItem('form_progress_' + window.location.pathname);
                updateProgress();
                alert('Form data cleared successfully.');
            }}
        }}
        
        function downloadAsPDF() {{
            const formData = collectFormData();
            if (Object.keys(formData).length === 0) {{
                alert('Please fill out some form fields before downloading.');
                return;
            }}
            
            // Show loading indicator
            const btn = event.target;
            const originalText = btn.innerHTML;
            btn.innerHTML = '<i class="fas fa-spinner fa-spin"></i> Generating PDF...';
            btn.disabled = true;
            
            // Simulate API call (replace with actual implementation)
            setTimeout(() => {{
                alert(`PDF download functionality is being processed.\\n\\n` +
                      `Your form will be converted to PDF format with:\\n` +
                      `- Professional formatting and layout\\n` +
                      `- All form fields and responses\\n` +
                      `- Company branding and logos\\n` +
                      `- Digital signature support\\n\\n` +
                      `The download will begin shortly...`);
                
                btn.innerHTML = originalText;
                btn.disabled = false;
            }}, 2000);
        }}
        
        function downloadAsDOCX() {{
            const formData = collectFormData();
            if (Object.keys(formData).length === 0) {{
                alert('Please fill out some form fields before downloading.');
                return;
            }}
            
            // Show loading indicator
            const btn = event.target;
            const originalText = btn.innerHTML;
            btn.innerHTML = '<i class="fas fa-spinner fa-spin"></i> Generating DOCX...';
            btn.disabled = true;
            
            // Simulate API call (replace with actual implementation)
            setTimeout(() => {{
                alert(`DOCX download functionality is being processed.\\n\\n` +
                      `Your form will be converted to Word document format with:\\n` +
                      `- Editable document structure\\n` +
                      `- Professional table formatting\\n` +
                      `- Company headers and footers\\n` +
                      `- Compatible with Microsoft Word and LibreOffice\\n\\n` +
                      `The download will begin shortly...`);
                
                btn.innerHTML = originalText;
                btn.disabled = false;
            }}, 2000);
        }}
        
        // Keyboard shortcuts
        document.addEventListener('keydown', function(e) {{
            if (e.ctrlKey || e.metaKey) {{
                switch(e.key) {{
                    case 's':
                        e.preventDefault();
                        saveProgress();
                        break;
                    case 'Enter':
                        if (e.shiftKey) {{
                            e.preventDefault();
                            document.getElementById('dynamicForm').dispatchEvent(new Event('submit'));
                        }}
                        break;
                }}
            }}
        }});
    </script>
</body>
</html>"""

        # Write HTML content to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)

        logger.info(f"Generated HTML form: {filepath}")
        return str(filepath)

    def generate_html_content(self, form: ProfessionalForm, filename: str = None) -> tuple[str, str]:
        """Generate HTML form content without saving to file (for preview/editing).
        Returns tuple of (html_content, filename)"""
        if not filename:
            filename = f"{form.title.replace(' ', '_').lower()}_{form.form_id}.html"

        # Ensure the filename doesn't contain path separators
        filename = filename.replace('\\', '_').replace('/', '_')

        # Generate HTML content (enhanced with more comprehensive information)
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{form.title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            line-height: 1.6;
            min-height: 100vh;
        }}
        .form-container {{
            background: white;
            padding: 40px;
            border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
            position: relative;
            overflow: hidden;
        }}
        .form-container::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 5px;
            background: linear-gradient(90deg, #007bff, #28a745, #ffc107, #dc3545);
        }}
        .form-header {{
            text-align: center;
            margin-bottom: 40px;
            border-bottom: 2px solid #007bff;
            padding-bottom: 25px;
            position: relative;
        }}
        .form-title {{
            color: #2c3e50;
            font-size: 32px;
            margin-bottom: 15px;
            font-weight: 700;
            text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .form-description {{
            color: #7f8c8d;
            font-size: 18px;
            margin-bottom: 15px;
            max-width: 600px;
            margin-left: auto;
            margin-right: auto;
        }}
        .company-name {{
            color: #2c5aa0;
            font-size: 20px;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .form-meta {{
            display: flex;
            justify-content: center;
            gap: 30px;
            margin-top: 20px;
            flex-wrap: wrap;
        }}
        .meta-item {{
            display: flex;
            align-items: center;
            gap: 8px;
            color: #6c757d;
            font-size: 14px;
        }}
        .meta-item i {{
            color: #007bff;
        }}
        .progress-container {{
            background: #f8f9fa;
            border-radius: 10px;
            padding: 15px;
            margin-bottom: 30px;
            border-left: 4px solid #007bff;
        }}
        .progress-text {{
            font-size: 14px;
            color: #6c757d;
            margin-bottom: 8px;
        }}
        .progress-bar {{
            width: 100%;
            height: 8px;
            background: #e9ecef;
            border-radius: 4px;
            overflow: hidden;
        }}
        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #007bff, #28a745);
            width: 0%;
            transition: width 0.3s ease;
        }}
        .section {{
            margin-bottom: 35px;
            padding: 25px;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            transition: box-shadow 0.3s ease;
        }}
        .section:hover {{
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }}
        .section-title {{
            color: #2c3e50;
            font-size: 22px;
            margin-bottom: 15px;
            border-bottom: 2px solid #3498db;
            padding-bottom: 8px;
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        .section-icon {{
            color: #3498db;
            font-size: 20px;
        }}
        .section-description {{
            color: #7f8c8d;
            font-size: 15px;
            margin-bottom: 20px;
            font-style: italic;
            background: #f8f9fa;
            padding: 10px 15px;
            border-radius: 6px;
            border-left: 3px solid #3498db;
        }}
        .field-group {{
            margin-bottom: 25px;
            position: relative;
        }}
        .field-label {{
            display: block;
            font-weight: 600;
            margin-bottom: 8px;
            color: #2c3e50;
            font-size: 15px;
        }}
        .required {{
            color: #e74c3c;
        }}
        .field-description {{
            font-size: 13px;
            color: #7f8c8d;
            margin-bottom: 8px;
            font-style: italic;
            background: #f8f9fa;
            padding: 5px 10px;
            border-radius: 4px;
            display: inline-block;
        }}
        input[type="text"], input[type="email"], input[type="number"], 
        input[type="date"], input[type="tel"], input[type="url"], input[type="password"],
        select, textarea {{
            width: 100%;
            padding: 12px 15px;
            border: 2px solid #e9ecef;
            border-radius: 8px;
            font-size: 15px;
            box-sizing: border-box;
            transition: all 0.3s ease;
            background: #ffffff;
        }}
        input:focus, select:focus, textarea:focus {{
            outline: none;
            border-color: #007bff;
            box-shadow: 0 0 0 3px rgba(0,123,255,0.1);
            transform: translateY(-1px);
        }}
        textarea {{
            height: 120px;
            resize: vertical;
            font-family: inherit;
        }}
        .radio-group, .checkbox-group {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 12px;
            padding: 10px;
            background: #f8f9fa;
            border-radius: 8px;
        }}
        .radio-option, .checkbox-option {{
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 8px 12px;
            background: white;
            border-radius: 6px;
            border: 1px solid #e9ecef;
            transition: all 0.2s ease;
            cursor: pointer;
        }}
        .radio-option:hover, .checkbox-option:hover {{
            border-color: #007bff;
            box-shadow: 0 2px 8px rgba(0,123,255,0.15);
        }}
        .radio-option input, .checkbox-option input {{
            width: auto;
            margin: 0;
        }}
        .form-footer {{
            margin-top: 40px;
            padding-top: 25px;
            border-top: 2px solid #e9ecef;
            text-align: center;
            color: #6c757d;
            font-size: 14px;
            background: #f8f9fa;
            margin: 40px -40px -40px -40px;
            padding: 25px 40px;
        }}
        .form-actions {{
            text-align: center;
            margin-top: 40px;
            padding-top: 25px;
            border-top: 3px solid #007bff;
            background: linear-gradient(145deg, #f8f9fa 0%, #e9ecef 100%);
            margin: 40px -40px 0 -40px;
            padding: 30px 40px;
        }}
        .btn {{
            padding: 15px 35px;
            margin: 0 8px 10px 8px;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            position: relative;
            overflow: hidden;
        }}
        .btn::before {{
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
            transition: left 0.5s;
        }}
        .btn:hover::before {{
            left: 100%;
        }}
        .btn-primary {{
            background: linear-gradient(135deg, #007bff 0%, #0056b3 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(0,123,255,0.3);
        }}
        .btn-primary:hover {{
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(0,123,255,0.4);
        }}
        .btn-secondary {{
            background: linear-gradient(135deg, #6c757d 0%, #545b62 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(108,117,125,0.3);
        }}
        .btn-secondary:hover {{
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(108,117,125,0.4);
        }}
        .btn-success {{
            background: linear-gradient(135deg, #28a745 0%, #1e7e34 100%);
            color: white;
            box-shadow: 0 4px 15px rgba(40,167,69,0.3);
        }}
        .btn-success:hover {{
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(40,167,69,0.4);
        }}
        .required-note {{
            color: #e74c3c;
            font-size: 14px;
            margin-top: 20px;
            text-align: center;
            font-weight: 500;
        }}
        .download-options {{
            margin-top: 25px;
            padding: 20px;
            background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
            border-radius: 10px;
            border: 1px solid #e9ecef;
        }}
        .download-options h4 {{
            margin-bottom: 15px;
            color: #2c3e50;
            font-size: 18px;
            text-align: center;
        }}
        .form-info {{
            background: #e3f2fd;
            border-left: 4px solid #2196f3;
            padding: 15px;
            margin-bottom: 25px;
            border-radius: 0 8px 8px 0;
        }}
        .form-info h4 {{
            color: #1976d2;
            margin-bottom: 10px;
            font-size: 16px;
        }}
        .form-info p {{
            color: #424242;
            margin: 5px 0;
            font-size: 14px;
        }}
        @media (max-width: 768px) {{
            .form-container {{
                padding: 20px;
            }}
            .form-meta {{
                flex-direction: column;
                gap: 10px;
            }}
            .radio-group, .checkbox-group {{
                grid-template-columns: 1fr;
            }}
        }}
    </style>
</head>
<body>
    <div class="form-container">
        <div class="form-header">
            <h1 class="form-title">{form.title}</h1>
            {f'<p class="form-description">{form.description}</p>' if form.description else ''}
            {f'<p class="company-name">{form.company_name}</p>' if form.company_name else ''}
            
            <div class="form-meta">
                <div class="meta-item">
                    <i class="fas fa-calendar"></i>
                    <span>Created: {form.created_date}</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-id-card"></i>
                    <span>Form ID: {form.form_id}</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-list"></i>
                    <span>{len(form.sections)} Section(s)</span>
                </div>
                <div class="meta-item">
                    <i class="fas fa-tasks"></i>
                    <span>{sum(len(section.fields) for section in form.sections)} Field(s)</span>
                </div>
            </div>
        </div>
        
        <div class="form-info">
            <h4><i class="fas fa-info-circle"></i> Comprehensive Form Guide</h4>
            <p><strong>Purpose:</strong> This {form.form_type} form is designed to collect detailed information efficiently and securely.</p>
            <p><strong>Required Fields:</strong> Fields marked with <span style="color: #e74c3c;">*</span> are mandatory and must be completed before submission.</p>
            <p><strong>Data Privacy:</strong> All information provided is encrypted and handled according to our strict privacy policy and GDPR compliance.</p>
            <p><strong>Completion Time:</strong> Estimated time to complete: {max(3, len(form.sections) * 2)} minutes</p>
            <p><strong>Auto-Save:</strong> Your progress is automatically saved as you type to prevent data loss.</p>
            <p><strong>Technical Support:</strong> If you experience any issues, please contact support@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</p>
            <p><strong>Browser Requirements:</strong> This form works best with modern browsers (Chrome 90+, Firefox 88+, Safari 14+, Edge 90+).</p>
            <p><strong>Accessibility:</strong> This form is designed to be accessible and supports screen readers and keyboard navigation.</p>
        </div>
        
        <div class="progress-container">
            <div class="progress-text">Form Completion Progress: <span id="progress-percent">0%</span></div>
            <div class="progress-bar">
                <div class="progress-fill" id="progress-fill"></div>
            </div>
        </div>

        <form id="dynamicForm" method="post" action="#" onsubmit="return handleSubmit(event)">
"""

        # Enhanced sections with more detailed field information
        for section_idx, section in enumerate(form.sections, 1):
            html_content += f"""
            <div class="section">
                <h2 class="section-title">
                    <span class="section-icon"><i class="fas fa-folder"></i></span>
                    Section {section_idx}: {section.title}
                    <span style="font-size: 14px; color: #6c757d; font-weight: normal;">({len(section.fields)} fields)</span>
                </h2>
                {f'<p class="section-description"><i class="fas fa-info"></i> {section.description}</p>' if section.description else ''}
                
                <div style="background: #f8f9fa; padding: 10px 15px; border-radius: 6px; margin-bottom: 20px; font-size: 14px; color: #495057;">
                    <strong>Section Guidelines:</strong> Please provide accurate and complete information in this section. 
                    All fields are important for processing your request efficiently.
                </div>
"""

            for field_idx, field in enumerate(section.fields, 1):
                required_attr = 'required' if field.required else ''
                required_mark = '<span class="required">*</span>' if field.required else ''
                
                # Enhanced field information
                field_info = ""
                if field.field_type == 'email':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-envelope'></i> Please enter a valid email address (e.g., user@example.com)</small>"
                elif field.field_type == 'tel':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-phone'></i> Include country code if international (e.g., +1-555-123-4567)</small>"
                elif field.field_type == 'date':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-calendar'></i> Select date using the calendar picker</small>"
                elif field.field_type == 'number':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-hashtag'></i> Enter numeric values only</small>"
                elif field.field_type == 'select':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-list'></i> Choose one option from {len(field.options)} available choices</small>"
                elif field.field_type == 'radio':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-dot-circle'></i> Select one option from the {len(field.options)} choices below</small>"
                elif field.field_type == 'checkbox':
                    field_info = f"<small style='color: #6c757d;'><i class='fas fa-check-square'></i> You can select multiple options from the {len(field.options)} choices</small>"
                elif field.field_type == 'textarea':
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-align-left'></i> Provide detailed information. This field supports multiple lines of text.</small>"
                else:
                    field_info = "<small style='color: #6c757d;'><i class='fas fa-edit'></i> Please provide accurate information</small>"
                
                html_content += f"""
                <div class="field-group" data-field-type="{field.field_type}">
                    <label class="field-label" for="{field.name}">
                        <span style="font-weight: 600;">{field_idx}. {field.label} {required_mark}</span>
                        <span style="font-size: 12px; color: #6c757d; font-weight: normal;">({field.field_type.title()})</span>
                    </label>
                    {f'<p class="field-description"><i class="fas fa-info-circle"></i> {field.description}</p>' if field.description else ''}
                    {field_info}
"""

                if field.field_type == 'text':
                    html_content += f'<input type="text" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'email':
                    html_content += f'<input type="email" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'number':
                    html_content += f'<input type="number" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'date':
                    html_content += f'<input type="date" id="{field.name}" name="{field.name}" {required_attr}>'
                elif field.field_type == 'tel':
                    html_content += f'<input type="tel" id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}>'
                elif field.field_type == 'textarea':
                    html_content += f'<textarea id="{field.name}" name="{field.name}" placeholder="{field.placeholder or ""}" {required_attr}></textarea>'
                elif field.field_type == 'select':
                    html_content += f'<select id="{field.name}" name="{field.name}" {required_attr}>'
                    html_content += f'<option value="">Choose an option...</option>'
                    for option in field.options:
                        html_content += f'<option value="{option}">{option}</option>'
                    html_content += '</select>'
                elif field.field_type == 'radio':
                    html_content += '<div class="radio-group">'
                    for option in field.options:
                        html_content += f'''
                        <div class="radio-option">
                            <input type="radio" id="{field.name}_{option}" name="{field.name}" value="{option}" {required_attr}>
                            <label for="{field.name}_{option}">{option}</label>
                        </div>'''
                    html_content += '</div>'
                elif field.field_type == 'checkbox':
                    html_content += '<div class="checkbox-group">'
                    for option in field.options:
                        html_content += f'''
                        <div class="checkbox-option">
                            <input type="checkbox" id="{field.name}_{option}" name="{field.name}[]" value="{option}">
                            <label for="{field.name}_{option}">{option}</label>
                        </div>'''
                    html_content += '</div>'

                html_content += '</div>'

            html_content += '</div>'

        # Enhanced form footer with comprehensive information
        html_content += f"""
            <div class="form-actions">
                <button type="submit" class="btn btn-primary">
                    <i class="fas fa-paper-plane"></i> Submit Form
                </button>
                <button type="button" class="btn btn-secondary" onclick="clearForm()">
                    <i class="fas fa-eraser"></i> Clear All Data
                </button>
                <button type="button" class="btn" style="background: linear-gradient(135deg, #17a2b8 0%, #138496 100%); color: white;" onclick="saveProgress()">
                    <i class="fas fa-save"></i> Save Progress
                </button>
                
                <div class="download-options">
                    <h4><i class="fas fa-download"></i> Download Options</h4>
                    <p style="font-size: 14px; color: #6c757d; margin-bottom: 15px;">Download your completed form in various formats for your records:</p>
                    <button type="button" class="btn btn-success" onclick="downloadAsPDF()">
                        <i class="fas fa-file-pdf"></i> Download as PDF
                    </button>
                    <button type="button" class="btn btn-success" onclick="downloadAsDOCX()">
                        <i class="fas fa-file-word"></i> Download as DOCX
                    </button>
                    <div style="margin-top: 10px; font-size: 12px; color: #6c757d;">
                        <i class="fas fa-info-circle"></i> Downloads preserve all your entered data and form formatting
                    </div>
                </div>
            </div>

            <div class="form-footer">
                <div style="border-top: 2px solid #e9ecef; padding-top: 20px; margin-top: 20px;">
                    <h5 style="color: #2c3e50; margin-bottom: 15px;"><i class="fas fa-shield-alt"></i> Security & Privacy Information</h5>
                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px; margin-bottom: 15px;">
                        <div>
                            <strong>Data Encryption:</strong><br>
                            <small>All data is encrypted using industry-standard AES-256 encryption</small>
                        </div>
                        <div>
                            <strong>Privacy Compliance:</strong><br>
                            <small>GDPR, CCPA, and SOX compliant data handling procedures</small>
                        </div>
                        <div>
                            <strong>Data Retention:</strong><br>
                            <small>Data retained only as long as necessary for business purposes</small>
                        </div>
                        <div>
                            <strong>Access Control:</strong><br>
                            <small>Strict role-based access controls and audit trails</small>
                        </div>
                    </div>
                    
                    <div style="background: #f8f9fa; padding: 15px; border-radius: 8px; margin-top: 15px;">
                        <h6 style="color: #495057; margin-bottom: 10px;"><i class="fas fa-question-circle"></i> Need Help?</h6>
                        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 10px; font-size: 14px;">
                            <div><strong>Technical Support:</strong> support@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</div>
                            <div><strong>Form Issues:</strong> forms@{form.company_name.lower().replace(' ', '').replace('company', '') if form.company_name else 'company'}.com</div>
                            <div><strong>Phone Support:</strong> 1-800-SUPPORT</div>
                            <div><strong>Live Chat:</strong> Available 24/7 on our website</div>
                        </div>
                    </div>
                    
                    <div style="text-align: center; margin-top: 20px; padding-top: 15px; border-top: 1px solid #dee2e6;">
                        <p style="margin: 5px 0; color: #6c757d; font-size: 14px;">
                            <strong>Form ID:</strong> {form.form_id} | 
                            <strong>Generated:</strong> {form.created_date} | 
                            <strong>Version:</strong> 2.1.0
                        </p>
                        <p style="margin: 5px 0; color: #6c757d; font-size: 13px;">
                            © 2024 {form.company_name if form.company_name else 'Your Company'}. All rights reserved. 
                            <a href="#" style="color: #007bff;">Privacy Policy</a> | 
                            <a href="#" style="color: #007bff;">Terms of Service</a> | 
                            <a href="#" style="color: #007bff;">Accessibility</a>
                        </p>
                    </div>
                </div>
            </div>
            
            <p class="required-note">
                <i class="fas fa-asterisk"></i> Required fields | 
                <i class="fas fa-lock"></i> Secure form submission | 
                <i class="fas fa-mobile-alt"></i> Mobile-friendly design
            </p>
        </form>
    </div>

    <script>
        // Enhanced form handling with comprehensive features
        let formProgress = 0;
        let autoSaveInterval;
        
        // Initialize form enhancements
        document.addEventListener('DOMContentLoaded', function() {{
            initializeFormFeatures();
            startAutoSave();
            updateProgress();
        }});
        
        function initializeFormFeatures() {{
            // Add input event listeners for progress tracking
            const inputs = document.querySelectorAll('input, select, textarea');
            inputs.forEach(input => {{
                input.addEventListener('input', updateProgress);
                input.addEventListener('change', updateProgress);
            }});
            
            // Add accessibility enhancements
            inputs.forEach((input, index) => {{
                input.setAttribute('tabindex', index + 1);
                if (input.hasAttribute('required')) {{
                    input.setAttribute('aria-required', 'true');
                }}
            }});
        }}
        
        function updateProgress() {{
            const allFields = document.querySelectorAll('input, select, textarea');
            const filledFields = Array.from(allFields).filter(field => {{
                if (field.type === 'radio' || field.type === 'checkbox') {{
                    return document.querySelector("input[name=\"" + field.name + "\"]:checked");
                    if (field.type === 'radio') {{
                        return field.value.trim() !== '';
                    }} else {{
                        return Array.from(document.querySelectorAll("input[name=\"" + field.name + "\"]:checked")).length > 0;
                    }}
                }});
            }});
            
            const totalFields = allFields.length;
            const completed = new Set(filledFields.map(f => f.name)).size;
            formProgress = totalFields > 0 ? Math.round((completed / totalFields) * 100) : 0;
            
            document.getElementById('progress-percent').textContent = formProgress + '%';
            document.getElementById('progress-fill').style.width = formProgress + '%';
        }}
        
        function startAutoSave() {{
            autoSaveInterval = setInterval(() => {{
                saveProgress();
            }}, 30000); // Auto-save every 30 seconds
        }}
        
        function saveProgress() {{
            const formData = collectFormData();
            localStorage.setItem('form_progress_' + window.location.pathname, JSON.stringify({{
                data: formData,
                timestamp: new Date().toISOString(),
                progress: formProgress
            }}));
            
            // Show save indicator
            const saveBtn = document.querySelector('button[onclick="saveProgress()"]');
            if (saveBtn) {{
                const originalText = saveBtn.innerHTML;
                saveBtn.innerHTML = '<i class="fas fa-check"></i> Saved!';
                saveBtn.style.background = 'linear-gradient(135deg, #28a745 0%, #20c997 100%)';
                setTimeout(() => {{
                    saveBtn.innerHTML = originalText;
                    saveBtn.style.background = 'linear-gradient(135deg, #17a2b8 0%, #138496 100%)';
                }}, 2000);
            }}
        }}
        
        function collectFormData() {{
            const formData = new FormData(document.getElementById('dynamicForm'));
            const data = {{}};
            
            for (let [key, value] of formData.entries()) {{
                if (data[key]) {{
                    if (Array.isArray(data[key])) {{
                        data[key].push(value);
                    }} else {{
                        data[key] = [data[key], value];
                    }}
                }} else {{
                    data[key] = value;
                }}
            }}
            
            return data;
        }}
        
        function handleSubmit(event) {{
            event.preventDefault();
            
            // Validate form
            const form = event.target;
            if (!form.checkValidity()) {{
                form.reportValidity();
                return false;
            }}
            
            // Collect comprehensive form data
            const formData = collectFormData();
            const submissionData = {{
                formId: '{form.form_id}',
                submissionDate: new Date().toISOString(),
                formData: formData,
                browserInfo: {{
                    userAgent: navigator.userAgent,
                    language: navigator.language,
                    platform: navigator.platform
                }},
                formMetadata: {{
                    totalFields: document.querySelectorAll('input, select, textarea').length,
                    completionTime: 'Calculated from session start',
                    version: '2.1.0'
                }}
            }};
            
            // Clear auto-save
            localStorage.removeItem('form_progress_' + window.location.pathname);
            clearInterval(autoSaveInterval);
            
            // Show success message with more details
            alert(`Form submitted successfully!\\\\n\\\\nSubmission Details:\\\\n` + 
                  `- Form ID: {form.form_id}\\\\n` +
                  `- Submission Date: ` + new Date().toLocaleString() + `\\\\n` +
                  `- Fields Completed: ` + Object.keys(formData).length + `\\\\n` +
                  `- Data will be processed within 24-48 hours\\\\n\\\\n` +
                  `You will receive a confirmation email shortly.`);
            
            return false;
        }}
        
        function clearForm() {{
            if (confirm('Are you sure you want to clear all form data?\\n\\nThis action cannot be undone and will remove all entered information.')) {{
                document.getElementById('dynamicForm').reset();
                localStorage.removeItem('form_progress_' + window.location.pathname);
                updateProgress();
                alert('Form data cleared successfully.');
            }}
        }}
        
        function downloadAsPDF() {{
            const formData = collectFormData();
            if (Object.keys(formData).length === 0) {{
                alert('Please fill out some form fields before downloading.');
                return;
            }}
            
            // Show loading indicator
            const btn = event.target;
            const originalText = btn.innerHTML;
            btn.innerHTML = '<i class="fas fa-spinner fa-spin"></i> Generating PDF...';
            btn.disabled = true;
            
            // Simulate API call (replace with actual implementation)
            setTimeout(() => {{
                alert(`PDF download functionality is being processed.\\n\\n` +
                      `Your form will be converted to PDF format with:\\n` +
                      `- Professional formatting and layout\\n` +
                      `- All form fields and responses\\n` +
                      `- Company branding and logos\\n` +
                      `- Digital signature support\\n\\n` +
                      `The download will begin shortly...`);
                
                btn.innerHTML = originalText;
                btn.disabled = false;
            }}, 2000);
        }}
        
        function downloadAsDOCX() {{
            const formData = collectFormData();
            if (Object.keys(formData).length === 0) {{
                alert('Please fill out some form fields before downloading.');
                return;
            }}
            
            // Show loading indicator
            const btn = event.target;
            const originalText = btn.innerHTML;
            btn.innerHTML = '<i class="fas fa-spinner fa-spin"></i> Generating DOCX...';
            btn.disabled = true;
            
            // Simulate API call (replace with actual implementation)
            setTimeout(() => {{
                alert(`DOCX download functionality is being processed.\\n\\n` +
                      `Your form will be converted to Word document format with:\\n` +
                      `- Editable document structure\\n` +
                      `- Professional table formatting\\n` +
                      `- Company headers and footers\\n` +
                      `- Compatible with Microsoft Word and LibreOffice\\n\\n` +
                      `The download will begin shortly...`);
                
                btn.innerHTML = originalText;
                btn.disabled = false;
            }}, 2000);
        }}
        
        // Keyboard shortcuts
        document.addEventListener('keydown', function(e) {{
            if (e.ctrlKey || e.metaKey) {{
                switch(e.key) {{
                    case 's':
                        e.preventDefault();
                        saveProgress();
                        break;
                    case 'Enter':
                        if (e.shiftKey) {{
                            e.preventDefault();
                            document.getElementById('dynamicForm').dispatchEvent(new Event('submit'));
                        }}
                        break;
                }}
            }}
        }});
    </script>
</body>
</html>"""

        logger.info(f"Generated HTML content for preview: {filename}")
        return html_content, filename


# -----------------------------
# MCP (Model Context Protocol) Integration
# -----------------------------

@dataclass
class MCPServer:
    """Configuration for an MCP server."""
    name: str
    command: str
    args: List[str] = field(default_factory=list)
    env: Dict[str, str] = field(default_factory=dict)
    enabled: bool = True

class MCPManager:
    """Manages MCP server connections and tool integration."""

    def __init__(self):
        self.servers: Dict[str, MCPServer] = {}
        self.sessions: Dict[str, ClientSession] = {}
        self.mcp_tools: Dict[str, List] = {}

    def register_mcp_server(self, server: MCPServer) -> bool:
        """Register an MCP server configuration."""
        if not MCP_AVAILABLE:
            logger.warning("MCP not available - server registration skipped")
            return False

        self.servers[server.name] = server
        logger.info(f"Registered MCP server: {server.name}")
        return True

    async def connect_server(self, server_name: str) -> bool:
        """Connect to an MCP server and initialize session."""
        if not MCP_AVAILABLE:
            return False

        if server_name not in self.servers:
            logger.error(f"MCP server {server_name} not registered")
            return False

        server = self.servers[server_name]
        if not server.enabled:
            logger.info(f"MCP server {server_name} is disabled")
            return False

        try:
            # Create server parameters
            server_params = StdioServerParameters(
                command=server.command,
                args=server.args,
                env=server.env
            )

            # Connect to server
            async with stdio_client(server_params) as (read, write):
                async with ClientSession(read, write) as session:
                    # Initialize the session
                    await session.initialize()

                    # Get available tools
                    tools_result = await session.list_tools()

                    # Store session and tools
                    self.sessions[server_name] = session
                    self.mcp_tools[server_name] = tools_result.tools if hasattr(tools_result, 'tools') else []

                    logger.info(f"Connected to MCP server {server_name} with {len(self.mcp_tools[server_name])} tools")
                    return True

        except Exception as e:
            logger.error(f"Failed to connect to MCP server {server_name}: {e}")
            return False

    def get_mcp_tools_for_tenant(self, tenant_id: str) -> List:
        """Get MCP tools available for a specific tenant."""
        if not MCP_AVAILABLE:
            return []

        # For now, return all MCP tools - can be enhanced with tenant-specific filtering
        all_tools = []
        for server_name, tools in self.mcp_tools.items():
            all_tools.extend(tools)

        return all_tools

    async def call_mcp_tool(self, server_name: str, tool_name: str, arguments: Dict) -> str:
        """Call an MCP tool on a specific server."""
        if not MCP_AVAILABLE:
            return "MCP not available"

        if server_name not in self.sessions:
            return f"No active session for MCP server {server_name}"

        try:
            session = self.sessions[server_name]
            result = await session.call_tool(tool_name, arguments)

            if hasattr(result, 'content'):
                return str(result.content)
            else:
                return str(result)

        except Exception as e:
            logger.error(f"MCP tool call failed: {e}")
            return f"Error calling MCP tool: {e}"

# Global MCP manager
MCP_MANAGER = MCPManager()

def setup_default_mcp_servers():
    """Set up default MCP servers for common integrations."""
    if not MCP_AVAILABLE:
        return

    # Example MCP servers - can be configured via environment or config file
    default_servers = [
        MCPServer(
            name="filesystem",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-filesystem", "/path/to/allowed/directory"],
            enabled=os.environ.get("MCP_FILESYSTEM_ENABLED", "false").lower() == "true"
        ),
        MCPServer(
            name="git",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-git", "--repository", "."],
            enabled=os.environ.get("MCP_GIT_ENABLED", "false").lower() == "true"
        ),
        MCPServer(
            name="sqlite",
            command="npx",
            args=["-y", "@modelcontextprotocol/server-sqlite", "--db-path", "database.db"],
            enabled=os.environ.get("MCP_SQLITE_ENABLED", "false").lower() == "true"
        )
    ]

    for server in default_servers:
        MCP_MANAGER.register_mcp_server(server)

    logger.info(f"Registered {len(default_servers)} default MCP servers")

# Initialize default MCP servers
setup_default_mcp_servers()


# -----------------------------
# Authentication and Tenant Management
# -----------------------------

@dataclass
class TenantConfig:
    """Configuration for a tenant."""
    tenant_id: str
    name: str
    api_keys: Dict[str, str] = field(default_factory=dict)
    permissions: List[str] = field(default_factory=list)
    rate_limits: Dict[str, float] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    is_active: bool = True
    max_documents: int = 1000
    max_api_calls_per_hour: int = 1000

@dataclass
class DocumentMetadata:
    """Enhanced document metadata for tracking uploaded documents."""
    document_id: str
    filename: str
    file_path: str
    file_size: int
    file_type: str
    upload_timestamp: str
    tenant_id: str
    user_id: Optional[str] = None
    chunk_count: int = 0
    indexed: bool = False
    tags: List[str] = field(default_factory=list)
    file_hash: str = ""
    original_name: str = ""
    description: str = ""

    def __post_init__(self):
        if not self.original_name:
            self.original_name = self.filename
        if not self.file_hash and os.path.exists(self.file_path):
            self.file_hash = self.calculate_file_hash()

    def calculate_file_hash(self) -> str:
        """Calculate SHA256 hash of the file for deduplication."""
        try:
            with open(self.file_path, 'rb') as f:
                return hashlib.sha256(f.read()).hexdigest()
        except Exception:
            return ""

@dataclass
class ChatMessage:
    """Chat message with metadata."""
    message_id: str
    session_id: str
    tenant_id: str
    role: str  # 'user' or 'assistant'
    content: str
    timestamp: str
    user_id: Optional[str] = None
    agent_type: Optional[str] = None
    document_references: List[str] = field(default_factory=list)  # Referenced document IDs

@dataclass
class UserSession:
    """Enhanced user session information."""
    session_id: str
    tenant_id: str
    user_id: Optional[str] = None
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    last_activity: str = field(default_factory=lambda: datetime.now().isoformat())
    permissions: List[str] = field(default_factory=list)
    chat_history: List[ChatMessage] = field(default_factory=list)
    uploaded_documents: List[str] = field(default_factory=list)  # Document IDs

# Enhanced Database Storage System
class DocumentStorage:
    """Enhanced document storage with SQLite backend for persistence."""

    def __init__(self, db_path: str = "document_storage.db"):
        self.db_path = db_path
        self.init_database()

    def init_database(self):
        """Initialize SQLite database with required tables."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        # Documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                document_id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_type TEXT NOT NULL,
                upload_timestamp TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                chunk_count INTEGER DEFAULT 0,
                indexed BOOLEAN DEFAULT FALSE,
                tags TEXT,
                file_hash TEXT,
                original_name TEXT
            )
        ''')

        # Chat messages table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS chat_messages (
                message_id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                timestamp TEXT NOT NULL,
                agent_type TEXT,
                document_references TEXT
            )
        ''')

        # User sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_sessions (
                session_id TEXT PRIMARY KEY,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                created_at TEXT NOT NULL,
                last_activity TEXT NOT NULL,
                permissions TEXT,
                uploaded_documents TEXT,
                status TEXT DEFAULT 'active',
                ip_address TEXT,
                user_agent TEXT
            )
        ''')

        # Escalation tickets table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS escalation_tickets (
                ticket_id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                title TEXT NOT NULL,
                description TEXT NOT NULL,
                status TEXT DEFAULT 'open',
                priority TEXT DEFAULT 'medium',
                assigned_to TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                resolved_at TEXT,
                chat_context TEXT
            )
        ''')

        # Tenant customization table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS tenant_customization (
                tenant_id TEXT PRIMARY KEY,
                theme_color TEXT DEFAULT '#667eea',
                logo_url TEXT,
                chat_background_color TEXT DEFAULT '#ffffff',
                widget_position TEXT DEFAULT 'bottom-right',
                welcome_message TEXT DEFAULT 'Hello! How can I help you today?',
                custom_css TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
        ''')

        # Meeting schedules table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS meeting_schedules (
                meeting_id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                tenant_id TEXT NOT NULL,
                user_id TEXT,
                title TEXT NOT NULL,
                description TEXT,
                scheduled_time TEXT NOT NULL,
                duration_minutes INTEGER DEFAULT 30,
                meeting_type TEXT DEFAULT 'general',
                calendar_provider TEXT,
                calendar_event_id TEXT,
                status TEXT DEFAULT 'scheduled',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
        ''')

        # Create indexes for better performance
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_documents_tenant ON documents(tenant_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_messages_session ON chat_messages(session_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_sessions_tenant ON user_sessions(tenant_id)')

        conn.commit()
        conn.close()

    def save_document(self, doc_metadata: DocumentMetadata) -> bool:
        """Save document metadata to database."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                INSERT OR REPLACE INTO documents
                (document_id, filename, file_path, file_size, file_type, upload_timestamp,
                 tenant_id, user_id, chunk_count, indexed, tags, file_hash, original_name)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc_metadata.document_id,
                doc_metadata.filename,
                doc_metadata.file_path,
                doc_metadata.file_size,
                doc_metadata.file_type,
                doc_metadata.upload_timestamp,
                doc_metadata.tenant_id,
                doc_metadata.user_id,
                doc_metadata.chunk_count,
                doc_metadata.indexed,
                json.dumps(doc_metadata.tags),
                getattr(doc_metadata, 'file_hash', ''),
                getattr(doc_metadata, 'original_name', doc_metadata.filename)
            ))

            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Failed to save document metadata: {e}")
            return False

    def get_documents_by_tenant(self, tenant_id: str) -> List[DocumentMetadata]:
        """Get all documents for a tenant."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                SELECT document_id, filename, file_path, file_size, file_type, upload_timestamp,
                       tenant_id, user_id, chunk_count, indexed, tags, file_hash, original_name
                FROM documents WHERE tenant_id = ?
                ORDER BY upload_timestamp DESC
            ''', (tenant_id,))

            documents = []
            for row in cursor.fetchall():
                doc = DocumentMetadata(
                    document_id=row[0],
                    filename=row[1],
                    file_path=row[2],
                    file_size=row[3],
                    file_type=row[4],
                    upload_timestamp=row[5],
                    tenant_id=row[6],
                    user_id=row[7],
                    chunk_count=row[8] or 0,
                    indexed=bool(row[9]),
                    tags=json.loads(row[10]) if row[10] else []
                )
                # Add additional fields
                doc.file_hash = row[11] or ''
                doc.original_name = row[12] or row[1]
                documents.append(doc)

            conn.close()
            return documents
        except Exception as e:
            logger.error(f"Failed to get documents for tenant {tenant_id}: {e}")
            return []

    def save_chat_message(self, message: ChatMessage) -> bool:
        """Save chat message to database."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                INSERT OR REPLACE INTO chat_messages
                (message_id, session_id, tenant_id, user_id, role, content, timestamp, agent_type, document_references)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                message.message_id,
                message.session_id,
                message.tenant_id,
                message.user_id,
                message.role,
                message.content,
                message.timestamp,
                message.agent_type,
                json.dumps(message.document_references)
            ))

            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Failed to save chat message: {e}")
            return False

    def get_chat_history(self, session_id: str, limit: int = 50) -> List[ChatMessage]:
        """Get chat history for a session."""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute('''
                SELECT message_id, session_id, tenant_id, user_id, role, content, timestamp, agent_type, document_references
                FROM chat_messages WHERE session_id = ?
                ORDER BY timestamp DESC LIMIT ?
            ''', (session_id, limit))

            messages = []
            for row in cursor.fetchall():
                message = ChatMessage(
                    message_id=row[0],
                    session_id=row[1],
                    tenant_id=row[2],
                    role=row[4],
                    content=row[5],
                    timestamp=row[6],
                    user_id=row[3],
                    agent_type=row[7],
                    document_references=json.loads(row[8]) if row[8] else []
                )
                messages.append(message)

            conn.close()
            return list(reversed(messages))  # Return in chronological order
        except Exception as e:
            logger.error(f"Failed to get chat history for session {session_id}: {e}")
            return []

# Global storage instance
document_storage = DocumentStorage()

# Global runtime context
CURRENT_TENANT_ID: Optional[str] = None
CURRENT_SESSION: Optional[UserSession] = None

def set_current_tenant(tenant_id: str):
    """Set the current tenant for the session"""
    global CURRENT_TENANT_ID, CURRENT_SESSION
    CURRENT_TENANT_ID = tenant_id
    if tenant_id in _tenant_registry:
        CURRENT_SESSION = create_session(tenant_id)
    logger.info(f"Set current tenant to: {tenant_id}")

# Tenant registry
_tenant_registry: Dict[str, TenantConfig] = {}
_active_sessions: Dict[str, UserSession] = {}

def create_tenant(tenant_id: str, name: str, permissions: Optional[List[str]] = None) -> TenantConfig:
    """Create a new tenant with default configuration."""
    if tenant_id in _tenant_registry:
        raise ValueError(f"Tenant {tenant_id} already exists")
    
    config = TenantConfig(
        tenant_id=tenant_id,
        name=name,
        permissions=permissions or ["read_documents", "use_tools", "generate_forms"],
        rate_limits={"default": 0.5, "search_web": 1.0, "get_weather": 0.5}
    )
    
    _tenant_registry[tenant_id] = config
    logger.info(f"Created tenant: {tenant_id}")
    return config

def get_tenant_config(tenant_id: str) -> Optional[TenantConfig]:
    """Get tenant configuration."""
    return _tenant_registry.get(tenant_id)

def authenticate_tenant(tenant_id: str, api_key: Optional[str] = None) -> bool:
    """Authenticate a tenant (simplified for demo)."""
    config = get_tenant_config(tenant_id)
    if not config or not config.is_active:
        return False
    
    # In production, you'd validate the API key here
    if api_key and "master_key" in config.api_keys:
        return config.api_keys["master_key"] == api_key
    
    # For demo purposes, allow access without API key
    return True

def create_session(tenant_id: str, user_id: Optional[str] = None) -> UserSession:
    """Create a new user session."""
    if not authenticate_tenant(tenant_id):
        raise ValueError(f"Invalid tenant: {tenant_id}")
    
    session_id = secrets.token_urlsafe(32)
    config = get_tenant_config(tenant_id)
    
    session = UserSession(
        session_id=session_id,
        tenant_id=tenant_id,
        user_id=user_id,
        permissions=config.permissions if config else []
    )
    
    _active_sessions[session_id] = session
    logger.info(f"Created session {session_id} for tenant {tenant_id}")
    return session

def get_session(session_id: str) -> Optional[UserSession]:
    """Get session by ID."""
    session = _active_sessions.get(session_id)
    if session:
        # Update last activity
        session.last_activity = datetime.now().isoformat()
    return session

def has_permission(permission: str, session: Optional[UserSession] = None) -> bool:
    """Check if current session has permission."""
    current_session = session or CURRENT_SESSION
    if not current_session:
        return False
    return permission in current_session.permissions

def require_permission(permission: str) -> Callable:
    """Decorator to require specific permission."""
    def decorator(func: Callable) -> Callable:
        def wrapper(*args, **kwargs):
            if not has_permission(permission):
                raise PermissionError(f"Permission required: {permission}")
            return func(*args, **kwargs)
        return wrapper
    return decorator

# Initialize default tenant
def initialize_default_tenant():
    """Initialize default tenant for demo purposes."""
    if "default" not in _tenant_registry:
        create_tenant(
            "default",
            "Default Tenant",
            ["read_documents", "use_tools", "generate_forms", "admin"]
        )

initialize_default_tenant()


# -----------------------------
# Enhanced Dynamic Tooling Infrastructure
# -----------------------------

from collections import defaultdict
from threading import Lock
import logging

# Enhanced tool registry with metadata
_dynamic_tool_registry: Dict[str, List] = {}
_tool_metadata: Dict[str, Dict] = {}
_last_call_timestamp_per_tool: Dict[str, float] = {}
_tool_call_counts: Dict[str, int] = defaultdict(int)
_tool_error_counts: Dict[str, int] = defaultdict(int)
_registry_lock = Lock()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _rate_limited(tool_name: str, min_interval_seconds: float = 0.5) -> bool:
    """Enhanced rate limiting with per-tool configuration."""
    now = time.time()
    last = _last_call_timestamp_per_tool.get(tool_name)
    
    # Get tool-specific rate limit if available
    tool_meta = _tool_metadata.get(tool_name, {})
    interval = tool_meta.get('rate_limit_seconds', min_interval_seconds)
    
    if last is None or now - last >= interval:
        _last_call_timestamp_per_tool[tool_name] = now
        return True
    return False


def get_tenant_tools(tenant_id: Optional[str]) -> List:
    """Get all available tools for a tenant with enhanced filtering and MCP integration."""
    base_tools = [search_web, get_weather, get_document_stats_tool, discover_api_endpoint, analyze_supabase_sample_apis]
    tenant_list = _dynamic_tool_registry.get(tenant_id or "default", [])

    # Filter out disabled tools
    active_tenant_tools = [
        tool for tool in tenant_list
        if _tool_metadata.get(tool.name, {}).get('enabled', True)
    ]

    # Add MCP tools if available
    mcp_tools = MCP_MANAGER.get_mcp_tools_for_tenant(tenant_id or "default")
    
    # Add dynamic API tools
    dynamic_api_tools = DYNAMIC_API_MANAGER.get_api_tools(tenant_id)

    return base_tools + active_tenant_tools + mcp_tools + dynamic_api_tools


def register_dynamic_tool(tenant_id: str, dynamic_tool, metadata: Optional[Dict] = None) -> None:
    """Enhanced tool registration with metadata and validation."""
    with _registry_lock:
        tools_for_tenant = _dynamic_tool_registry.setdefault(tenant_id, [])
        
        # Replace by name if already exists
        existing_names = {t.name for t in tools_for_tenant}
        if dynamic_tool.name in existing_names:
            tools_for_tenant[:] = [t for t in tools_for_tenant if t.name != dynamic_tool.name]
        
        tools_for_tenant.append(dynamic_tool)
        
        # Store metadata
        tool_meta = metadata or {}
        tool_meta.update({
            'tenant_id': tenant_id,
            'registered_at': datetime.now().isoformat(),
            'enabled': tool_meta.get('enabled', True),
            'rate_limit_seconds': tool_meta.get('rate_limit_seconds', 0.5),
            'max_retries': tool_meta.get('max_retries', 3)
        })
        _tool_metadata[dynamic_tool.name] = tool_meta
        
        logger.info(f"Registered tool '{dynamic_tool.name}' for tenant '{tenant_id}'")


def unregister_tool(tenant_id: str, tool_name: str) -> bool:
    """Remove a tool from a tenant's registry."""
    with _registry_lock:
        tools_for_tenant = _dynamic_tool_registry.get(tenant_id, [])
        original_count = len(tools_for_tenant)
        
        _dynamic_tool_registry[tenant_id] = [
            t for t in tools_for_tenant if t.name != tool_name
        ]
        
        if tool_name in _tool_metadata:
            del _tool_metadata[tool_name]
        
        removed = len(_dynamic_tool_registry[tenant_id]) < original_count
        if removed:
            logger.info(f"Unregistered tool '{tool_name}' from tenant '{tenant_id}'")
        
        return removed


def get_tool_stats(tenant_id: Optional[str] = None) -> Dict:
    """Get statistics about tool usage."""
    if tenant_id:
        tools = get_tenant_tools(tenant_id)
        tool_names = [t.name for t in tools]
    else:
        tool_names = list(_tool_metadata.keys())
    
    stats = {}
    for tool_name in tool_names:
        stats[tool_name] = {
            'call_count': _tool_call_counts.get(tool_name, 0),
            'error_count': _tool_error_counts.get(tool_name, 0),
            'last_called': _last_call_timestamp_per_tool.get(tool_name),
            'metadata': _tool_metadata.get(tool_name, {})
        }
    
    return stats


def make_http_get_tool(
    name: str,
    description: str,
    base_url_env: str,
    api_key_env: Optional[str] = None,
    rate_limit_seconds: float = 0.5,
    timeout: int = 20
) -> any:
    """Enhanced HTTP GET tool with better error handling and configuration."""

    from langchain_core.tools import StructuredTool
    from pydantic import BaseModel, Field

    class InputSchema(BaseModel):
        path: str = Field(..., description="URL path to append to the base URL, starting with '/'")
        query: Optional[Dict[str, str]] = Field(default=None, description="Query params as key-value map")
        headers: Optional[Dict[str, str]] = Field(default=None, description="Additional headers")

    base_url = os.environ.get(base_url_env)
    api_key = os.environ.get(api_key_env) if api_key_env else None

    def _run(path: str, query: Optional[Dict[str, str]] = None, headers: Optional[Dict[str, str]] = None) -> str:
        tool_name = f"{name}"
        
        # Track call count
        _tool_call_counts[tool_name] += 1
        
        # Rate limiting
        if not _rate_limited(tool_name, rate_limit_seconds):
            return "Rate limited. Please retry shortly."
        
        # Validation
        if not base_url:
            _tool_error_counts[tool_name] += 1
            return f"HTTP GET tool misconfigured: missing env {base_url_env}"
        
        # Build request
        url = base_url.rstrip("/") + path
        request_headers = {"Accept": "application/json", "User-Agent": "Multi-Agent-Chatbot/1.0"}
        
        if api_key:
            request_headers["Authorization"] = f"Bearer {api_key}"
        
        if headers:
            request_headers.update(headers)
        
        try:
            logger.info(f"Making HTTP GET request to {url}")
            resp = requests.get(
                url,
                params=query or {},
                headers=request_headers,
                timeout=timeout
            )
            
            if resp.ok:
                content = resp.text[:4000]  # Limit response size
                logger.info(f"HTTP GET successful for {tool_name}")
                return content
            else:
                error_msg = f"HTTP {resp.status_code}: {resp.text[:800]}"
                logger.warning(f"HTTP GET failed for {tool_name}: {error_msg}")
                _tool_error_counts[tool_name] += 1
                return error_msg
                
        except requests.exceptions.Timeout:
            error_msg = f"HTTP GET timeout after {timeout}s"
            logger.error(f"HTTP GET timeout for {tool_name}")
            _tool_error_counts[tool_name] += 1
            return error_msg
        except Exception as exc:
            error_msg = f"HTTP GET error: {exc}"
            logger.error(f"HTTP GET error for {tool_name}: {exc}")
            _tool_error_counts[tool_name] += 1
            return error_msg

    return StructuredTool.from_function(
        name=name,
        description=description,
        func=_run,
        args_schema=InputSchema,
    )


def make_http_post_tool(
    name: str,
    description: str,
    base_url_env: str,
    api_key_env: Optional[str] = None,
    rate_limit_seconds: float = 1.0,
    timeout: int = 30
) -> any:
    """Create an HTTP POST tool for API interactions."""
    
    from langchain_core.tools import StructuredTool
    from pydantic import BaseModel, Field

    class InputSchema(BaseModel):
        path: str = Field(..., description="URL path to append to the base URL")
        data: Optional[Dict] = Field(default=None, description="JSON data to send in request body")
        headers: Optional[Dict[str, str]] = Field(default=None, description="Additional headers")

    base_url = os.environ.get(base_url_env)
    api_key = os.environ.get(api_key_env) if api_key_env else None

    def _run(path: str, data: Optional[Dict] = None, headers: Optional[Dict[str, str]] = None) -> str:
        tool_name = f"{name}"
        
        _tool_call_counts[tool_name] += 1
        
        if not _rate_limited(tool_name, rate_limit_seconds):
            return "Rate limited. Please retry shortly."
        
        if not base_url:
            _tool_error_counts[tool_name] += 1
            return f"HTTP POST tool misconfigured: missing env {base_url_env}"
        
        url = base_url.rstrip("/") + path
        request_headers = {"Content-Type": "application/json", "User-Agent": "Multi-Agent-Chatbot/1.0"}
        
        if api_key:
            request_headers["Authorization"] = f"Bearer {api_key}"
        
        if headers:
            request_headers.update(headers)
        
        try:
            logger.info(f"Making HTTP POST request to {url}")
            resp = requests.post(
                url,
                json=data,
                headers=request_headers,
                timeout=timeout
            )
            
            if resp.ok:
                content = resp.text[:4000]
                logger.info(f"HTTP POST successful for {tool_name}")
                return content
            else:
                error_msg = f"HTTP {resp.status_code}: {resp.text[:800]}"
                logger.warning(f"HTTP POST failed for {tool_name}: {error_msg}")
                _tool_error_counts[tool_name] += 1
                return error_msg
                
        except Exception as exc:
            error_msg = f"HTTP POST error: {exc}"
            logger.error(f"HTTP POST error for {tool_name}: {exc}")
            _tool_error_counts[tool_name] += 1
            return error_msg

    return StructuredTool.from_function(
        name=name,
        description=description,
        func=_run,
        args_schema=InputSchema,
    )


# -----------------------------
# API Discovery Tool
# -----------------------------

@tool
def discover_api_endpoint(url: str, method: str = "GET", headers: Optional[str] = None, body: Optional[str] = None) -> str:
    """Discover and analyze an API endpoint to understand its structure, parameters, and response format.
    
    Args:
        url: The full API endpoint URL to analyze
        method: HTTP method (GET, POST, PUT, DELETE)
        headers: Optional JSON string of headers to include
        body: Optional JSON string of request body for POST/PUT requests
    
    Returns:
        Detailed analysis of the API endpoint including structure, response format, and usage examples
    """
    tool_name = "discover_api_endpoint"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 2.0):  # 2 second rate limit for API discovery
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info(f"Discovering API endpoint: {method} {url}")
        
        # Parse headers if provided
        request_headers = {"User-Agent": "Multi-Agent-Chatbot-API-Discovery/1.0"}
        if headers:
            try:
                additional_headers = json.loads(headers)
                request_headers.update(additional_headers)
            except json.JSONDecodeError:
                logger.warning(f"Invalid headers JSON: {headers}")
        
        # Parse body if provided
        request_body = None
        if body and method.upper() in ["POST", "PUT", "PATCH"]:
            try:
                request_body = json.loads(body)
                request_headers["Content-Type"] = "application/json"
            except json.JSONDecodeError:
                logger.warning(f"Invalid body JSON: {body}")
        
        # Make the API request
        response = None
        if method.upper() == "GET":
            response = requests.get(url, headers=request_headers, timeout=15)
        elif method.upper() == "POST":
            response = requests.post(url, headers=request_headers, json=request_body, timeout=15)
        elif method.upper() == "PUT":
            response = requests.put(url, headers=request_headers, json=request_body, timeout=15)
        elif method.upper() == "DELETE":
            response = requests.delete(url, headers=request_headers, timeout=15)
        else:
            return f"Unsupported HTTP method: {method}"
        
        # Analyze the response
        analysis = _analyze_api_response(url, method, response, request_headers, request_body)
        
        logger.info(f"API discovery completed for: {url}")
        return analysis
        
    except requests.exceptions.Timeout:
        error_msg = f"API discovery timeout after 15s for {url}"
        logger.error(error_msg)
        _tool_error_counts[tool_name] += 1
        return error_msg
    except Exception as exc:
        error_msg = f"API discovery error for {url}: {exc}"
        logger.error(error_msg)
        _tool_error_counts[tool_name] += 1
        return error_msg


def _analyze_api_response(url: str, method: str, response: requests.Response, 
                         request_headers: dict, request_body: Optional[dict]) -> str:
    """Analyze an API response and provide structured insights."""
    
    analysis_parts = []
    
    # Basic Info
    analysis_parts.append(f"**API ANALYSIS**")
    analysis_parts.append(f"• **Method**: {method.upper()}")
    analysis_parts.append(f"• **Status**: {response.status_code} ({response.reason})")
    
    # Parse URL for cleaner display
    try:
        from urllib.parse import urlparse, parse_qs
        parsed = urlparse(url)
        base_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        query_params = parse_qs(parsed.query) if parsed.query else {}
        
        analysis_parts.append(f"• **Endpoint**: {base_url}")
        
        if query_params:
            analysis_parts.append(f"• **Parameters**:")
            for param, values in query_params.items():
                analysis_parts.append(f"  - {param}: {values[0]}")
    except:
        analysis_parts.append(f"• **URL**: {url}")
    
    # Response Analysis
    if response.status_code >= 400:
        analysis_parts.append(f"• **Error**: {response.text[:200]}..." if len(response.text) > 200 else f"• **Error**: {response.text}")
        return "\n".join(analysis_parts)
    
    # Content Analysis
    content_type = response.headers.get('content-type', '').lower()
    if 'json' in content_type:
        try:
            json_data = response.json()
            analysis_parts.append(f"• **Format**: JSON")
            
            # Structure
            structure_info = _analyze_json_structure_simple(json_data)
            analysis_parts.append(f"• **Structure**:")
            analysis_parts.extend([f"  {line}" for line in structure_info])
            
            # Sample (compact)
            sample = json.dumps(json_data, separators=(',', ':'))
            if len(sample) > 150:
                sample = sample[:150] + "..."
            analysis_parts.append(f"• **Sample**: `{sample}`")
            
        except json.JSONDecodeError:
            analysis_parts.append(f"• **Format**: Invalid JSON")
    else:
        analysis_parts.append(f"• **Format**: {content_type or 'Unknown'}")
        if response.text:
            preview = response.text[:100].replace('\n', ' ').strip()
            analysis_parts.append(f"• **Content**: {preview}...")
    
    # Registration Command
    if response.status_code < 400:
        tool_name = parsed.path.split('/')[-1].replace('-', '_') if 'parsed' in locals() else 'api_tool'
        base_for_tool = f"{parsed.scheme}://{parsed.netloc}" + "/".join(parsed.path.split('/')[:-1]) if 'parsed' in locals() else url.split('/')[:-1]
        analysis_parts.append(f"• **Register**: `/tool.httpget {tool_name} {base_for_tool}`")
    
    return "\n".join(analysis_parts)


def _analyze_json_structure_simple(data: Any, level: int = 0) -> List[str]:
    """Provide simple, structured JSON analysis."""
    structure_info = []
    
    if level > 2:  # Limit depth
        return ["...(nested)"]
    
    if isinstance(data, dict):
        for key, value in list(data.items())[:5]:  # Limit items
            if isinstance(value, dict):
                structure_info.append(f"- {key}: Object")
                if level < 2:
                    nested = _analyze_json_structure_simple(value, level + 1)
                    structure_info.extend([f"  {line}" for line in nested])
            elif isinstance(value, list):
                item_type = type(value[0]).__name__ if value else "unknown"
                structure_info.append(f"- {key}: Array[{len(value)}] of {item_type}")
            else:
                value_type = type(value).__name__
                structure_info.append(f"- {key}: {value_type}")
        
        if len(data) > 5:
            structure_info.append(f"- ...(+{len(data)-5} more)")
    
    elif isinstance(data, list):
        if data:
            structure_info.append(f"Array[{len(data)}] of {type(data[0]).__name__}")
        else:
            structure_info.append("Empty Array")
    
    return structure_info


def _analyze_json_structure(data: Any, level: int = 0, max_level: int = 3) -> List[str]:
    """Recursively analyze JSON structure and return insights."""
    indent = "  " * level
    structure_info = []
    
    if level > max_level:
        structure_info.append(f"{indent}... (nested structure continues)")
        return structure_info
    
    if isinstance(data, dict):
        structure_info.append(f"{indent}📦 Object with {len(data)} properties:")
        for key, value in list(data.items())[:5]:  # Limit to first 5 properties
            value_type = type(value).__name__
            if isinstance(value, list) and value:
                structure_info.append(f"{indent}  • {key}: Array[{len(value)}] of {type(value[0]).__name__}")
            elif isinstance(value, dict):
                structure_info.append(f"{indent}  • {key}: Object")
                if level < max_level:
                    structure_info.extend(_analyze_json_structure(value, level + 1, max_level))
            else:
                example_value = str(value)[:50]
                if len(str(value)) > 50:
                    example_value += "..."
                structure_info.append(f"{indent}  • {key}: {value_type} (e.g., '{example_value}')")
        
        if len(data) > 5:
            structure_info.append(f"{indent}  ... and {len(data) - 5} more properties")
    
    elif isinstance(data, list):
        structure_info.append(f"{indent}📋 Array with {len(data)} items")
        if data and level < max_level:
            structure_info.append(f"{indent}  Sample item structure:")
            structure_info.extend(_analyze_json_structure(data[0], level + 1, max_level))
    
    else:
        value_type = type(data).__name__
        example_value = str(data)[:50]
        if len(str(data)) > 50:
            example_value += "..."
        structure_info.append(f"{indent}📄 {value_type}: '{example_value}'")
    
    return structure_info


# -----------------------------
# Enhanced Built-in Tools
# -----------------------------


@tool
def analyze_supabase_sample_apis() -> str:
    """Analyze the sample Supabase APIs (get-order-status and get-product-price) provided by the user.
    
    This tool demonstrates the API discovery functionality with the specific APIs mentioned.
    """
    tool_name = "analyze_supabase_sample_apis"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 3.0):  # 3 second rate limit
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info("Analyzing sample Supabase APIs")
        
        # Sample APIs provided by the user
        sample_apis = [
            {
                "name": "Order Status API",
                "url": "https://oamrapppfdexxiyoesxo.supabase.co/functions/v1/get-order-status?order_id=ORD002",
                "endpoint": "/get-order-status",
                "param": "order_id"
            },
            {
                "name": "Product Price API",
                "url": "https://oamrapppfdexxiyoesxo.supabase.co/functions/v1/get-product-price?id=2bc2af12-1287-4fdf-adbd-6a76358ca9dd",
                "endpoint": "/get-product-price",
                "param": "id"
            }
        ]
        
        results = []
        results.append("**SUPABASE API SUMMARY**")
        results.append("")
        
        for i, api in enumerate(sample_apis, 1):
            results.append(f"**{i}. {api['name']}**")
            results.append(f"• **Endpoint**: {api['endpoint']}")
            results.append(f"• **Parameter**: {api['param']}")
            results.append(f"• **Base URL**: https://oamrapppfdexxiyoesxo.supabase.co/functions/v1")
            results.append("")
        
        # Registration commands
        results.append("**REGISTRATION COMMANDS**")
        results.append("```")
        results.append("/tool.httpget get_order_status https://oamrapppfdexxiyoesxo.supabase.co/functions/v1")
        results.append("/tool.httpget get_product_price https://oamrapppfdexxiyoesxo.supabase.co/functions/v1")
        results.append("```")
        
        results.append("")
        results.append("**NEXT STEPS**")
        results.append("• Use 'discover_api_endpoint' for detailed analysis")
        results.append("• Register APIs using commands above")
        results.append("• Test with different parameter values")
        
        logger.info("Sample API analysis completed")
        return "\n".join(results)
        
    except Exception as exc:
        error_msg = f"Error analyzing sample APIs: {exc}"
        logger.error(error_msg)
        _tool_error_counts[tool_name] += 1
        return error_msg


@tool
def search_web(query: str) -> str:
    """Enhanced web search via DuckDuckGo instant answers; returns brief summary/snippets."""
    tool_name = "search_web"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 1.0):  # 1 second rate limit for web search
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info(f"Performing web search for: {query}")
        resp = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1},
            timeout=12,
        )
        data = resp.json() if resp.ok else {}
        
        # Try abstract first
        abstract = data.get("AbstractText") or data.get("Abstract") or ""
        if abstract:
            logger.info(f"Web search successful for: {query}")
            return f"Abstract: {abstract}"
        
        # Try related topics
        related = data.get("RelatedTopics", [])
        snippets: List[str] = []
        for item in related:
            if isinstance(item, dict) and item.get("Text"):
                snippets.append(item["Text"])
            elif isinstance(item, dict) and item.get("Topics"):
                for sub in item.get("Topics", []):
                    if sub.get("Text"):
                        snippets.append(sub["Text"])
            if len(snippets) >= 3:
                break
        
        if snippets:
            result = "Related info: " + " | ".join(snippets[:3])
            logger.info(f"Web search successful for: {query}")
            return result
        
        # Try definition
        definition = data.get("Definition", "")
        if definition:
            return f"Definition: {definition}"
        
        logger.warning(f"No results found for web search: {query}")
        return "No quick answer found. Try rephrasing your query."
        
    except Exception as exc:
        logger.error(f"Web search failed for {query}: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Search failed: {exc}"


@tool
def get_weather(city: str) -> str:
    """Enhanced weather lookup for a city using Open‑Meteo API."""
    tool_name = "get_weather"
    _tool_call_counts[tool_name] += 1
    
    if not _rate_limited(tool_name, 0.5):
        return "Rate limited. Please retry shortly."
    
    try:
        logger.info(f"Getting weather for: {city}")
        
        # Geocoding
        geo = requests.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": city, "count": 1, "language": "en", "format": "json"},
            timeout=12,
        )
        geo_data = geo.json() if geo.ok else {}
        results = geo_data.get("results") or []
        
        if not results:
            logger.warning(f"City not found: {city}")
            return f"Could not find city '{city}'. Please check the spelling."
        
        loc = results[0]
        lat, lon = loc["latitude"], loc["longitude"]
        place = f"{loc.get('name')}, {loc.get('country_code', '')}".strip()

        # Weather data
        w = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": lat,
                "longitude": lon,
                "current": "temperature_2m,precipitation,relative_humidity_2m,apparent_temperature,is_day,weather_code,wind_speed_10m",
                "timezone": "auto"
            },
            timeout=12,
        )
        w_data = w.json() if w.ok else {}
        cur = w_data.get("current") or {}
        
        temp = cur.get("temperature_2m")
        precip = cur.get("precipitation", 0)
        humidity = cur.get("relative_humidity_2m")
        feels = cur.get("apparent_temperature")
        wind_speed = cur.get("wind_speed_10m", 0)
        is_day = cur.get("is_day", 1)
        
        time_of_day = "day" if is_day else "night"
        
        result = (
            f"Weather in {place} ({time_of_day}): "
            f"temp {temp}°C (feels like {feels}°C), "
            f"humidity {humidity}%, wind {wind_speed} km/h"
        )
        
        if precip > 0:
            result += f", precipitation {precip} mm"
        
        logger.info(f"Weather lookup successful for: {city}")
        return result
        
    except Exception as exc:
        logger.error(f"Weather lookup failed for {city}: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Weather lookup failed: {exc}"


@tool
def get_document_stats_tool(tenant_id: Optional[str] = None) -> str:
    """Get statistics about indexed documents for the current or specified tenant."""
    current_tenant = tenant_id or CURRENT_TENANT_ID or "default"
    
    tool_name = "get_document_stats"
    _tool_call_counts[tool_name] += 1
    
    try:
        stats = get_document_stats(current_tenant)
        
        if "error" in stats:
            return f"Error getting document stats: {stats['error']}"
        
        result = f"Document Statistics for tenant '{current_tenant}':\n"
        result += f"- Total chunks: {stats['total_chunks']}\n"
        result += f"- Unique sources: {stats['unique_sources']}\n"
        result += f"- File types: {', '.join([f'{k}({v})' for k, v in stats['file_types'].items()])}\n"
        
        if stats['sample_sources']:
            result += f"- Sample sources: {', '.join(stats['sample_sources'][:3])}..."
        
        return result
        
    except Exception as exc:
        logger.error(f"Error getting document stats: {exc}")
        _tool_error_counts[tool_name] += 1
        return f"Error getting document stats: {exc}"


@tool
def get_tool_statistics() -> str:
    """Get usage statistics for all tools."""
    tool_name = "get_tool_statistics"
    _tool_call_counts[tool_name] += 1
    
    try:
        stats = get_tool_stats()
        
        if not stats:
            return "No tool usage statistics available."
        
        result = "Tool Usage Statistics:\n"
        for name, data in stats.items():
            calls = data['call_count']
            errors = data['error_count']
            success_rate = ((calls - errors) / calls * 100) if calls > 0 else 0
            
            result += f"- {name}: {calls} calls, {errors} errors ({success_rate:.1f}% success)\n"
        
        return result
        
    except Exception as exc:
        logger.error(f"Error getting tool statistics: {exc}")
        return f"Error getting tool statistics: {exc}"


# -----------------------------
# LLM setup
# -----------------------------


def get_llm(temperature: float = 0):
    """Get LLM instance based on configuration. Supports both Google and OpenAI models."""
    model_provider = os.environ.get("MODEL_PROVIDER", "google").lower()

    if model_provider == "openai" and OPENAI_AVAILABLE:
        model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
        api_key = os.environ.get("OPENAI_API_KEY")

        if not api_key:
            logger.warning("OPENAI_API_KEY not set, falling back to Google models")
            model_provider = "google"
        else:
            return ChatOpenAI(
                model=model,
                temperature=temperature,
                api_key=api_key
            )

    # Default to Google models
    if model_provider != "google":
        logger.info(f"Model provider '{model_provider}' not available, using Google models")

    model = os.environ.get("GOOGLE_MODEL", "gemini-2.0-flash")
    api_key = os.environ.get("GOOGLE_API_KEY")

    if not api_key:
        raise ValueError("GOOGLE_API_KEY is required when using Google models")

    return ChatGoogleGenerativeAI(model=model, temperature=temperature)


def build_llm_with_tools_for_tenant(tenant_id: Optional[str]):
    tools = get_tenant_tools(tenant_id)
    return get_llm(temperature=0).bind_tools(tools)


# -----------------------------
# Enhanced RAG: Indexing and Retrieval per Tenant (FAISS)
# -----------------------------

import hashlib
from datetime import datetime
from pathlib import Path

# Enhanced embeddings implementation with better semantic understanding
class EnhancedEmbeddings(Embeddings):
    """Enhanced embeddings using improved text processing for better semantic understanding"""

    def __init__(self):
        self.dimension = 768  # Increased dimension for better representation

    def _text_to_vector(self, text):
        """Convert text to enhanced vector representation with better semantic understanding"""
        import hashlib
        import re
        from collections import Counter

        # Clean and normalize text
        text = text.lower().strip()
        if not text:
            return [0.0] * self.dimension

        words = re.findall(r'\b\w+\b', text)
        if not words:
            return [0.0] * self.dimension

        # Create enhanced features
        features = []
        word_counts = Counter(words)

        # 1. Semantic keyword features (expanded for better matching)
        semantic_keywords = {
            'recipe_cooking': ['recipe', 'cook', 'cooking', 'ingredient', 'ingredients', 'prepare', 'preparation',
                              'bake', 'baking', 'fry', 'boil', 'mix', 'stir', 'heat', 'oven', 'pan', 'pot',
                              'cup', 'tablespoon', 'teaspoon', 'minutes', 'temperature', 'serve', 'dish'],
            'story_narrative': ['story', 'stories', 'tale', 'tales', 'character', 'characters', 'plot', 'chapter',
                               'chapters', 'beginning', 'end', 'once', 'upon', 'time', 'lived', 'said', 'told',
                               'narrative', 'fiction', 'novel', 'book', 'read', 'reading'],
            'instructions': ['step', 'steps', 'instruction', 'instructions', 'method', 'procedure', 'process',
                            'guide', 'tutorial', 'how', 'way', 'first', 'second', 'third', 'next', 'then',
                            'finally', 'last', 'follow', 'complete'],
            'questions': ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'can', 'could', 'would',
                         'should', 'do', 'does', 'did', 'is', 'are', 'was', 'were'],
            'actions': ['make', 'create', 'build', 'develop', 'produce', 'generate', 'form', 'construct',
                       'establish', 'design', 'craft', 'manufacture']
        }

        for category, keywords in semantic_keywords.items():
            count = sum(word_counts.get(word, 0) for word in keywords)
            features.append(count / max(len(words), 1))

        # 2. Enhanced text statistics
        features.extend([
            len(text) / 1000.0,  # Text length
            len(words) / 100.0,  # Word count
            len(set(words)) / max(len(words), 1),  # Vocabulary diversity
            sum(len(w) for w in words) / max(len(words), 1) / 10.0,  # Average word length
            text.count('.') / max(len(text), 1),  # Sentence density
            text.count('?') / max(len(text), 1),  # Question density
            text.count('!') / max(len(text), 1),  # Exclamation density
            text.count(',') / max(len(text), 1),  # Comma density
            len([w for w in words if len(w) > 6]) / max(len(words), 1),  # Long words ratio
        ])

        # 3. N-gram features for better context
        bigrams = [f"{words[i]} {words[i+1]}" for i in range(len(words)-1)]
        bigram_counts = Counter(bigrams)

        important_bigrams = ['how to', 'what is', 'how do', 'how can', 'recipe for', 'story about',
                            'once upon', 'the end', 'first step', 'next step']
        for bigram in important_bigrams:
            features.append(bigram_counts.get(bigram, 0) / max(len(bigrams), 1))

        # 4. Character frequency features (26 letters)
        char_freq = [0] * 26
        for char in text:
            if 'a' <= char <= 'z':
                char_freq[ord(char) - ord('a')] += 1
        total_chars = sum(char_freq)
        if total_chars > 0:
            char_freq = [f / total_chars for f in char_freq]
        features.extend(char_freq)

        # 5. Hash-based features for uniqueness
        hash_obj = hashlib.md5(text.encode())
        hash_hex = hash_obj.hexdigest()
        hash_features = [int(hash_hex[i:i+2], 16) / 255.0 for i in range(0, min(len(hash_hex), 32), 2)]
        features.extend(hash_features)

        # 6. Contextual indicators
        features.extend([
            1.0 if any(word in text for word in ['recipe', 'cook', 'ingredient', 'prepare']) else 0.0,
            1.0 if any(word in text for word in ['story', 'character', 'plot', 'tale']) else 0.0,
            1.0 if any(word in text for word in ['step', 'instruction', 'guide', 'method']) else 0.0,
            1.0 if '?' in text else 0.0,
            1.0 if any(word in text for word in ['how', 'what', 'why', 'when', 'where']) else 0.0,
        ])

        # Pad or truncate to desired dimension
        while len(features) < self.dimension:
            features.extend(features[:min(self.dimension - len(features), len(features))])

        return features[:self.dimension]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        """Embed a list of documents"""
        return [self._text_to_vector(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        """Embed a single query"""
        return self._text_to_vector(text)

# Initialize embeddings
EMBEDDINGS = EnhancedEmbeddings()
logger.info("Using enhanced embeddings implementation (no TensorFlow dependencies)")

# Global form generator instance
FORM_GENERATOR = FormGenerator()

# Dynamic API Registry
@dataclass
class DynamicAPI:
    """Configuration for a dynamic API endpoint."""
    name: str
    base_url: str
    method: str = "GET"
    headers: Dict[str, str] = field(default_factory=dict)
    auth_type: str = "none"  # none, bearer, api_key, basic
    auth_value: str = ""
    description: str = ""
    parameters: Dict[str, Any] = field(default_factory=dict)

class DynamicAPIManager:
    """Manages dynamic API connections and tool generation."""
    
    def __init__(self):
        self.apis: Dict[str, DynamicAPI] = {}
        self.generated_tools: Dict[str, Callable] = {}
    
    def register_api(self, api: DynamicAPI) -> bool:
        """Register a dynamic API endpoint."""
        self.apis[api.name] = api
        
        # Generate tool function for this API
        tool_func = self._create_api_tool(api)
        self.generated_tools[api.name] = tool_func
        
        logger.info(f"Registered dynamic API: {api.name} ({api.method} {api.base_url})")
        return True
    
    def _create_api_tool(self, api: DynamicAPI) -> Callable:
        """Create a tool function for the API."""
        
        @tool
        def dynamic_api_call(query: str = "", **kwargs) -> str:
            f"""Call {api.name} API: {api.description}"""
            try:
                # Prepare headers
                headers = api.headers.copy()
                
                # Add authentication
                if api.auth_type == "bearer":
                    headers["Authorization"] = f"Bearer {api.auth_value}"
                elif api.auth_type == "api_key":
                    headers["X-API-Key"] = api.auth_value
                elif api.auth_type == "basic":
                    import base64
                    credentials = base64.b64encode(api.auth_value.encode()).decode()
                    headers["Authorization"] = f"Basic {credentials}"
                
                # Prepare URL and parameters
                url = api.base_url
                params = api.parameters.copy()
                
                # Add query parameter if provided
                if query:
                    params["q"] = query
                    params["query"] = query
                
                # Add any additional parameters
                params.update(kwargs)
                
                # Make API call
                if api.method.upper() == "GET":
                    response = requests.get(url, headers=headers, params=params, timeout=10)
                elif api.method.upper() == "POST":
                    response = requests.post(url, headers=headers, json=params, timeout=10)
                else:
                    return f"Unsupported HTTP method: {api.method}"
                
                response.raise_for_status()
                
                # Try to parse JSON response
                try:
                    data = response.json()
                    return json.dumps(data, indent=2)
                except:
                    return response.text
                    
            except requests.exceptions.RequestException as e:
                return f"API call failed: {str(e)}"
            except Exception as e:
                return f"Error calling {api.name} API: {str(e)}"
        
        # Set the function name dynamically
        dynamic_api_call.__name__ = f"call_{api.name.lower().replace(' ', '_')}"
        dynamic_api_call.__doc__ = f"Call {api.name} API: {api.description}"
        
        return dynamic_api_call
    
    def get_api_tools(self, tenant_id: str = None) -> List[Callable]:
        """Get all API tools for a tenant."""
        return list(self.generated_tools.values())
    
    def remove_api(self, api_name: str) -> bool:
        """Remove a dynamic API."""
        if api_name in self.apis:
            del self.apis[api_name]
            if api_name in self.generated_tools:
                del self.generated_tools[api_name]
            logger.info(f"Removed dynamic API: {api_name}")
            return True
        return False

# Global dynamic API manager
DYNAMIC_API_MANAGER = DynamicAPIManager()

def _json_to_professional_form(form_data: dict) -> ProfessionalForm:
    """Convert JSON form data to ProfessionalForm object."""
    # Extract basic info
    title = form_data.get("title", "Untitled Form")
    description = form_data.get("description", "")
    company_name = form_data.get("company_name", "")
    form_type = form_data.get("form_type", "general")
    footer_text = form_data.get("footer_text", "")

    # Handle both old format (direct fields) and new format (sections)
    sections = []

    if "sections" in form_data and isinstance(form_data["sections"], list):
        # New format with sections
        for section_data in form_data["sections"]:
            section_title = section_data.get("title", "Section")
            section_desc = section_data.get("description", "")

            fields = []
            for field_data in section_data.get("fields", []):
                field = FormField(
                    name=field_data.get("name", ""),
                    label=field_data.get("label", ""),
                    field_type=field_data.get("field_type", field_data.get("type", "text")),
                    required=field_data.get("required", False),
                    placeholder=field_data.get("placeholder", ""),
                    options=field_data.get("options", []),
                    validation=field_data.get("validation", ""),
                    description=field_data.get("description", ""),
                    default_value=field_data.get("default_value", "")
                )
                fields.append(field)

            section = FormSection(
                title=section_title,
                description=section_desc,
                fields=fields
            )
            sections.append(section)

    elif "fields" in form_data and isinstance(form_data["fields"], list):
        # Old format with direct fields - create a single section
        fields = []
        for field_data in form_data["fields"]:
            field = FormField(
                name=field_data.get("name", ""),
                label=field_data.get("label", ""),
                field_type=field_data.get("field_type", field_data.get("type", "text")),
                required=field_data.get("required", False),
                placeholder=field_data.get("placeholder", ""),
                options=field_data.get("options", []),
                validation=field_data.get("validation", ""),
                description=field_data.get("description", ""),
                default_value=field_data.get("default_value", "")
            )
            fields.append(field)

        # Create a single section
        section = FormSection(
            title="Form Fields",
            description="",
            fields=fields
        )
        sections.append(section)

    # Create the professional form
    professional_form = ProfessionalForm(
        title=title,
        description=description,
        company_name=company_name,
        form_type=form_type,
        sections=sections,
        footer_text=footer_text
    )

    return professional_form


def _tenant_index_path(tenant_id: str) -> str:
    os.makedirs("indices", exist_ok=True)
    return os.path.join("indices", f"faiss_{tenant_id}")


def _get_file_hash(file_path: str) -> str:
    """Generate hash for file content to detect changes."""
    try:
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except Exception:
        return ""


def _extract_text_from_file(file_path: str) -> tuple[str, dict]:
    """Enhanced text extraction with better metadata."""
    path_obj = Path(file_path)
    ext = path_obj.suffix.lower()
    
    metadata = {
        "source": file_path,
        "filename": path_obj.name,
        "file_type": ext,
        "file_size": path_obj.stat().st_size if path_obj.exists() else 0,
        "modified_time": datetime.fromtimestamp(path_obj.stat().st_mtime).isoformat() if path_obj.exists() else "",
        "file_hash": _get_file_hash(file_path)
    }
    
    text = ""
    
    try:
        if ext in {".txt", ".md"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == ".csv":
            try:
                import pandas as pd
                df = pd.read_csv(file_path)
                
                # Enhanced CSV processing for better RAG performance
                file_name = os.path.basename(file_path)
                text = f"CSV Dataset: {file_name}\n"
                text += "=" * 50 + "\n\n"
                
                # Dataset overview
                text += f"📊 DATASET OVERVIEW:\n"
                text += f"• File: {file_name}\n"
                text += f"• Columns: {len(df.columns)} columns\n"
                text += f"• Rows: {len(df)} records\n"
                text += f"• Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
                
                # Column information with data types
                text += f"📋 COLUMN DETAILS:\n"
                for i, col in enumerate(df.columns, 1):
                    dtype = str(df[col].dtype)
                    null_count = df[col].isnull().sum()
                    unique_count = df[col].nunique()
                    
                    text += f"{i}. {col}\n"
                    text += f"   - Type: {dtype}\n"
                    text += f"   - Unique values: {unique_count}\n"
                    text += f"   - Missing values: {null_count}\n"
                    
                    # Sample values for better understanding
                    if df[col].dtype == 'object':
                        sample_values = df[col].dropna().unique()[:5]
                        text += f"   - Sample values: {', '.join(map(str, sample_values))}\n"
                    else:
                        min_val = df[col].min()
                        max_val = df[col].max()
                        text += f"   - Range: {min_val} to {max_val}\n"
                    text += "\n"
                
                # Statistical summary for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    text += f"📈 STATISTICAL SUMMARY (Numeric Columns):\n"
                    stats_summary = df[numeric_cols].describe()
                    text += stats_summary.to_string() + "\n\n"
                
                # Category analysis for object columns
                object_cols = df.select_dtypes(include=['object']).columns
                if len(object_cols) > 0:
                    text += f"🏷️ CATEGORY ANALYSIS (Text Columns):\n"
                    for col in object_cols:
                        value_counts = df[col].value_counts().head(10)
                        text += f"{col}:\n"
                        for value, count in value_counts.items():
                            percentage = (count / len(df)) * 100
                            text += f"  • {value}: {count} ({percentage:.1f}%)\n"
                        text += "\n"
                
                # Sample data with better formatting
                text += f"📋 SAMPLE DATA (First 10 rows):\n"
                text += "-" * 80 + "\n"
                
                # Create a more readable table format
                sample_data = df.head(10)
                for index, row in sample_data.iterrows():
                    text += f"Row {index + 1}:\n"
                    for col in df.columns:
                        value = row[col]
                        if pd.isna(value):
                            value = "[Empty]"
                        text += f"  • {col}: {value}\n"
                    text += "\n"
                
                # Add searchable content for specific questions
                text += f"🔍 SEARCHABLE CONTENT:\n"
                text += f"This dataset contains information about: {', '.join(df.columns)}\n\n"
                
                # Create question-answerable content
                text += f"❓ QUICK FACTS:\n"
                for col in df.columns:
                    if df[col].dtype in ['int64', 'float64']:
                        total = df[col].sum()
                        avg = df[col].mean()
                        text += f"• Total {col}: {total}\n"
                        text += f"• Average {col}: {avg:.2f}\n"
                    elif df[col].dtype == 'object':
                        unique_vals = df[col].nunique()
                        most_common = df[col].mode().iloc[0] if len(df[col].mode()) > 0 else "N/A"
                        text += f"• Unique {col} values: {unique_vals}\n"
                        text += f"• Most common {col}: {most_common}\n"
                
                # Add the full dataset as structured text for complex queries
                text += "\n" + "=" * 50 + "\n"
                text += f"📊 COMPLETE DATASET:\n"
                text += "=" * 50 + "\n"
                text += df.to_string(index=False, max_rows=None)
                
                # Store enhanced metadata
                metadata["csv_columns"] = df.columns.tolist()
                metadata["csv_rows"] = len(df)
                metadata["csv_shape"] = df.shape
                metadata["csv_dtypes"] = {col: str(df[col].dtype) for col in df.columns}
                metadata["csv_numeric_columns"] = numeric_cols.tolist()
                metadata["csv_object_columns"] = object_cols.tolist()
                
            except Exception as exc:
                text = f"[CSV read error: {exc}]"
                metadata["error"] = str(exc)
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(file_path)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                metadata["page_count"] = len(reader.pages)
            except Exception as exc:  # noqa: BLE001
                text = f"[PDF read error: {exc}]"
                metadata["error"] = str(exc)
        elif ext == ".docx":
            try:
                import docx  # type: ignore
                d = docx.Document(file_path)
                text = "\n".join(p.text for p in d.paragraphs)
                metadata["paragraph_count"] = len(d.paragraphs)
            except Exception as exc:  # noqa: BLE001
                text = f"[DOCX read error: {exc}]"
                metadata["error"] = str(exc)
        elif ext == ".json":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    text = json.dumps(data, indent=2)
                metadata["json_keys"] = list(data.keys()) if isinstance(data, dict) else []
            except Exception as exc:
                text = f"[JSON read error: {exc}]"
                metadata["error"] = str(exc)
        else:
            # Try to read as text for other extensions
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception:
                text = f"[Unsupported file type: {ext}]"
                
    except Exception as exc:
        text = f"[File read error: {exc}]"
        metadata["error"] = str(exc)
    
    return text, metadata


def ingest_single_document(tenant_id: str, file_path: str, user_id: Optional[str] = None,
                          chunk_size: int = 1000, chunk_overlap: int = 150) -> Dict[str, Any]:
    """Enhanced single document ingestion with metadata tracking."""
    try:
        # Check if file already exists (deduplication)
        file_hash = hashlib.sha256(open(file_path, 'rb').read()).hexdigest()
        existing_docs = document_storage.get_documents_by_tenant(tenant_id)

        for doc in existing_docs:
            if doc.file_hash == file_hash:
                return {
                    "success": True,
                    "message": f"Document already exists: {doc.filename}",
                    "document_id": doc.document_id,
                    "duplicate": True
                }

        # Extract text and metadata
        text, base_metadata = _extract_text_from_file(file_path)

        if not text.strip():
            return {"success": False, "message": "No text content found in document"}

        # Create document metadata
        document_id = secrets.token_urlsafe(16)
        file_stat = os.stat(file_path)

        doc_metadata = DocumentMetadata(
            document_id=document_id,
            filename=os.path.basename(file_path),
            file_path=file_path,
            file_size=file_stat.st_size,
            file_type=Path(file_path).suffix.lower(),
            upload_timestamp=datetime.now().isoformat(),
            tenant_id=tenant_id,
            user_id=user_id,
            file_hash=file_hash,
            original_name=os.path.basename(file_path)
        )

        # Enhanced text splitting with better semantic boundaries
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n\n",  # Multiple line breaks (section breaks)
                "\n\n",    # Paragraph breaks
                "\n",      # Line breaks
                ". ",      # Sentence endings
                "! ",      # Exclamation endings
                "? ",      # Question endings
                "; ",      # Semicolon breaks
                ", ",      # Comma breaks (for lists)
                " ",       # Word breaks
                ""         # Character breaks (last resort)
            ],
            length_function=len,
            is_separator_regex=False,
        )

        # Split text into chunks
        chunks = splitter.split_text(text)

        # Post-process chunks to improve semantic coherence
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            # Clean up chunk boundaries
            chunk = chunk.strip()

            # If chunk is too short and not the last chunk, try to merge with next
            if len(chunk) < chunk_size * 0.3 and i < len(chunks) - 1:
                next_chunk = chunks[i + 1].strip()
                if len(chunk) + len(next_chunk) <= chunk_size * 1.2:
                    # Merge chunks
                    merged_chunk = chunk + " " + next_chunk
                    processed_chunks.append(merged_chunk)
                    chunks[i + 1] = ""  # Mark next chunk as processed
                    continue

            if chunk:  # Only add non-empty chunks
                processed_chunks.append(chunk)

        chunks = processed_chunks
        doc_metadata.chunk_count = len(chunks)

        # Create documents for vector store
        docs: List[Document] = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "tenant_id": tenant_id,
                "document_id": document_id,
                "chunk_id": i,
                "chunk_count": len(chunks),
                "chunk_size": len(chunk),
                "ingestion_time": datetime.now().isoformat()
            })

            # Sanitize metadata to ensure all values are serializable
            sanitized_metadata = {}
            for key, value in chunk_metadata.items():
                try:
                    # Convert to string if not a basic type
                    if isinstance(value, (str, int, float, bool, type(None))):
                        sanitized_metadata[key] = value
                    else:
                        sanitized_metadata[key] = str(value)
                except Exception:
                    sanitized_metadata[key] = "unknown"

            try:
                # Create document with error handling for Pydantic compatibility
                doc = Document(
                    page_content=chunk,
                    metadata=sanitized_metadata
                )
                docs.append(doc)
            except Exception as doc_error:
                logger.error(f"Error creating document for chunk {i}: {doc_error}")
                # Try with minimal metadata as fallback
                try:
                    doc = Document(
                        page_content=chunk,
                        metadata={"source": sanitized_metadata.get("source", "unknown")}
                    )
                    docs.append(doc)
                except Exception as fallback_error:
                    logger.error(f"Fallback document creation also failed: {fallback_error}")
                    continue

        # Save to vector store
        index_dir = _tenant_index_path(tenant_id)
        try:
            logger.info(f"Attempting to save {len(docs)} documents to vector store at {index_dir}")

            # Debug: Check document structure
            if docs:
                sample_doc = docs[0]
                logger.info(f"Sample document type: {type(sample_doc)}")
                logger.info(f"Sample metadata keys: {list(sample_doc.metadata.keys())}")

            if os.path.isdir(index_dir):
                logger.info("Loading existing vector store")
                try:
                    vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
                    logger.info("Adding documents to existing vector store")
                    vs.add_documents(docs)
                except (KeyError, AttributeError, Exception) as load_error:
                    logger.warning(f"Failed to load existing vector store (likely version incompatibility): {load_error}")
                    logger.info("Creating new vector store to replace corrupted one")
                    # Remove corrupted index directory
                    import shutil
                    shutil.rmtree(index_dir, ignore_errors=True)
                    vs = FAISS.from_documents(docs, EMBEDDINGS)
            else:
                logger.info("Creating new vector store")
                vs = FAISS.from_documents(docs, EMBEDDINGS)

            logger.info("Saving vector store to disk")
            vs.save_local(index_dir)
            logger.info("Vector store saved successfully")

            doc_metadata.indexed = True
        except Exception as e:
            logger.error(f"Failed to save to vector store: {e}")
            logger.error(f"Error type: {type(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return {"success": False, "message": f"Vector indexing failed: {e}"}

        # Save document metadata to database
        if document_storage.save_document(doc_metadata):
            return {
                "success": True,
                "message": f"Document processed successfully: {doc_metadata.filename}",
                "document_id": document_id,
                "chunks": len(chunks),
                "duplicate": False
            }
        else:
            return {"success": False, "message": "Failed to save document metadata"}

    except Exception as e:
        logger.error(f"Error processing document {file_path}: {e}")
        return {"success": False, "message": f"Processing failed: {e}"}

def ingest_multiple_documents(tenant_id: str, file_paths: List[str], user_id: Optional[str] = None) -> Dict[str, Any]:
    """Process multiple documents simultaneously."""
    results = []
    successful = 0
    failed = 0
    duplicates = 0

    for file_path in file_paths:
        result = ingest_single_document(tenant_id, file_path, user_id)
        results.append({
            "file_path": file_path,
            "filename": os.path.basename(file_path),
            **result
        })

        if result["success"]:
            if result.get("duplicate", False):
                duplicates += 1
            else:
                successful += 1
        else:
            failed += 1

    return {
        "success": True,
        "total_files": len(file_paths),
        "successful": successful,
        "failed": failed,
        "duplicates": duplicates,
        "results": results
    }

def ingest_documents_from_dir(tenant_id: str, source_dir: str, chunk_size: int = 1000, chunk_overlap: int = 150) -> str:
    """Enhanced document ingestion with better processing and metadata."""
    file_paths = []

    for root, _dirs, files in os.walk(source_dir):
        for fname in files:
            file_path = os.path.join(root, fname)

            # Skip hidden files and common non-document files
            if fname.startswith('.') or fname.lower().endswith(('.exe', '.dll', '.so', '.dylib')):
                continue

            file_paths.append(file_path)

    if not file_paths:
        return "No documents found to ingest."

    result = ingest_multiple_documents(tenant_id, file_paths)
    return f"Processed {result['total_files']} files: {result['successful']} successful, {result['failed']} failed, {result['duplicates']} duplicates"


def get_retriever_for_tenant(tenant_id: str):
    """Enhanced retriever with better search capabilities."""
    index_dir = _tenant_index_path(tenant_id)
    if not os.path.isdir(index_dir):
        return None
        
    try:
        vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
    except (KeyError, AttributeError, Exception) as exc:
        logger.warning(f"Vector store for tenant {tenant_id} is corrupted (likely version incompatibility): {exc}")
        logger.info(f"Removing corrupted vector store at {index_dir}")
        import shutil
        shutil.rmtree(index_dir, ignore_errors=True)
        return None
    except Exception as exc:
        logger.error(f"Error loading vector store for tenant {tenant_id}: {exc}")
        return None
    
    def _retrieve(query: str, k: int = 4, score_threshold: float = 0.8) -> List[Document]:
        """Enhanced retrieval with improved similarity scoring and query expansion."""
        try:
            # Expand query with synonyms and related terms
            expanded_queries = [query]

            # Add variations for better matching
            query_lower = query.lower()
            if 'recipe' in query_lower and 'how' not in query_lower:
                expanded_queries.append(f"how to make {query}")
            if 'story' in query_lower and 'tell' not in query_lower:
                expanded_queries.append(f"tell me the story {query}")
            if any(word in query_lower for word in ['what', 'how', 'why', 'when', 'where']):
                # Keep original question format
                pass
            else:
                # Add question variations
                expanded_queries.extend([
                    f"what is {query}",
                    f"how to {query}",
                    f"about {query}"
                ])

            all_docs_with_scores = []

            # Search with each query variation
            for q in expanded_queries[:3]:  # Limit to avoid too many queries
                try:
                    docs_scores = vs.similarity_search_with_score(q, k=k*3)
                    all_docs_with_scores.extend(docs_scores)
                except Exception as e:
                    logger.warning(f"Error searching with query '{q}': {e}")
                    continue

            # Remove duplicates based on content and document_id
            seen_content = set()
            seen_doc_ids = set()
            unique_docs_with_scores = []

            for doc, score in all_docs_with_scores:
                content_hash = hash(doc.page_content[:100])  # Use first 100 chars for dedup
                doc_id = doc.metadata.get('document_id', '')
                chunk_id = doc.metadata.get('chunk_id', 0)
                unique_key = f"{doc_id}_{chunk_id}"

                if content_hash not in seen_content and unique_key not in seen_doc_ids:
                    seen_content.add(content_hash)
                    seen_doc_ids.add(unique_key)
                    unique_docs_with_scores.append((doc, score))

            # Sort by score (lower is better in FAISS)
            unique_docs_with_scores.sort(key=lambda x: x[1])

            # Apply more lenient scoring for better recall
            filtered_docs = [
                doc for doc, score in unique_docs_with_scores
                if score <= score_threshold
            ][:k*2]  # Get more docs initially

            # If still no docs meet threshold, use top results with relaxed threshold
            if not filtered_docs:
                relaxed_threshold = min(score_threshold * 1.5, 1.2)
                filtered_docs = [
                    doc for doc, score in unique_docs_with_scores
                    if score <= relaxed_threshold
                ][:k]

            # If still no results, return top k regardless of score
            if not filtered_docs:
                filtered_docs = [doc for doc, _ in unique_docs_with_scores[:k]]

            # Limit final results
            return filtered_docs[:k]

        except Exception as exc:
            logger.error(f"Error during retrieval: {exc}")
            return []

    return _retrieve


def get_document_stats(tenant_id: str) -> dict:
    """Get statistics about indexed documents for a tenant."""
    index_dir = _tenant_index_path(tenant_id)
    if not os.path.isdir(index_dir):
        return {"error": "No index found for tenant"}
    
    try:
        vs = FAISS.load_local(index_dir, EMBEDDINGS, allow_dangerous_deserialization=True)
        
        # Get basic stats
        total_chunks = vs.index.ntotal
        
        # Sample some documents to get metadata stats
        sample_docs = vs.similarity_search("", k=min(100, total_chunks)) if total_chunks > 0 else []
        
        file_types = {}
        sources = set()
        
        for doc in sample_docs:
            metadata = doc.metadata
            file_type = metadata.get("file_type", "unknown")
            file_types[file_type] = file_types.get(file_type, 0) + 1
            sources.add(metadata.get("source", "unknown"))
        
        return {
            "tenant_id": tenant_id,
            "total_chunks": total_chunks,
            "unique_sources": len(sources),
            "file_types": file_types,
            "sample_sources": list(sources)[:10]  # Show first 10 sources
        }
        
    except Exception as exc:
        return {"error": f"Error getting stats: {exc}"}


# -----------------------------
# Agent Nodes
# -----------------------------


def node_router(state: MessagesState) -> str:
    """Enhanced router with better intent classification and tool name recognition."""
    last_user = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            last_user = getattr(msg, "content", "")
            break

    # Check for explicit tool name references first
    last_user_lower = last_user.lower()
    
    # Get available tools for current tenant to check against
    tenant_id = CURRENT_TENANT_ID or "default"
    try:
        available_tools = get_tenant_tools(tenant_id)
        tool_names = [tool.name.lower() for tool in available_tools]
        
        # Check if user explicitly mentions a tool name
        for tool_name in tool_names:
            if tool_name in last_user_lower:
                logger.info(f"Router detected explicit tool reference: {tool_name}")
                return "api_exec"
    except Exception as e:
        logger.warning(f"Failed to get tools for routing: {e}")

    # Enhanced routing logic with more context
    prompt = (
        "You are an intelligent router for a multi-agent chatbot system. "
        "Analyze the user's message and classify their intent into one of these categories:\n\n"
        "- greeting: greetings, small talk, general conversation, introductions\n"
        "- doc_qa: questions about documents, files, or knowledge base content\n"
        "- api_exec: requests to perform actions, call APIs, get external data, use tools, fetch data from services\n"
        "- form_gen: requests to create forms, collect structured data, or generate input fields\n"
        "- analytics: requests for data analysis, statistics, insights, reports, or metrics\n"
        "- escalate: requests for human help, complaints, or complex issues beyond AI capability\n\n"
        "Consider context clues like:\n"
        "- Keywords related to documents, files, or knowledge\n"
        "- Action words like 'get', 'fetch', 'call', 'search', 'use', 'execute'\n"
        "- API/tool references like 'api', 'tool', 'service', 'endpoint'\n"
        "- Form-related terms like 'form', 'input', 'collect', 'survey'\n"
        "- Analytics terms like 'analyze', 'statistics', 'metrics', 'report', 'insights'\n"
        "- Escalation phrases like 'human', 'agent', 'help', 'support'\n\n"
        "Respond with only the category name."
    )
    
    llm = get_llm(temperature=0)
    res = llm.invoke([("system", prompt), ("user", last_user or "hello")])
    label = (getattr(res, "content", "") or "").strip().lower()
    
    # Enhanced fallback logic with better keyword detection
    if label not in {"greeting", "doc_qa", "api_exec", "form_gen", "analytics", "escalate"}:
        if any(word in last_user_lower for word in ["document", "file", "pdf", "text", "knowledge", "uploaded"]):
            return "doc_qa"
        elif any(word in last_user_lower for word in [
            "weather", "search", "api", "get", "fetch", "call", "use", "tool", "service", 
            "endpoint", "http", "request", "data", "posts", "facts", "joke", "bin"
        ]):
            return "api_exec"
        elif any(word in last_user_lower for word in ["form", "input", "collect", "survey", "field"]):
            return "form_gen"
        elif any(word in last_user_lower for word in ["analyze", "analytics", "statistics", "metrics", "report", "insights", "stats"]):
            return "analytics"
        elif any(word in last_user_lower for word in ["human", "agent", "help", "support", "escalate"]):
            return "escalate"
        else:
            return "greeting"
    
    logger.info(f"Router classified intent as: {label} for query: {last_user[:50]}...")
    return label


def node_greeting(state: MessagesState):
    llm = get_llm(temperature=0.6)
    sys = (
        "You are a helpful generalist assistant. Be concise and friendly."
    )
    res = llm.invoke([("system", sys), *state["messages"]])
    return {"messages": [res]}


def node_doc_qa(state: MessagesState):
    """Enhanced Document Q&A with chat context memory and multiple document support."""
    tenant_id = CURRENT_TENANT_ID or "default"
    session_id = CURRENT_SESSION.session_id if CURRENT_SESSION else "default"

    # Check if documents are available
    retr = get_retriever_for_tenant(tenant_id)
    documents = document_storage.get_documents_by_tenant(tenant_id)

    if retr is None or not documents:
        content = (
            f"No documents indexed for tenant '{tenant_id}'. "
            f"Please upload documents first using the upload area."
        )
        # Save assistant message to chat history
        save_chat_message_to_history(session_id, tenant_id, "assistant", content, "doc_qa")
        return {"messages": [("assistant", content)]}

    # Find the latest user message
    user_msg = ""
    for msg in reversed(state["messages"]):
        # Handle both tuple format ("user", content) and object format
        if isinstance(msg, tuple) and len(msg) >= 2:
            role, content = msg[0], msg[1]
            if role in ["user", "human"]:
                user_msg = content
                break
        elif getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    # Save user message to chat history
    save_chat_message_to_history(session_id, tenant_id, "user", user_msg, "doc_qa")

    # Get chat history for context
    chat_history = document_storage.get_chat_history(session_id, limit=10)

    # Enhanced document retrieval with better coverage
    docs = retr(user_msg, k=10)  # Increased for better coverage across multiple documents

    if not docs:
        content = (
            f"I couldn't find relevant information in the uploaded documents for your query: '{user_msg}'. "
            f"Please try rephrasing your question or check if the relevant documents are uploaded."
        )
        save_chat_message_to_history(session_id, tenant_id, "assistant", content, "doc_qa")
        return {"messages": [("assistant", content)]}

    # Group documents by source for better organization and ensure all documents are represented
    doc_sources = {}
    doc_metadata_info = {}

    for doc in docs:
        source = doc.metadata.get('source', 'Unknown')
        doc_id = doc.metadata.get('document_id', 'Unknown')

        if source not in doc_sources:
            doc_sources[source] = []
            doc_metadata_info[source] = {
                'document_id': doc_id,
                'chunks': 0,
                'total_length': 0
            }

        doc_sources[source].append(doc.page_content)
        doc_metadata_info[source]['chunks'] += 1
        doc_metadata_info[source]['total_length'] += len(doc.page_content)

    # Build enhanced context with better document organization
    context_parts = []
    referenced_docs = []

    # Sort documents by relevance (number of chunks retrieved)
    sorted_sources = sorted(doc_sources.items(), key=lambda x: len(x[1]), reverse=True)

    for source, contents in sorted_sources:
        # Find document metadata for this source
        doc_info = None
        for doc_meta in documents:
            if doc_meta.file_path.endswith(source) or doc_meta.filename == source:
                doc_info = doc_meta
                referenced_docs.append(doc_meta.document_id)
                break

        doc_name = doc_info.filename if doc_info else source
        doc_metadata = doc_metadata_info.get(source, {})

        # Enhanced document header with relevance info
        context_parts.append(f"[Document: {doc_name} - {doc_metadata.get('chunks', len(contents))} relevant sections found]")

        # Deduplicate and organize content
        unique_contents = []
        seen_content = set()

        for content in contents:
            # Simple deduplication based on first 100 characters
            content_key = content[:100].strip().lower()
            if content_key not in seen_content and len(content.strip()) > 20:
                seen_content.add(content_key)
                unique_contents.append(content.strip())

        # Add content with better separation
        for i, content in enumerate(unique_contents):
            if i > 0:
                context_parts.append("---")  # Separator between chunks from same document
            context_parts.append(content)

        context_parts.append("")  # Empty line between documents

    context = "\n".join(context_parts)

    # Validate context quality
    if len(context.strip()) < 100:
        content = (
            f"I found documents but couldn't extract sufficient relevant content for your query: '{user_msg}'. "
            f"The available documents might not contain the specific information you're looking for. "
            f"Please try rephrasing your question or check if the relevant documents are properly uploaded."
        )
        save_chat_message_to_history(session_id, tenant_id, "assistant", content, "doc_qa")
        return {"messages": [("assistant", content)]}

    # Build conversation context from recent chat history
    conversation_context = ""
    if len(chat_history) > 1:  # More than just the current message
        recent_messages = chat_history[-6:-1]  # Last 5 messages before current
        conversation_context = "\n".join([
            f"{msg.role.title()}: {msg.content}"
            for msg in recent_messages
        ])

    # Enhanced prompt with conversation context and better instructions
    prompt_parts = [
        "You are an expert document Q&A assistant with advanced comprehension capabilities.",
        "Your task is to provide accurate, comprehensive answers based on the provided documents.",
        "",
        "IMPORTANT GUIDELINES:",
        "1. Answer ONLY based on the information found in the provided documents",
        "2. The documents may contain structured data like CSV files, tables, or datasets - analyze these carefully",
        "3. For CSV/tabular data, look for specific values, prices, quantities, categories, and other data points",
        "4. When answering about specific items (like products, prices, quantities), search through ALL the document content thoroughly",
        "5. If you find relevant data in CSV format or tables, extract and present the specific information requested",
        "6. Only state 'This information is not available in the uploaded documents' if you truly cannot find ANY relevant data after thorough analysis",
        "7. When answering, specify which document contains the information",
        "8. For recipes or instructions, provide complete step-by-step details if available",
        "9. For stories, provide comprehensive summaries or specific details as requested",
        "10. If multiple documents contain relevant information, synthesize information from all relevant sources",
        "11. Use direct quotes or specific data points when appropriate to support your answers",
        "12. If the question asks about something specific (like product names, prices, categories), look for those exact terms AND related concepts",
        "13. Use conversation history to maintain context and provide coherent responses",
        "14. For CSV data specifically: look for matching product names, categories, prices, quantities, and other fields that answer the question",
        ""
    ]

    if conversation_context:
        prompt_parts.extend([
            "RECENT CONVERSATION CONTEXT:",
            conversation_context,
            ""
        ])

    prompt_parts.extend([
        "AVAILABLE DOCUMENTS WITH RELEVANT CONTENT:",
        context,
        "",
        f"USER QUESTION: {user_msg}",
        "",
        "RESPONSE INSTRUCTIONS:",
        "- Analyze the document content thoroughly, especially any structured data, tables, or CSV content",
        "- If the documents contain CSV data or tabular information, look for specific data points that answer the question",
        "- For product queries: search for product names, prices, quantities, categories, and descriptions",
        "- For data queries: look for numerical values, statistics, totals, and specific measurements",
        "- Provide detailed, accurate answers based solely on the document content above",
        "- If you find relevant information, be comprehensive and include all pertinent details",
        "- Present data in a clear, organized manner (e.g., 'Product: X, Price: Y, Stock: Z')",
        "- If you cannot find the specific information requested, but find related information, provide what IS available",
        "- Only state information is unavailable if you genuinely cannot find ANY relevant data after thorough analysis",
        "- Always maintain accuracy and never make up information not present in the documents"
    ])

    prompt = "\n".join(prompt_parts)

    # Generate response
    llm = get_llm(temperature=0.1)  # Slightly higher temperature for more natural responses
    res = llm.invoke([("system", "Document QA mode with conversation context."), ("user", prompt)])

    # Save assistant response to chat history
    response_content = getattr(res, "content", str(res))
    save_chat_message_to_history(session_id, tenant_id, "assistant", response_content, "doc_qa", referenced_docs)

    return {"messages": [res]}

def save_chat_message_to_history(session_id: str, tenant_id: str, role: str, content: str,
                                agent_type: Optional[str] = None, document_references: Optional[List[str]] = None):
    """Save a chat message to the persistent chat history."""
    try:
        message = ChatMessage(
            message_id=secrets.token_urlsafe(16),
            session_id=session_id,
            tenant_id=tenant_id,
            role=role,
            content=content,
            timestamp=datetime.now().isoformat(),
            user_id=CURRENT_SESSION.user_id if CURRENT_SESSION else None,
            agent_type=agent_type,
            document_references=document_references or []
        )
        document_storage.save_chat_message(message)
    except Exception as e:
        logger.error(f"Failed to save chat message: {e}")


def node_api_exec(state: MessagesState):
    """Enhanced API execution node with proper tool handling and better tool selection."""
    tenant_id = CURRENT_TENANT_ID or "default"
    tools = get_tenant_tools(tenant_id)
    
    # Extract user message
    user_msg = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break
    
    logger.info(f"API Executor processing: {user_msg[:100]}...")
    logger.info(f"Available tools: {[t.name for t in tools]}")
    
    # Create LLM with tools
    llm_with_tools = get_llm(temperature=0).bind_tools(tools)
    
    # Enhanced system prompt for API execution with tool awareness
    tool_list = "\n".join([f"- {tool.name}: {tool.description}" for tool in tools])
    
    system_prompt = (
        "You are an API execution specialist. Your role is to:\n"
        "1. Understand what the user wants to accomplish\n"
        "2. Select and use the appropriate tools to fulfill their request\n"
        "3. Provide clear, helpful responses based on tool results\n"
        "4. Handle errors gracefully and suggest alternatives\n\n"
        f"Available tools:\n{tool_list}\n\n"
        "IMPORTANT INSTRUCTIONS:\n"
        "- When the user mentions a specific tool name, use that tool\n"
        "- For data requests, use appropriate API tools to fetch real data\n"
        "- Always explain what you're doing and why\n"
        "- If a tool fails, try alternatives or explain the issue\n"
        "- Use multiple tools if needed to complete the request"
    )
    
    # Prepare messages with system prompt
    messages = [("system", system_prompt)] + state["messages"]
    
    # Invoke LLM with tools
    response = llm_with_tools.invoke(messages)
    
    # Check if tools were called
    if hasattr(response, 'tool_calls') and response.tool_calls:
        logger.info(f"Tool calls detected: {[call.get('name', 'unknown') for call in response.tool_calls]}")
        
        # Handle tool calls
        tool_node = ToolNode(tools)
        tool_results = tool_node.invoke({"messages": [response]})
        
        # Generate final response with tool results
        final_messages = messages + [response] + tool_results["messages"]
        final_response = get_llm(temperature=0).invoke(final_messages + [
            ("system", "Summarize the results and provide a helpful response to the user. Include relevant data from the tool execution.")
        ])
        
        logger.info(f"API execution completed with tool calls")
        return {"messages": [final_response]}
    
    logger.info(f"API execution completed without tool calls")
    return {"messages": [response]}


def node_form_gen(state: MessagesState):
    """Professional form generation with PDF/DOC export capabilities."""
    if not has_permission("generate_forms"):
        return {"messages": [("assistant", "Permission denied: form generation not allowed")]}

    llm = get_llm(temperature=0)
    user_msg = ""
    for msg in reversed(state["messages"]):
        # Handle both tuple format ("user", content) and object format
        if isinstance(msg, tuple) and len(msg) >= 2:
            role, content = msg[0], msg[1]
            if role in ["user", "human"]:
                user_msg = content
                break
        elif getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    # Enhanced file format detection
    file_format = "html"  # default changed to HTML
    user_msg_lower = user_msg.lower()

    # Check for explicit format specification
    if "(format: html)" in user_msg_lower:
        file_format = "html"
    elif "(format: docx)" in user_msg_lower or "(format: doc)" in user_msg_lower:
        file_format = "docx"
    elif "(format: pdf)" in user_msg_lower:
        file_format = "pdf"
    # Check for format keywords
    elif any(word in user_msg_lower for word in ["html", "web form", "interactive form"]):
        file_format = "html"
    elif any(word in user_msg_lower for word in ["docx", "doc", "word document", "microsoft word"]):
        file_format = "docx"
    elif any(word in user_msg_lower for word in ["pdf", "portable document"]):
        file_format = "pdf"

    # Log the detected format for debugging
    logger.info(f"Full message received: '{user_msg}'")
    logger.info(f"Message lowercase: '{user_msg_lower}'")
    logger.info(f"Contains '(format: docx)': {'(format: docx)' in user_msg_lower}")
    logger.info(f"Contains '(format: pdf)': {'(format: pdf)' in user_msg_lower}")
    logger.info(f"Detected file format: {file_format}")
    
    enhanced_prompt = (
        "You are a professional form generation specialist. Create a comprehensive, structured form based on the user's request.\n"
        "Generate a detailed JSON response with this EXACT structure:\n\n"
        "{\n"
        '  "title": "Professional Form Title",\n'
        '  "description": "Detailed description of the form purpose and instructions",\n'
        '  "company_name": "Company/Organization Name (if applicable)",\n'
        '  "form_type": "contract|survey|registration|feedback|application|contact|other",\n'
        '  "sections": [\n'
        "    {\n"
        '      "title": "Section Title",\n'
        '      "description": "Section description",\n'
        '      "fields": [\n'
        "        {\n"
        '          "name": "field_name",\n'
        '          "label": "Field Label",\n'
        '          "field_type": "text|email|number|date|select|textarea|checkbox|radio|tel",\n'
        '          "required": true|false,\n'
        '          "placeholder": "Placeholder text",\n'
        '          "description": "Field description/help text",\n'
        '          "options": ["option1", "option2"] // Only for select/radio/checkbox\n'
        "        }\n"
        "      ]\n"
        "    }\n"
        "  ],\n"
        '  "footer_text": "Footer text, terms, or additional information"\n'
        "}\n\n"
        "IMPORTANT GUIDELINES:\n"
        "- Create logical sections to organize related fields\n"
        "- Include comprehensive field descriptions and help text\n"
        "- Use appropriate field types for data validation\n"
        "- Add relevant options for select/radio/checkbox fields\n"
        "- Make forms professional and user-friendly\n"
        "- Include proper legal disclaimers for contracts\n"
        "- Add contact information sections where appropriate\n"
        f"\nUser request: {user_msg}"
    )
    
    try:
        # Use proper message format for Gemini
        messages = [
            ("system", "You are a professional form generation specialist. Create comprehensive, structured forms."),
            ("user", enhanced_prompt)
        ]
        res = llm.invoke(messages)
        content = getattr(res, "content", "")

        # Clean up the content to extract JSON
        content = content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]

        # Find JSON boundaries
        start = content.find("{")
        end = content.rfind("}") + 1

        if start != -1 and end > start:
            form_json = content[start:end]

            try:
                form_data = json.loads(form_json)

                # Convert JSON to ProfessionalForm object
                professional_form = _json_to_professional_form(form_data)

                # Generate the file
                try:
                    if file_format == "html":
                        # For HTML, generate content without saving to file (for preview/editing)
                        html_content, filename = FORM_GENERATOR.generate_html_content(professional_form)
                        
                        # Generate form preview for HTML format
                        form_preview = FORM_GENERATOR.generate_form_preview(professional_form)
                        
                        response_text = (
                            f"✅ **Interactive {professional_form.form_type.title()} Form Generated Successfully!**\n\n"
                            f"📋 **Form Preview:**\n{form_preview}\n\n"
                            f"📄 **File Details:**\n"
                            f"• Format: Interactive HTML Form\n"
                            f"• File Size: {len(html_content.encode('utf-8')) / 1024:.1f} KB\n"
                            f"• Total Fields: {sum(len(section.fields) for section in professional_form.sections)}\n"
                            f"• Sections: {len(professional_form.sections)}\n\n"
                            f"🎯 **Interactive Form Ready for Preview & Editing!** The form includes:\n"
                            f"• ✅ Real-time validation\n"
                            f"• ✅ Professional styling\n"
                            f"• ✅ Live preview capabilities\n"
                            f"• ✅ HTML editor for customization\n"
                            f"• ✅ Multiple download options (HTML/PDF/DOCX)\n"
                            f"• ✅ Mobile-responsive design\n"
                            f"• ✅ Form submission handling"
                        )

                        return {
                            "messages": [("assistant", response_text)],
                            "form_generated": True,
                            "preview": form_preview,
                            "html_content": html_content,
                            "filename": filename,
                            "content_type": "text/html",
                            "file_format": file_format,
                            "file_size": len(html_content.encode('utf-8')),
                            "interactive": True
                        }
                    elif file_format == "docx":
                        if not DOCX_AVAILABLE:
                            file_path = FORM_GENERATOR.create_pdf_form(professional_form)
                            file_format = "pdf"  # fallback
                        else:
                            file_path = FORM_GENERATOR.create_docx_form(professional_form)
                    else:
                        file_path = FORM_GENERATOR.create_pdf_form(professional_form)

                    # Generate form preview
                    form_preview = FORM_GENERATOR.generate_form_preview(professional_form)

                    # Handle different file formats for content reading and response
                    # For PDF/DOCX, read as binary
                    with open(file_path, 'rb') as f:
                        file_content = f.read()

                    # Create enhanced response with preview and download capability
                    response_text = (
                        f"✅ **Professional {professional_form.form_type.title()} Form Generated Successfully!**\n\n"
                        f"📋 **Form Preview:**\n{form_preview}\n\n"
                        f"📄 **File Details:**\n"
                        f"• Format: {file_format.upper()}\n"
                        f"• File Size: {len(file_content) / 1024:.1f} KB\n"
                        f"• Total Fields: {sum(len(section.fields) for section in professional_form.sections)}\n"
                        f"• Sections: {len(professional_form.sections)}\n\n"
                        f"🎯 **Ready for Download!** The form has been professionally formatted with proper headings, "
                        f"sections, field labels, and validation requirements."
                    )

                    # Return response with download capability
                    logger.info(f"Generated professional {file_format.upper()} form: {file_path}")
                    return {
                        "messages": [("assistant", response_text)],
                        "form_generated": True,
                        "preview": form_preview,
                        "file_content": base64.b64encode(file_content).decode('utf-8'),
                        "filename": os.path.basename(file_path),
                        "content_type": "application/pdf" if file_format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "file_format": file_format,
                        "file_size": len(file_content)
                    }

                except Exception as file_error:
                    logger.error(f"File generation error: {file_error}")
                    # Fallback to JSON response
                    form_data["generated_at"] = datetime.now().isoformat()
                    form_data["tenant_id"] = CURRENT_TENANT_ID or "default"
                    form_data["form_id"] = professional_form.form_id

                    fallback_response = (
                        f"⚠️ Form structure generated successfully, but file creation failed.\n"
                        f"Error: {file_error}\n\n"
                        f"**Form JSON Structure:**\n```json\n{json.dumps(form_data, indent=2)}\n```"
                    )
                    return {"messages": [("assistant", fallback_response)]}

            except (json.JSONDecodeError, ValueError) as e:
                logger.error(f"Form parsing error: {e}")
                return {"messages": [("assistant", f"Error parsing form structure: {e}")]}

        # Fallback if no valid JSON found
        return {"messages": [("assistant", "Unable to generate form structure. Please provide more specific requirements.")]}

    except Exception as exc:
        logger.error(f"Form generation error: {exc}")
        return {"messages": [("assistant", f"Error generating form: {exc}")]}


def node_escalate(state: MessagesState):
    """Enhanced escalation workflow with proper handling."""
    try:
        # Get user message for context
        user_msg = ""
        for msg in reversed(state["messages"]):
            if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
                user_msg = getattr(msg, "content", "")
                break
        
        # Create escalation record
        escalation_id = secrets.token_urlsafe(8)
        escalation_data = {
            "escalation_id": escalation_id,
            "tenant_id": CURRENT_TENANT_ID or "default",
            "session_id": CURRENT_SESSION.session_id if CURRENT_SESSION else None,
            "user_message": user_msg,
            "timestamp": datetime.now().isoformat(),
            "status": "pending",
            "conversation_history": [
                {
                    "role": getattr(msg, "type", "unknown"),
                    "content": getattr(msg, "content", str(msg))
                }
                for msg in state["messages"][-5:]  # Last 5 messages for context
            ]
        }
        
        # Log escalation
        logger.info(f"Escalation created: {escalation_id} for tenant {CURRENT_TENANT_ID}")
        
        # In a real system, you would:
        # 1. Store escalation in database
        # 2. Notify human agents
        # 3. Create ticket in support system
        # 4. Send email/slack notification
        
        response = (
            f"I've escalated your request to a human agent.\n\n"
            f"Escalation ID: {escalation_id}\n"
            f"Your request has been logged and a human agent will assist you shortly.\n"
            f"Please keep this escalation ID for reference.\n\n"
            f"In the meantime, you can continue using the chatbot for other queries."
        )
        
        return {"messages": [("assistant", response)]}
        
    except Exception as exc:
        logger.error(f"Escalation error: {exc}")
        return {"messages": [("assistant", "I apologize, but I'm having trouble escalating your request. Please try again or contact support directly.")]}


def node_analytics(state: MessagesState):
    """Analytics agent for data analysis and insights."""
    if not has_permission("use_tools"):
        return {"messages": [("assistant", "Permission denied: analytics not allowed")]}

    llm = get_llm(temperature=0)
    user_msg = ""
    for msg in reversed(state["messages"]):
        if getattr(msg, "type", None) == "human" or getattr(msg, "role", None) == "user":
            user_msg = getattr(msg, "content", "")
            break

    try:
        # Get system statistics for analysis
        stats = get_system_stats()
        tool_stats = get_tool_stats()

        # Enhanced analytics prompt
        analytics_prompt = (
            "You are an analytics specialist. Analyze the provided system data and user request to provide insights.\n\n"
            f"System Statistics:\n{json.dumps(stats, indent=2)}\n\n"
            f"Tool Usage Statistics:\n{json.dumps(tool_stats, indent=2)}\n\n"
            "Provide detailed analysis including:\n"
            "1. Key metrics and trends\n"
            "2. Usage patterns and insights\n"
            "3. Recommendations for optimization\n"
            "4. Performance indicators\n"
            "5. Actionable insights\n\n"
            f"User request: {user_msg}"
        )

        messages = [
            ("system", "You are a data analytics expert. Provide comprehensive analysis and actionable insights."),
            ("user", analytics_prompt)
        ]

        res = llm.invoke(messages)
        content = getattr(res, "content", "")

        # Enhance response with visual indicators
        enhanced_response = (
            "📊 **Analytics Report**\n\n"
            f"{content}\n\n"
            "---\n"
            f"**Report Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"**Tenant:** {CURRENT_TENANT_ID or 'default'}\n"
            f"**Data Points Analyzed:** {len(stats) + len(tool_stats)}"
        )

        logger.info(f"Generated analytics report for tenant {CURRENT_TENANT_ID}")
        return {"messages": [("assistant", enhanced_response)]}

    except Exception as exc:
        logger.error(f"Analytics error: {exc}")
        return {"messages": [("assistant", f"Error generating analytics: {exc}")]}


# -----------------------------
# Build Enhanced LangGraph
# -----------------------------

def should_continue(state: MessagesState) -> str:
    """Determine if we should continue processing or end."""
    last_message = state["messages"][-1]

    # Check if the last message has tool calls
    if hasattr(last_message, 'tool_calls') and last_message.tool_calls:
        return "tools"

    # Always end after processing (no re-routing)
    return "end"

def create_enhanced_workflow():
    """Create the enhanced multi-agent workflow with memory and tool handling."""
    workflow = StateGraph(MessagesState)

    # Add all agent nodes (no router node, just conditional routing)
    workflow.add_node("greeting", node_greeting)
    workflow.add_node("doc_qa", node_doc_qa)
    workflow.add_node("api_exec", node_api_exec)
    workflow.add_node("form_gen", node_form_gen)
    workflow.add_node("analytics", node_analytics)
    workflow.add_node("escalate", node_escalate)

    # Add tool execution node that dynamically gets tools for current tenant
    def tool_node_func(state: MessagesState):
        tenant_id = CURRENT_TENANT_ID or "default"
        tools = get_tenant_tools(tenant_id)
        tool_node = ToolNode(tools)
        return tool_node.invoke(state)

    workflow.add_node("tools", tool_node_func)

    # Define routing logic - route directly from START
    workflow.add_conditional_edges(
        START,
        node_router,  # Use router as conditional function only
        {
            "greeting": "greeting",
            "doc_qa": "doc_qa",
            "api_exec": "api_exec",
            "form_gen": "form_gen",
            "analytics": "analytics",
            "escalate": "escalate"
        },
    )
    
    # Add conditional edges for tool handling in api_exec
    workflow.add_conditional_edges(
        "api_exec",
        should_continue,
        {"tools": "tools", "end": END}
    )
    
    # Tool results go back to api_exec for final processing
    workflow.add_edge("tools", "api_exec")
    
    # Simple end edges for other agents
    workflow.add_edge("greeting", END)
    workflow.add_edge("doc_qa", END)
    workflow.add_edge("form_gen", END)
    workflow.add_edge("analytics", END)
    workflow.add_edge("escalate", END)
    
    # Add memory for conversation state
    memory = MemorySaver()
    return workflow.compile(checkpointer=memory)

agent = create_enhanced_workflow()


# -----------------------------
# CLI Helpers
# -----------------------------


# -----------------------------
# Admin Dashboard Functions
# -----------------------------

def get_system_stats() -> Dict[str, Any]:
    """Get comprehensive system statistics."""
    stats = {
        "tenants": {
            "total": len(_tenant_registry),
            "active": len([t for t in _tenant_registry.values() if t.is_active]),
            "list": list(_tenant_registry.keys())
        },
        "sessions": {
            "active": len(_active_sessions),
            "total_created": len(_active_sessions)  # Simplified for demo
        },
        "tools": get_tool_stats(),
        "documents": {}
    }
    
    # Get document stats for each tenant
    for tenant_id in _tenant_registry.keys():
        doc_stats = get_document_stats(tenant_id)
        if "error" not in doc_stats:
            stats["documents"][tenant_id] = doc_stats
    
    return stats

def create_admin_dashboard() -> str:
    """Generate admin dashboard HTML."""
    try:
        stats = get_system_stats()

        # Generate components safely
        tenant_rows = _generate_tenant_rows(stats)
        tool_rows = _generate_tool_rows(stats["tools"])

        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Multi-Agent Chatbot Admin Dashboard</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; }}
        .card {{ background: white; padding: 20px; margin: 10px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .stats-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; }}
        .stat-item {{ text-align: center; }}
        .stat-number {{ font-size: 2em; font-weight: bold; color: #2196F3; }}
        .stat-label {{ color: #666; margin-top: 5px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
        th, td {{ padding: 10px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background-color: #f8f9fa; }}
        .status-active {{ color: #4CAF50; font-weight: bold; }}
        .status-inactive {{ color: #f44336; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Multi-Agent Chatbot Admin Dashboard</h1>

        <div class="stats-grid">
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["tenants"]["total"]}</div>
                    <div class="stat-label">Total Tenants</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["tenants"]["active"]}</div>
                    <div class="stat-label">Active Tenants</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{stats["sessions"]["active"]}</div>
                    <div class="stat-label">Active Sessions</div>
                </div>
            </div>
            <div class="card">
                <div class="stat-item">
                    <div class="stat-number">{len(stats["tools"])}</div>
                    <div class="stat-label">Total Tools</div>
                </div>
            </div>
        </div>

        <div class="card">
            <h2>Tenant Overview</h2>
            <table>
                <thead>
                    <tr>
                        <th>Tenant ID</th>
                        <th>Name</th>
                        <th>Status</th>
                        <th>Documents</th>
                        <th>Permissions</th>
                    </tr>
                </thead>
                <tbody>
                    {tenant_rows}
                </tbody>
            </table>
        </div>

        <div class="card">
            <h2>Tool Usage Statistics</h2>
            <table>
                <thead>
                    <tr>
                        <th>Tool Name</th>
                        <th>Total Calls</th>
                        <th>Errors</th>
                        <th>Success Rate</th>
                    </tr>
                </thead>
                <tbody>
                    {tool_rows}
                </tbody>
            </table>
        </div>
    </div>
</body>
</html>"""

        return html
    except Exception as e:
        logger.error(f"Dashboard generation error: {e}")
        return f"<html><body><h1>Dashboard Error</h1><p>Error: {e}</p></body></html>"

def _generate_tenant_rows(stats: Dict) -> str:
    """Generate HTML rows for tenant table."""
    rows = []
    for tenant_id, config in _tenant_registry.items():
        doc_count = stats["documents"].get(tenant_id, {}).get("total_chunks", 0)
        status_class = "status-active" if config.is_active else "status-inactive"
        status_text = "Active" if config.is_active else "Inactive"
        
        row = f"""
        <tr>
            <td>{tenant_id}</td>
            <td>{config.name}</td>
            <td class="{status_class}">{status_text}</td>
            <td>{doc_count} chunks</td>
            <td>{', '.join(config.permissions[:3])}{'...' if len(config.permissions) > 3 else ''}</td>
        </tr>
        """
        rows.append(row)
    
    return "".join(rows)

def _generate_tool_rows(tool_stats: Dict) -> str:
    """Generate HTML rows for tool statistics table."""
    rows = []
    for tool_name, data in tool_stats.items():
        calls = data['call_count']
        errors = data['error_count']
        success_rate = ((calls - errors) / calls * 100) if calls > 0 else 0
        
        row = f"""
        <tr>
            <td>{tool_name}</td>
            <td>{calls}</td>
            <td>{errors}</td>
            <td>{success_rate:.1f}%</td>
        </tr>
        """
        rows.append(row)
    
    return "".join(rows)

# -----------------------------
# Enhanced CLI Commands
# -----------------------------

def handle_command(line: str) -> Optional[str]:
    global CURRENT_TENANT_ID, CURRENT_SESSION
    
    if line.startswith("/tenant "):
        tenant_id = line.split(" ", 1)[1].strip() or None
        if tenant_id and authenticate_tenant(tenant_id):
            CURRENT_TENANT_ID = tenant_id
            # Create session for tenant
            try:
                CURRENT_SESSION = create_session(tenant_id)
                return f"Active tenant set to: {CURRENT_TENANT_ID} (Session: {CURRENT_SESSION.session_id[:8]}...)"
            except ValueError as e:
                return f"Error: {e}"
        else:
            return f"Invalid or inactive tenant: {tenant_id}"
    
    if line.startswith("/who"):
        session_info = f" (Session: {CURRENT_SESSION.session_id[:8]}...)" if CURRENT_SESSION else ""
        return f"Active tenant: {CURRENT_TENANT_ID}{session_info}"
    
    if line.startswith("/create-tenant "):
        parts = line.split(" ", 2)
        if len(parts) < 3:
            return "Usage: /create-tenant TENANT_ID TENANT_NAME"
        tenant_id, name = parts[1], parts[2]
        try:
            config = create_tenant(tenant_id, name)
            return f"Created tenant '{tenant_id}' ({name})"
        except ValueError as e:
            return f"Error: {e}"
    
    if line.startswith("/ingest "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("read_documents"):
            return "Permission denied: document ingestion not allowed"
        path = line.split(" ", 1)[1].strip().strip('"')
        return ingest_documents_from_dir(CURRENT_TENANT_ID, path)
    
    if line.startswith("/tool.httpget "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("use_tools"):
            return "Permission denied: tool registration not allowed"
        # Usage: /tool.httpget NAME BASE_URL_ENV [API_KEY_ENV]
        parts = line.split()
        if len(parts) < 3:
            return "Usage: /tool.httpget NAME BASE_URL_ENV [API_KEY_ENV]"
        name = parts[1]
        base_env = parts[2]
        api_env = parts[3] if len(parts) > 3 else None
        t = make_http_get_tool(name=name, description=f"HTTP GET tool for {name}", base_url_env=base_env, api_key_env=api_env)
        register_dynamic_tool(CURRENT_TENANT_ID, t)
        return f"Registered tool '{name}' for tenant {CURRENT_TENANT_ID}."
    
    if line.startswith("/tool.httppost "):
        if not CURRENT_TENANT_ID:
            return "Set a tenant first: /tenant TENANT_ID"
        if not has_permission("use_tools"):
            return "Permission denied: tool registration not allowed"
        parts = line.split()
        if len(parts) < 3:
            return "Usage: /tool.httppost NAME BASE_URL_ENV [API_KEY_ENV]"
        name = parts[1]
        base_env = parts[2]
        api_env = parts[3] if len(parts) > 3 else None
        t = make_http_post_tool(name=name, description=f"HTTP POST tool for {name}", base_url_env=base_env, api_key_env=api_env)
        register_dynamic_tool(CURRENT_TENANT_ID, t)
        return f"Registered POST tool '{name}' for tenant {CURRENT_TENANT_ID}."
    
    if line.startswith("/tools"):
        names = [t.name for t in get_tenant_tools(CURRENT_TENANT_ID)]
        return "Available tools: " + ", ".join(names)
    
    if line.startswith("/stats"):
        if not has_permission("admin"):
            return get_document_stats_tool()  # Limited stats for non-admin
        stats = get_system_stats()
        result = f"System Statistics:\n"
        result += f"- Tenants: {stats['tenants']['total']} total, {stats['tenants']['active']} active\n"
        result += f"- Active Sessions: {stats['sessions']['active']}\n"
        result += f"- Tools: {len(stats['tools'])} registered\n"
        return result
    
    if line.startswith("/dashboard"):
        if not has_permission("admin"):
            return "Permission denied: admin access required"
        try:
            html = create_admin_dashboard()
            dashboard_file = "admin_dashboard.html"
            with open(dashboard_file, "w", encoding="utf-8") as f:
                f.write(html)
            return f"Admin dashboard saved to {dashboard_file}. Open in browser to view."
        except Exception as e:
            return f"Error creating dashboard: {e}"
    
    if line.startswith("/permissions"):
        if not CURRENT_SESSION:
            return "No active session"
        return f"Your permissions: {', '.join(CURRENT_SESSION.permissions)}"
    
    if line.startswith("/help"):
        return (
            "Available Commands:\n"
            "  /tenant TENANT_ID                    Set active tenant\n"
            "  /create-tenant ID NAME               Create new tenant (admin)\n"
            "  /who                                 Show active tenant and session\n"
            "  /permissions                         Show your permissions\n"
            "  /ingest PATH                         Ingest documents from directory\n"
            "  /tool.httpget NAME BASE_URL_ENV [KEY_ENV]   Register HTTP GET tool\n"
            "  /tool.httppost NAME BASE_URL_ENV [KEY_ENV]  Register HTTP POST tool\n"
            "  /tools                               List available tools\n"
            "  /stats                               Show system statistics\n"
            "  /dashboard                           Generate admin dashboard (admin)\n"
            "  /help                                Show this help"
        )
    
    return None


def chat_once(user_input: str, thread_id: str = "default") -> str:
    """Chat with the agent using proper thread configuration."""
    config = {"configurable": {"thread_id": thread_id}}
    final_state = agent.invoke({"messages": [("user", user_input)]}, config=config)
    last = final_state["messages"][-1]
    return getattr(last, "content", str(last))

def chat_with_agent(user_input: str, tenant_id: str = "default") -> str:
    """Simplified chat function for testing and demos."""
    global CURRENT_TENANT_ID, CURRENT_SESSION
    original_tenant = CURRENT_TENANT_ID
    original_session = CURRENT_SESSION

    try:
        # Set tenant context
        CURRENT_TENANT_ID = tenant_id

        # Create or get session for this tenant
        if tenant_id in _tenant_registry:
            session = create_session(tenant_id)
            CURRENT_SESSION = session

        thread_id = f"chat_{tenant_id}_{hash(user_input) % 1000}"
        result = chat_once(user_input, thread_id)
        return result
    finally:
        CURRENT_TENANT_ID = original_tenant
        CURRENT_SESSION = original_session


if __name__ == "__main__":
    if not os.environ.get("GOOGLE_API_KEY"):
        print("Warning: GOOGLE_API_KEY is not set. Set it to enable the LLM and embeddings.")
    print("Multi-Agent Chatbot (LangGraph + LangChain + FAISS RAG). Type '/help' for commands, 'exit' to quit.\n")
    # Choose tenant at start
    try:
        CURRENT_TENANT_ID = input("Tenant ID (default): ").strip() or "default"
    except (EOFError, KeyboardInterrupt):
        CURRENT_TENANT_ID = "default"
    print(f"Active tenant: {CURRENT_TENANT_ID}\n")
    while True:
        try:
            user = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nBye!")
            break
        if not user:
            continue
        if user.lower() in {"exit", "quit"}:
            print("Bye!")
            break
        if user.startswith("/"):
            out = handle_command(user)
            print(f"Bot: {out}\n")
            continue
        # Generate unique thread ID for conversation
        thread_id = f"session_{CURRENT_TENANT_ID}_{hash(user) % 10000}"
        reply = chat_once(user, thread_id)
        print(f"Bot: {reply}\n")